from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = Path(r"C:\Users\Arslan\Desktop\yjd.docx")
MAIN_DOC = ROOT / "kakabalowa.docx"
OUT_DOC = ROOT / "ykdysady_iot_security_yjd_takyk.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_heading_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True


def replace_by_index(doc: Document, replacements: dict[int, str]) -> None:
    for index, text in replacements.items():
        if 0 <= index < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[index], text)


def build_exact_copy() -> Document:
    doc = Document(str(SOURCE))

    title = doc.paragraphs[0].insert_paragraph_before(
        "Taslamanyň ykdysady netijeliliginiň hasaplamasy"
    )
    try:
        title.style = "Heading 1"
    except KeyError:
        pass
    set_heading_font(title)

    # Original yjd.docx paragraphs are shifted by +1 because of the inserted section title.
    replacements = {
        108: "Programma üpjünçiligi taslamasynyň ykdysady – tehniki esaslandyrylyşy.",
        109: (
            "Programma üpjünçiligi taslamasynyň ykdysady – tehniki esaslandyrylyşy (YTE) – "
            "IoT Security Scanner programma toplumyny işläp taýýarlamak we amaly taýdan "
            "ulanmak üçin esasy taslama resminamasy bolup çykyş edýär. YTE-de bellenilen "
            "düzgünleriň esasynda programma arhitekturasy, ulanylýan tehnologiýalar, "
            "smeta hasaplamalary we synag işleri kesgitlenilýär."
        ),
        110: (
            "YTE-da programma toplumynyň funksional çözgütleri, tor skanirleme algoritmleri, "
            "maglumatlar bazasy, grafiki interfeýs, hasabat taýýarlamak ulgamy, zähmet we "
            "ekologik howpsuzlyk meseleleri kesgitlenýär, şeýle-de onuň ykdysady netijeliligi "
            "we durmuş-ykdysady peýdasy bahalandyrylýar."
        ),
        111: (
            "Programma üpjünçiligini işläp taýýarlamak üçin düýpli maýa goýum hökmünde "
            "bolup biljek maliýe çeşmeleri hökmünde şular çykyş edip biler:"
        ),
        117: "Programma taslamasynyň YTE-si şu bölümlerden düzülýär:",
        118: "- umumy düşündiriş ýazgysy;",
        119: "- predmet oblastynyň seljerilişi we tor howpsuzlygy boýunça maglumatlar;",
        120: "- programma arhitekturasy we tehnologiýa çözgütleri;",
        121: "- programma toplumyny dolandyrmak, ulanyjy interfeýsi we iş şertlerini goramak;",
        122: "- maglumatlar bazasy, hasabat we bildiriş ulgamlary;",
        123: "- synag işleri, skanirleme netijeleri we netijelilik bahalandyrmasy;",
        124: "- daşarky gurşawy goramak we kagyz serişdelerini tygşytlamak.",
        125: "- ykdysady netijeliligiň hasaplanylyşy.",
        127: (
            "IoT Security Scanner programma toplumyny işläp taýýarlamakda işgärleriň "
            "sanynyň hasaplanylyşy"
        ),
        128: (
            "Programma üpjünçiligi taslamasynyň ykdysady netijeliligini kesgitlemek üçin "
            "ilki bilen onda işlejek hünärmenleriň sanyny we kabul edilen normalaryň "
            "esasynda programma işlerine sarp ediljek wagty kesgitlemeli."
        ),
        129: (
            "Ilki işgär personalynyň iş tertibini we zähmet öndürijiliginiň ösüşini bilip "
            "işçileriň sanyny şu aşakdaky formula bilen kesgitlemeli"
        ),
        131: "Tdb – programma toplumyny işläp taýýarlamagyň zähmet talap edişi",
        138: "Programma toplumyny işläp taýýarlamagyň bahasyny hasaplap çykarmak.",
        140: "a) Tehniki serişdeler we programma üpjünçiligi",
        141: "b) Internet, aragatnaşyk we energiýa",
        142: "ç) Işçileriň esasy we goşmaça zähmet haky",
        143: "d) Sosial gorag üçin geçirimler",
        144: (
            "Enjamlary, kompýuterleri we tor serişdelerini ulanmak we saklamak üçin edilýän "
            "harajatlara amortizasiýa, gündelik hyzmat etmek hem-de käbir dürli harajatlar "
            "degişlidir."
        ),
        146: (
            "Bölüm boýunça harajatlar programma taslamasyny dolandyrmak işlerini geçirýän "
            "adamlary saklamak, iş ýerini üpjün etmek, kompýuter enjamlarynyň amortizasiýasy "
            "we beýleki harajatlar bilen baglanyşyklydyr."
        ),
        147: (
            "Taslama smetasynyň esasynda IoT Security Scanner programma toplumyny döretmek "
            "üçin gerek boljak tehniki serişdeleriň, programma gurallarynyň we goşmaça "
            "serişdeleriň sanawy kesgitlenip, olaryň bahasy bazar nyrhlary boýunça alynýar."
        ),
        151: (
            "Programma toplumyny işläp taýýarlamaga gatnaşýan hünärmenleriň kwalifikasiýa "
            "derejesini we bir sagada tölegini kesgitläp, zähmet hakyny hasaplamaly."
        ),
        152: (
            "Zähmet haky esasy we goşmaça zähmet haka bölünýär. Esasy zähmet haky tarif "
            "boýunça tölegden we baýrak fondundan ybaratdyr. Goşmaça zähmet haky bolsa "
            "rugsady wagtyndaky tölegden, jemgyýetçilik işlerine gatnaşýanlara zähmet "
            "hakyndan we ş.m. ybaratdyr."
        ),
        156: "Td.b – programma üpjünçiligini işläp taýýarlamagyň zähmet talap edişi",
        186: "Netijede IoT Security Scanner programma taslamasy boýunça ähli harajatlary bir tablisa jemlemeli",
    }
    replace_by_index(doc, replacements)

    # Keep the same tables, formulas, and numeric scheme, but rename the concrete project items.
    table0 = doc.tables[0]
    table0.cell(1, 1).text = "Programmist"
    table0.cell(2, 1).text = "Dizaýner programmist"
    table0.cell(3, 1).text = "Ulgam dolandyryjy"

    table1 = doc.tables[1]
    table1.cell(1, 1).text = "Kompýuter"
    table1.cell(2, 1).text = "Python, PyQt6 we Nmap gurallary"
    table1.cell(2, 2).text = "Toplum"
    table1.cell(3, 1).text = "Windows OU"
    table1.cell(4, 1).text = "SQLite we PDF hasabat moduly"
    table1.cell(4, 2).text = "Toplum"

    table2 = doc.tables[2]
    table2.cell(1, 1).text = "Materiallar, tehniki we programma serişdeleri"
    table2.cell(8, 1).text = "Enjamlara we programma gurşawyna çykdajylar"

    return doc


def backup_main() -> Path:
    backup = MAIN_DOC.with_name(
        f"{MAIN_DOC.stem}_before_ykdysady_yjd_exact_{datetime.now().strftime('%Y%m%d_%H%M%S')}{MAIN_DOC.suffix}"
    )
    shutil.copy2(MAIN_DOC, backup)
    return backup


def replace_main_section(section_doc: Document) -> Path:
    backup = backup_main()
    main = Document(str(MAIN_DOC))
    body = main._body._element

    start_idx = None
    end_idx = None

    candidates = []
    for paragraph in main.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            candidates.append(paragraph)

    if candidates:
        start_idx = body.index(candidates[-1]._element)

    if start_idx is not None:
        for paragraph in main.paragraphs:
            element_idx = body.index(paragraph._element)
            if element_idx <= start_idx:
                continue
            text = " ".join(paragraph.text.split())
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if text == "Netije" and style_name.startswith("Heading"):
                end_idx = element_idx
                break

    if start_idx is None or end_idx is None:
        raise RuntimeError("Esasy faýlda ykdysady bölüm ýa-da global Netije tapylmady.")

    elements = list(body)
    for element in elements[start_idx:end_idx]:
        body.remove(element)

    netije_element = list(body)[start_idx]
    source_body = section_doc._body._element
    for child in list(source_body):
        if child.tag.endswith("sectPr"):
            continue
        body.insert(body.index(netije_element), copy.deepcopy(child))

    main.save(str(MAIN_DOC))
    normalize_main_heading_styles()
    return backup


def normalize_main_heading_styles() -> None:
    main = Document(str(MAIN_DOC))
    start = None
    for idx, paragraph in enumerate(main.paragraphs):
        text = " ".join(paragraph.text.split())
        if idx > 300 and text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            start = idx
            break
    if start is None:
        return
    heading_offsets = {
        0: "Heading 1",
        1: "Heading 1",
        15: "Heading 2",
        30: "Heading 2",
        59: "Heading 2",
        92: "Heading 2",
        108: "Heading 2",
        127: "Heading 2",
        138: "Heading 2",
    }
    for offset, style_name in heading_offsets.items():
        paragraph_index = start + offset
        if 0 <= paragraph_index < len(main.paragraphs):
            try:
                main.paragraphs[paragraph_index].style = style_name
            except KeyError:
                pass
    main.save(str(MAIN_DOC))


def main() -> None:
    doc = build_exact_copy()
    doc.save(str(OUT_DOC))
    backup = replace_main_section(doc)
    print(f"Created: {OUT_DOC}")
    print(f"Updated: {MAIN_DOC}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
