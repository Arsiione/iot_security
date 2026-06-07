from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_1_5.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("• " + text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
                    if row_index == 0:
                        run.bold = True


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "1.5. Diplom işiniň maksady we wezipeleri")

    add_normal(
        doc,
        "Geçirilen nazary seljermäniň netijesinde IoT gurluşlarynyň häzirki zaman lokal torlarynda "
        "giňden ulanylýandygy, emma olaryň howpsuzlygynyň köplenç ýeterlik derejede üpjün edilmeýändigi "
        "anyklanyldy. Standart parollar, açyk portlar, köne firmware, goragsyz protokollar we tordaky "
        "gurluşlaryň doly hasaba alynmazlygy administrator üçin möhüm töwekgelçilik döredýär. Şu sebäpli "
        "diplom işiniň esasy ugry lokal torlarda IoT gurluşlaryny awtomatik ýüze çykarmak, olaryň "
        "açyk hyzmatlaryny seljermek we ulanyja howpsuzlyk boýunça düşnükli maslahat bermek bilen "
        "baglanyşyklydyr."
    )
    add_normal(
        doc,
        "Diplom işiniň maksady — lokal Wi‑Fi ýa-da Ethernet torunda ýerleşýän IoT gurluşlaryny we beýleki "
        "işjeň enjamlary awtomatik tapmaga, olaryň IP we MAC maglumatlaryny görkezmäge, açyk portlaryny "
        "barlamaga, töwekgelçilik derejesini kesgitlemäge, netijeleri maglumatlar bazasynda saklamaga we "
        "PDF görnüşinde hasabat taýýarlamaga mümkinçilik berýän IoT Security Scanner programma toplumyny "
        "taslamak we işläp taýýarlamakdyr."
    )
    add_normal(
        doc,
        "Işiň obýekti hökmünde IoT gurluşlary ulanylýan lokal kompýuter torlary kabul edilýär. Şeýle "
        "torlara öý Wi‑Fi ulgamlary, okuw edaralarynyň lokal segmentleri, kiçi kärhanalaryň router, kamera, "
        "printer, telefon, akylly öý enjamlary we beýleki tor gurluşlary girip biler."
    )
    add_normal(
        doc,
        "Işiň predmeti hökmünde lokal tordaky gurluşlary ýüze çykarmagyň, portlary skanirlemegiň, "
        "hyzmatlaryň howpsuzlyk töwekgelçiligini bahalandyrmagyň we ulanyja remediation, ýagny "
        "gowşaklyklary azaltmak boýunça maslahat bermegiň programma usullary öwrenilýär."
    )
    add_normal(
        doc,
        "Diplom işiniň maksadyna ýetmek üçin aşakdaky esasy wezipeler kesgitlenildi:"
    )

    bullets = [
        "IoT gurluşlarynyň görnüşlerini, olaryň lokal torlarda ulanylyşyny we howpsuzlyk aýratynlyklaryny öwrenmek;",
        "IoT torlaryna degişli esasy hüjüm wektorlaryny, şol sanda Telnet, HTTP/HTTPS, RTSP, MQTT, FTP we SSH bilen bagly töwekgelçilikleri seljermek;",
        "Nmap, Shodan, Nessus we OpenVAS ýaly bar bolan skanirleme gurallarynyň mümkinçiliklerini deňeşdirmek;",
        "Lokal tor üçin amatly skanirleme algoritmini taslamak we Wi‑Fi/Ethernet adapterini dogry saýlamak mehanizmini işläp düzmek;",
        "Nmap ping scan, ARP barlagy, ICMP ping sweep, Windows ARP keşi, lokal kompýuter we şlýuz maglumatlaryny birleşdirýän gurluş tapmak ulgamyny döretmek;",
        "Tapylan gurluşlaryň portlaryny barlamak we açyk hyzmatlara görä töwekgelçilik derejesini kesgitlemek;",
        "IoT gurluşlary üçin plugin esasly gowşaklyk barlaglaryny, şol sanda Telnet gowşak autentifikasiýasy we Hikvision bilen bagly alamatlary barlamak mümkinçiligini taýýarlamak;",
        "PyQt6 tehnologiýasy esasynda türkmen dilindäki grafiki interfeýsi döretmek;",
        "SQLite maglumatlar bazasynda skanirleme taryhyny, tapylan gurluşlary we netijeleri saklamak;",
        "Netijeleri PDF hasabat görnüşinde eksport etmek we ulanyja anyk howpsuzlyk maslahatlaryny bermek;",
        "Programmany test torunda barlap, onuň işjeň gurluşlary tapmak, portlary görkezmek we hasabat taýýarlamak mümkinçiliklerine baha bermek.",
    ]
    for item in bullets:
        add_bullet(doc, item)

    add_normal(
        doc,
        "Işiň ylmy-amaly ähmiýeti IoT howpsuzlyk auditiniň diňe çuňňur hünärmen gurallary arkaly däl, "
        "eýsem düşnükli grafiki programma arkaly hem geçirilip bilinýändigini görkezmekden ybaratdyr. "
        "Taslama skanirleme netijelerini ýönekeý ulanyjy üçin düşnükli görnüşde berýär: gurluşyň IP salgysy, "
        "ady, MAC salgysy, öndürijisi, görnüşi, tapylan usuly, açyk portlary, gowşaklyk ýagdaýy we maslahat "
        "bir tablisa içinde görkezilýär."
    )
    add_normal(
        doc,
        "Programma toplumynyň amaly ähmiýeti onuň diplom goragy wagtynda hakyky lokal tor bilen görkezilip "
        "bilinýändigindedir. Mysal üçin, programma Wi‑Fi toruny saýlap, 192.168.1.0/24 ýaly IP aralygynda "
        "routeri, kompýuteri, telefony ýa-da IoT kandidat gurluşyny tapýar. Eger enjamda açyk port bolmasa, "
        "ol netijelerden aýrylmaýar, pes töwekgelçilik derejesi bilen görkezilýär. Eger Telnet, HTTP, RTSP "
        "ýa-da beýleki hyzmatlar açyk bolsa, programma degişli töwekgelçiligi we maslahatlary görkezýär."
    )

    caption = doc.add_paragraph("Tablisa 1.5 - Diplom işiniň esasy wezipeleri we garaşylýan netijeler")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    headers = ["Wezipe", "Ýerine ýetiriş mazmuny", "Garaşylýan netije"]
    rows = [
        [
            "Predmet oblastyny seljermek",
            "IoT gurluşlary, howplar we skanirleme usullary öwrenilýär",
            "Nazary esaslandyrma döredilýär",
        ],
        [
            "Analoglary deňeşdirmek",
            "Nmap, Shodan, Nessus, OpenVAS mümkinçilikleri seljerilýär",
            "Taslamanyň tapawutly tarapy kesgitlenýär",
        ],
        [
            "Skanirleme algoritmini taslamak",
            "IP aralygy, adapter saýlamak, ARP/Nmap/Ping usullary birleşdirilýär",
            "Lokal tordaky gurluşlar has ygtybarly tapylýar",
        ],
        [
            "Programma ýadrosyny işläp düzmek",
            "Port barlagy, töwekgelçilik derejesi we plugin ulgamy taýýarlanýar",
            "IoT howpsuzlyk auditi awtomatlaşdyrylýar",
        ],
        [
            "Ulanyjy interfeýsini döretmek",
            "PyQt6 bilen türkmen dilindäki sahypalar taýýarlanýar",
            "Programma diplom goragynda düşnükli görkezilýär",
        ],
        [
            "Hasabat we taryh ulgamyny döretmek",
            "SQLite bazasy we PDF eksporty ulanylýar",
            "Skanirleme netijeleri saklanýar we resmileşdirilýär",
        ],
        [
            "Synag geçirmek",
            "Programma hakyky ýa-da test lokal torda barlanýar",
            "Netijelilik we işleýiş ukyby bahalandyrylýar",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)

    add_normal(
        doc,
        "Şeýlelikde, diplom işiniň birinji bölüminde IoT gurluşlarynyň aýratynlyklary, olaryň howpsuzlyk "
        "meseleleri, hüjüm wektorlary we bar bolan skanirleme ulgamlary seljerildi. Bu seljerme işiň "
        "indiki bölümlerinde programma toplumynyň arhitekturasyny, maglumatlar bazasyny, skanirleme "
        "algoritmini we ulanyjy interfeýsini taslamak üçin nazary binýat bolup hyzmat edýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_style(main_doc):
    for paragraph in main_doc.paragraphs:
        if paragraph.text.strip().startswith("1.5."):
            paragraph.style = main_doc.styles["Heading 2"]
            set_font(paragraph, 14, True)
            break


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("1.5.") for p in main_doc.paragraphs):
        print("Section 1.5 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_style(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
