from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path(r"C:\Users\Arslan\Desktop\yjd.docx")
MAIN_DOC = ROOT / "kakabalowa.docx"
OUT_DOC = ROOT / "ykdysady_iot_security.docx"


def money(value: float) -> str:
    return f"{value:.2f}".replace(".", ",")


MATERIALS = 3630.00
TRANSPORT = MATERIALS * 0.10
HOURS = 250
AVG_RATE = 18.33
Z_TAR = 4582.50
Z_BAYR = 458.25
Z_ES = 5040.75
Z_GOSM = 504.08
Z_UM = 5544.83
S_SOC = 1108.97
C_SEH = 756.11
C_KARH = 1512.23
C_ENJAM = 2268.34
TOTAL_COST = 15183.48
YEAR_EFFECT = 6000.00
PAYBACK = 2.53


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_font(paragraph, size: int = 12, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def add_p(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph, bold=bold)


def add_formula(doc: Document, text: str) -> None:
    style_names = {style.name for style in doc.styles}
    if "Preformat" in style_names:
        paragraph = doc.add_paragraph(text, style="Preformat")
    else:
        paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph, bold=False)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style_name = f"Heading {level}"
    style_names = {style.name for style in doc.styles}
    if style_name in style_names:
        paragraph = doc.add_paragraph(text, style=style_name)
    else:
        paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14 if level == 1 else 13)
        run.bold = True


def fill_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                set_font(paragraph, size=11, bold=(r_idx == 0))


def build_ykdysady_doc(template_path: Path = TEMPLATE) -> Document:
    doc = Document(str(template_path))
    clear_document(doc)

    add_heading(doc, "Taslamanyň ykdysady netijeliliginiň hasaplamasy", 1)

    add_heading(doc, "Maýa goýum taslamalarynyň netijeliliginiň esasy ýörelgeleri", 2)
    add_p(
        doc,
        "Maýa goýum taslamalarynyň netijeliligine baha bermekde taslamanyň maksady, "
        "ony durmuşa geçirmek üçin sarp edilýän serişdeler we taslamadan alynýan ykdysady "
        "netije bilelikde seljerilýär. Bu diplom işinde ykdysady bahalandyrma IoT Security "
        "Scanner programma toplumyny döretmäge, synagdan geçirmäge we ony lokal torlaryň "
        "başlangyç howpsuzlyk auditi üçin ulanmagyň peýdasyny kesgitlemäge gönükdirilýär.",
    )
    add_p(
        doc,
        "Taslamanyň netijeliligi programmanyň diňe tehniki taýdan işlemegi bilen däl, eýsem "
        "onuň hünärmeniň wagtyny tygşytlamagy, tor barlagyny standartlaşdyrmagy, hasabat "
        "taýýarlamagy çaltlandyrmagy we howpsuzlyk töwekgelçiliklerini ir ýüze çykarmagy "
        "bilen hem kesgitlenýär.",
    )
    add_p(
        doc,
        "Bütewi taslamanyň netijeliligi programma toplumyny işläp taýýarlamak üçin zerur "
        "bolan çykdajylary we ony ulanmak arkaly ýylda gazanylýan şertli ykdysady netijäni "
        "deňeşdirmek arkaly bahalandyrylýar. Şeýle çemeleşme taslamanyň okuw edarasynda, "
        "kiçi kärhanada ýa-da ofis torunda ulanmak üçin maksadalaýykdygyny görkezýär.",
    )
    add_p(doc, "Netijeliligi anyklamakda şu ýörelgeler göz öňünde tutulýar:")
    for item in [
        "taslamanyň ähli tapgyrlarynda çykdajylaryň hasaba alynmagy;",
        "iş wagtynyň tygşytlanyşynyň ykdysady netije hökmünde kabul edilmegi;",
        "el bilen geçirilýän audit bilen programma arkaly geçirilýän auditiň deňeşdirilmegi;",
        "programma üpjünçiliginiň açyk çeşmeli tehnologiýalara daýanmagy sebäpli lisenziýa çykdajylarynyň azaldylmagy;",
        "mümkin bolan töwekgelçilikleriň we çäklendirmeleriň hasaplamada görkezilmegi.",
    ]:
        add_p(doc, item)

    add_heading(doc, "Taslamanyň netijeliligini hasaplamak üçin başlangyç maglumatlar", 2)
    add_p(
        doc,
        "Başlangyç maglumatlar hökmünde programma toplumynyň işlenip taýýarlanylyşyna gatnaşýan "
        "hünärmenleriň sany, sarp edilýän iş wagty, ulanylýan tehniki serişdeler, programma "
        "üçin gerek bolan goşmaça serişdeler we lokal tor auditi geçirilende tygşytlanýan wagt "
        "alynýar.",
    )
    add_p(
        doc,
        "IoT Security Scanner Python programmirleme dili, PyQt6 grafiki interfeýs kitaphanasy, "
        "SQLite maglumat bazasy, Nmap tor skanirleme guraly, ReportLab we Matplotlib hasabat "
        "serişdeleri arkaly döredilýär. Bu tehnologiýalaryň köpüsi açyk çeşmeli bolandygy üçin "
        "taslamanyň programma üpjünçiligi boýunça çykdajylary peselýär.",
    )
    add_p(
        doc,
        "Hasaplamalarda programma toplumynyň bir lokal /24 toruny el bilen barlamak üçin ortaça "
        "6 sagat, programma arkaly barlamak we hasabat taýýarlamak üçin bolsa 1 sagat gerek "
        "diýlip kabul edildi. Bir ýylda 60 gezek audit geçirilende tygşytlanýan wagt ykdysady "
        "netijäniň esasy çeşmesi bolup durýar.",
    )

    add_heading(doc, "Kompauting we diskontirleme prosesleriniň düýp manysy", 2)
    add_p(
        doc,
        "Maýa goýum taslamalarynda pul serişdeleriniň wagt boýunça gymmatynyň üýtgemegi hem "
        "hasaba alynmalydyr. Şonuň üçin geljekki girdejini häzirki gymmata getirmek ýa-da "
        "häzirki çykdajynyň geljekki bahasyny kesgitlemek üçin kompauting we diskontirleme "
        "formulalary ulanylýar.",
    )
    add_formula(doc, "BC = HC(1 + K)^n")
    add_p(
        doc,
        "Bu ýerde BC – puluň geljekki gymmaty; HC – goýlan puluň häzirki gymmaty; "
        "K – göterim koeffisiýenti; n – ýyl sany. Bu formula goýlan serişdeleriň geljekki "
        "gymmatyny kesgitlemekde ulanylýar.",
    )
    add_formula(doc, "HC = BC / (1 + K)^n")
    add_p(
        doc,
        "Bu ýerde HC – geljekde alynjak girdejiniň häzirki gymmaty. Programma taslamalarynda "
        "bu usul çykdajylaryň we garaşylýan ykdysady netijäniň wagt boýunça deňeşdirilmegini "
        "esaslandyrmaga mümkinçilik berýär.",
    )

    add_heading(doc, "Programma toplumynyň ykdysady-tehniki esaslandyrylyşy", 2)
    add_p(
        doc,
        "IoT Security Scanner programma toplumy lokal tordaky kompýuterleri, marşrutizatorlary, "
        "telefonlary we IoT kandidat gurluşlaryny tapmaga, olaryň açyk portlaryny kesgitlemäge, "
        "töwekgelçilik derejesini görkezmäge we PDF hasabatyny taýýarlamaga niýetlenendir. "
        "Şonuň üçin taslamanyň ykdysady netijesi esasan howpsuzlyk auditiniň tizleşmegi we "
        "hünärmeniň gaýtalanýan işleriniň awtomatlaşdyrylmagy bilen baglanyşyklydyr.",
    )
    add_p(
        doc,
        "Taslama boýunça çykdajylar işgärleriň zähmet hakyndan, tehniki serişdelerden, "
        "transport we aragatnaşyk harajatlaryndan, sosial geçirimlerden, şeýle hem enjamlary "
        "ulanmak we saklamak bilen bagly çykdajylardan ybaratdyr.",
    )

    add_heading(doc, "Işgärleriň sanynyň we zähmet talap edişiniň hasaplanylyşy", 2)
    add_p(
        doc,
        "Programma toplumyny işläp taýýarlamak üçin programmist, interfeýs dizaýneri we tor "
        "howpsuzlygy boýunça testçi gatnaşýar. Işleriň umumy zähmet talap edişi aşakdaky "
        "formula boýunça kesgitlenilýär:",
    )
    add_formula(doc, "Tiş = Tseljerme + TUI + Tprogram + Tsynag + Tresminama")
    add_formula(doc, "Tiş = 40 + 40 + 120 + 30 + 20 = 250 sagat")

    fill_table(
        doc,
        [
            ["T№", "Işçileriň kategoriýasy", "Sany", "Razrýad", "Iş şerti", "1 sag. töleg.", "Baýrak %"],
            ["1", "Programmist", "1", "-", "Kadaly", "20", "10"],
            ["2", "UI/UX dizaýner", "1", "-", "Kadaly", "18", "10"],
            ["3", "Tor howpsuzlygy boýunça testçi", "1", "-", "Kadaly", "17", "10"],
            ["4", "", "", "", "", "", ""],
            ["5", "", "", "", "", "", ""],
            ["6", "", "", "", "", "", ""],
            ["7", "", "", "", "", "", ""],
            ["8", "", "", "", "", "", ""],
            ["9", "", "", "", "", "", ""],
            ["10", "", "", "", "", "", ""],
            ["11", "", "", "", "", "", ""],
            ["", "Jemi:", "3", "-", "", "55", "10"],
        ],
    )

    add_p(
        doc,
        "Ortaça bir sagatlyk töleg işgärleriň sagatlaýyn tölegleriniň ortaça bahasy hökmünde "
        "kesgitlenilýär:",
    )
    add_formula(doc, f"Corta = (20 + 18 + 17) / 3 = {money(AVG_RATE)} manat/sag.")
    add_formula(doc, f"Ztar = Corta × Tiş = {money(AVG_RATE)} × {HOURS} = {money(Z_TAR)} manat")
    add_p(doc, "Baýrak fondy tarif boýunça zähmet hakynyň 10%-i möçberinde kabul edilýär.")
    add_formula(doc, f"Zbaýr = Ztar × 0,10 = {money(Z_BAYR)} manat")
    add_formula(doc, f"Zes = Ztar + Zbaýr = {money(Z_ES)} manat")
    add_p(doc, "Goşmaça zähmet haky esasy zähmet hakynyň 10%-i möçberinde alynýar.")
    add_formula(doc, f"Zgoşm = Zes × 0,10 = {money(Z_GOSM)} manat")
    add_formula(doc, f"Zum = Zes + Zgoşm = {money(Z_UM)} manat")
    add_p(doc, "Sosial gorag üçin geçirimler umumy zähmet hakynyň 20%-i möçberinde hasaplanylýar.")
    add_formula(doc, f"Ssos.gor = Zum × 0,20 = {money(S_SOC)} manat")

    add_heading(doc, "Tehniki serişdeleriň we programma üpjünçiliginiň harajatlary", 2)
    add_p(
        doc,
        "Programma toplumyny döretmekde esasy çykdajylar iş stansiýasy, test tor enjamlary, "
        "operasion ulgam we resminama taýýarlamak üçin gerek bolan serişdeler bilen baglanyşyklydyr. "
        "Python, PyQt6, SQLite, Nmap, ReportLab, Matplotlib we PyInstaller ýaly serişdeler açyk "
        "çeşmeli bolansoň, lisenziýa harajaty hökmünde 0 manat kabul edilýär.",
    )
    fill_table(
        doc,
        [
            ["T№", "Ätiýaçlyk şaýlaryň görnüşi", "Ölçeg birligi", "Sany", "1-niň bahasy", "Umumy gymmaty"],
            ["1", "Kompýuter ýa-da noutbuk", "San", "1", "3000", "3000"],
            ["2", "Router ýa-da test switch", "San", "1", "350", "350"],
            ["3", "Windows operasion ulgamy", "San", "1", "100", "100"],
            ["4", "Python, PyQt6, SQLite, Nmap we beýleki açyk çeşmeli serişdeler", "Toplum", "1", "0", "0"],
            ["5", "Resminama we PDF hasabat taýýarlamak serişdeleri", "Toplum", "1", "100", "100"],
            ["6", "Tor kabeli, USB göteriji we ownuk serişdeler", "Toplum", "1", "80", "80"],
            ["", "Jemi", "", "6", "", money(MATERIALS)],
        ],
    )
    add_p(doc, "Transport we üpjünçilik harajatlary materiallaryň umumy bahasynyň 10%-i möçberinde kabul edilýär.")
    add_formula(doc, f"Ttr = M × 0,10 = {money(MATERIALS)} × 0,10 = {money(TRANSPORT)} manat")

    add_heading(doc, "Goşmaça çykdajylaryň hasaplanylyşy", 2)
    add_p(
        doc,
        "Seh boýunça çykdajylar, kärhana boýunça çykdajylar we enjamlary ulanmak bilen bagly "
        "harajatlar kabul edilen normalar boýunça esasy zähmet hakyndan hasaplanylýar.",
    )
    add_formula(doc, f"Çseh = Zes × 0,15 = {money(C_SEH)} manat")
    add_formula(doc, f"Çkärh = Zes × 0,30 = {money(C_KARH)} manat")
    add_formula(doc, f"Çenjam = Zes × 0,45 = {money(C_ENJAM)} manat")

    add_p(doc, "Netijede ähli harajatlar aşakdaky tablisa jemlenilýär.")
    fill_table(
        doc,
        [
            ["T№", "Maddanyň görnüşleri", "Gymmaty (Manat)"],
            ["1", "Materiallar we tehniki serişdeler", money(MATERIALS)],
            ["2", "Transport harajatlary", money(TRANSPORT)],
            ["3", "Esasy zähmet haky", money(Z_ES)],
            ["4", "Goşmaça zähmet haky", money(Z_GOSM)],
            ["5", "Sosial gorag üçin geçirimler", money(S_SOC)],
            ["6", "Seh boýunça çykdajylar", money(C_SEH)],
            ["7", "Kärhana boýunça çykdajylar", money(C_KARH)],
            ["8", "Enjamlary ulanmak we saklamak üçin çykdajylar", money(C_ENJAM)],
            ["9", "Jemi harajatlar", money(TOTAL_COST)],
            ["10", "Ýyllyk ykdysady netije", money(YEAR_EFFECT)],
            ["11", "Özüni ödeýän wagty (ýyl)", money(PAYBACK)],
        ],
    )

    add_heading(doc, "Ýyllyk ykdysady netijäniň we özüni ödeýiş möhletiniň kesgitlenilişi", 2)
    add_p(
        doc,
        "Programma toplumynyň ykdysady peýdasy lokal toruň howpsuzlyk auditini geçirmek üçin "
        "sarp edilýän wagtyň azalmagy bilen kesgitlenýär. El bilen audit geçirilende IP salgylary, "
        "portlar, gurluş görnüşleri we maslahatlar aýratynlykda barlanýar. Programma arkaly bu "
        "işleriň köp bölegi awtomatlaşdyrylýar.",
    )
    add_formula(doc, "Eýyl = (Tel - Tprog) × Csag × Nýyl")
    add_formula(doc, f"Eýyl = (6 - 1) × 20 × 60 = {money(YEAR_EFFECT)} manat")
    add_p(
        doc,
        "Bu ýerde Tel – el bilen audit üçin gerek bolan wagt; Tprog – programma arkaly audit "
        "üçin gerek bolan wagt; Csag – hünärmeniň bir sagatlyk tölegi; Nýyl – bir ýylda "
        "geçirilýän auditleriň sany.",
    )
    add_formula(doc, "K")
    add_formula(doc, "Töz.öd. = -------------")
    add_formula(doc, "Eýyl")
    add_formula(doc, f"Töz.öd. = {money(TOTAL_COST)} / {money(YEAR_EFFECT)} = {money(PAYBACK)} ýyl")
    add_p(
        doc,
        "Hasaplamadan görnüşi ýaly, IoT Security Scanner programma toplumyny işläp taýýarlamak "
        "üçin sarp edilýän çykdajylar takmynan 2,5 ýylyň dowamynda özüni ödeýär. Programma has "
        "köp lokal torda ýa-da ýygy-ýygydan audit geçirilýän gurşawda ulanylsa, özüni ödeýiş "
        "möhleti has hem gysgalýar.",
    )
    add_heading(doc, "Ykdysady bölüm boýunça netije", 2)
    add_p(
        doc,
        "Geçirilen hasaplamalar IoT Security Scanner programma toplumynyň ykdysady taýdan "
        "maksadalaýykdygyny görkezýär. Taslama açyk çeşmeli tehnologiýalara esaslanýandygy "
        "sebäpli lisenziýa çykdajylary pesdir, esasy çykdajylar bolsa iş wagtyna we tehniki "
        "serişdelere degişlidir. Programma lokal torlaryň howpsuzlyk barlagyny çaltlandyrýar, "
        "hasabat taýýarlamagy ýeňilleşdirýär we administratorlaryň iş wagtyny tygşytlaýar.",
    )

    return doc


def replace_section_in_main(section_doc: Document) -> None:
    backup = MAIN_DOC.with_name(
        f"{MAIN_DOC.stem}_before_ykdysady_yjd_{datetime.now().strftime('%Y%m%d_%H%M%S')}{MAIN_DOC.suffix}"
    )
    shutil.copy2(MAIN_DOC, backup)

    main = Document(str(MAIN_DOC))
    body = main._body._element

    start_idx = None
    end_idx = None
    paragraph_candidates = []
    fallback_candidates = []

    for p_idx, paragraph in enumerate(main.paragraphs):
        text = " ".join(paragraph.text.split())
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            if style_name.startswith("Heading"):
                paragraph_candidates.append(paragraph)
            elif p_idx > 80:
                fallback_candidates.append(paragraph)

    start_paragraph = (paragraph_candidates or fallback_candidates)[-1] if (paragraph_candidates or fallback_candidates) else None
    if start_paragraph is not None:
        start_idx = body.index(start_paragraph._element)

    for paragraph in main.paragraphs:
        if start_idx is None:
            break
        text = " ".join(paragraph.text.split())
        style_name = paragraph.style.name if paragraph.style is not None else ""
        element_idx = body.index(paragraph._element)
        if element_idx <= start_idx:
            continue
        if text == "Netije" and style_name.startswith("Heading"):
            end_idx = element_idx
            break

    if start_idx is None or end_idx is None:
        raise RuntimeError("Ykdysady bölüm ýa-da global Netije bölümi tapylmady.")

    elements = list(body)
    for element in elements[start_idx:end_idx]:
        body.remove(element)

    netije_element = list(body)[start_idx]
    source_body = section_doc._body._element
    for child in list(source_body):
        if child.tag.endswith("sectPr"):
            continue
        body.insert(body.index(netije_element), copy.deepcopy(child))

    main.save(str(MAIN_DOC))
    print(f"Backup: {backup}")


def main() -> None:
    standalone_doc = build_ykdysady_doc(TEMPLATE)
    standalone_doc.save(str(OUT_DOC))

    section_doc_for_main = build_ykdysady_doc(MAIN_DOC)
    replace_section_in_main(section_doc_for_main)
    print(f"Created: {OUT_DOC}")
    print(f"Updated: {MAIN_DOC}")
    print(f"Total cost: {money(TOTAL_COST)}")
    print(f"Payback: {money(PAYBACK)} years")


if __name__ == "__main__":
    main()
