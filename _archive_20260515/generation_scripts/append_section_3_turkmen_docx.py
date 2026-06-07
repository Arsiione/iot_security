from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_section_3.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def add_table(doc, caption_text, headers, rows):
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "3. Programma toplumynyň işlenip taýýarlanylyşy", level=1)
    add_normal(
        doc,
        "Bu bölümde IoT Security Scanner programma toplumynyň amaly taýdan işlenip taýýarlanylyşy "
        "beýan edilýär. Ikinji bölümde programma arhitekturasy we algoritmler taslanylan bolsa, üçünji "
        "bölümde şol taslamanyň Python we PyQt6 gurşawynda nähili durmuşa geçirilendigi düşündirilýär. "
        "Esasy üns skanirleme ýadrosynyň işleýşine, plugin ulgamy arkaly gowşaklyklary ýüze çykarmaga, "
        "töwekgelçiligi bahalandyrmaga, maglumatlary saklamaga we PDF hasabatyny döretmäge gönükdirilýär."
    )

    add_heading(doc, "3.1. Skanirleme ýadrosynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "Skanirleme ýadrosy core/scanner.py faýlynda ýerleşýär we ScanThread synpy arkaly durmuşa "
            "geçirilýär. Bu synp PyQt6 ulgamyndaky QThread synpyndan miras alýar. Şeýle çözgüt skanirleme "
            "amallarynyň aýratyn akymda ýerine ýetirilmegine mümkinçilik berýär. Netijede, programma "
            "tory barlaýarka grafiki interfeýs doňup galmaýar we ulanyjy progress bar, log ýazgylary hem-de "
            "tapylan gurluşlary real wagt režiminde görüp bilýär."
        ),
        (
            "ScanThread synpy progress, device_found, scan_finished we error atly signallary ulanýar. "
            "progress signaly skanirleme ýagdaýyny we göterim derejesini interfeýse iberýär. device_found "
            "signaly her täze tapylan gurluşy tablisa goşmak üçin ulanylýar. scan_finished skanirleme "
            "tamamlanandan soň ähli gurluşlaryň sanawyny berýär. error signaly bolsa näsazlyklar barada "
            "ulanyja habar bermek üçin niýetlenendir."
        ),
        (
            "Skanirleme başlananda programma ilki tor adapterini we IP aralygyny kesgitleýär. Ulanyjy "
            "interfeýsde Wi-Fi ýa-da Ethernet adapterini saýlap bilýär, programma bolsa şol adapteriň IP "
            "salgysyndan tor aralygyny awtomatik alýar. Bu aýratynlyk taslamanyň möhüm amaly bölegidir, "
            "sebäbi nädogry VPN aralygyny skanirlemek lokal Wi-Fi gurluşlarynyň görünmezligine sebäp bolup "
            "biler."
        ),
        (
            "Işjeň hostlary tapmak üçin birnäçe usul bilelikde ulanylýar. Programma Nmap arkaly -sn -n "
            "ping scan ýerine ýetirýär, lokal hususy torlarda bolsa -PR -sn -n ARP scan ulanýar. Mundan başga-da, "
            "Windows ARP keşi okalýar, ping sweep ýerine ýetirilýär, lokal kompýuter we router şlýuzy netijelere "
            "goşulýar. Şeýlelikde, diňe açyk porty bolan gurluşlar däl, eýsem port açmaýan, emma torda görünýän "
            "telefon ýa-da IoT gurluşlary hem ýüze çykarylyp bilýär."
        ),
        (
            "Her tapylan host üçin scan_device funksiýasy işledilýär. Bu funksiýa host adyny, MAC salgysyny, "
            "öndürijini, tapylan usuly, portlary, hyzmatlary we gurluş görnüşini kesgitleýär. Portlary barlamak "
            "üçin socket.connect_ex usuly ulanylýar. Programma IoT gurşawynda ýygy duş gelýän 21, 22, 23, 80, "
            "443, 554, 8000, 8080, 8888, 1883, 8883, 5683, 1900, 49152 we 5353 portlaryny barlaýar."
        ),
        (
            "Daşky komandalar işledilende Windows ulgamynda gara konsol penjireleriniň açylmazlygy üçin "
            "komandalar gizlin režimde ýerine ýetirilýär. Bu hem ulanyjy tejribesini gowulandyrýar, hem-de "
            "diplom goragy wagtynda programma has arassa we professional görünýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 3.1 - Skanirleme ýadrosynda ulanylýan esasy funksiýalar",
        ["Funksiýa ýa-da synp", "Ýerleşýän faýly", "Wezipesi"],
        [
            ["ScanThread", "core/scanner.py", "Skanirleme prosesini aýratyn akymda ýerine ýetirýär"],
            ["discover_hosts", "core/scanner.py", "Nmap, ARP, Ping we ARP cache arkaly hostlary tapýar"],
            ["scan_device", "core/scanner.py", "Bir IP salgy boýunça gurluş maglumatlaryny ýygnap berýär"],
            ["scan_ports", "core/scanner.py", "IoT üçin möhüm portlaryň açykdygyny barlaýar"],
            ["classify_device", "core/scanner.py", "Gurluşy Router, PC, IoT kandidat ýa-da Telefon/Unknown görnüşinde belleýär"],
            ["normalize_device", "core/scanner.py", "Netijeleri interfeýs we baza üçin bitewi görnüşe getirýär"],
        ],
    )

    add_heading(doc, "3.2. Gowşaklyklary ýüze çykarmak üçin plugin ulgamynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "IoT gurluşlarynyň gowşaklyklary öndüriji, model we firmware wersiýasyna baglylykda üýtgäp "
            "durýar. Şonuň üçin programma diňe bir umumy port barlagy bilen çäklenmän, plugin ulgamyny hem "
            "ulanýar. Plugin ulgamy aýratyn gowşaklyk barlaglaryny esasy skanirleme ýadrosyndan bölüp "
            "saklamaga mümkinçilik berýär. Şeýlelikde, geljekde täze howpsuzlyk barlagy gerek bolsa, "
            "programma gurluşyny düýpli üýtgetmän plugins bukjasyna täze faýl goşmak ýeterlik bolýar."
        ),
        (
            "core/scanner.py içindäki load_plugins funksiýasy plugins bukjasyndaky Python faýllaryny gözläp, "
            "olaryň içinde check atly funksiýanyň bardygyny barlaýar. Eger plugin dogry gurluşa eýe bolsa, "
            "ol skanirleme wagtynda her bir gurluş üçin işledilýär. Plugin gurluş barada maglumat alýar we "
            "eger gowşaklyk tapylsa, vulnerabilities we recommendations sanawlaryny doldurýar."
        ),
        (
            "plugins/hikvision.py faýly web interfeýsi açyk bolan kamera ýa-da wideo gözegçilik enjamlaryny "
            "goşmaça barlamak üçin niýetlenendir. Eger 80-nji port açyk bolsa, plugin HTTP interfeýsiniň "
            "bardygyny belläp, firmware ýagdaýyny, standart hasaplary we daşarky tordan elýeterliligi "
            "barlamagy maslahat berýär. Şeýle ýagdaýda töwekgelçilik derejesi orta derejä çykarylýar."
        ),
        (
            "plugins/telnet_weak_auth.py faýly Telnet hyzmatynda gowşak login/parol jübütlerini barlaýar. "
            "Eger 23-nji port açyk bolsa, plugin admin/admin, root/root, user/user, admin/password, "
            "root/123456 we admin/1234 ýaly kombinasiýalary synap görýär. Login we parol kabul edilip, shell "
            "alamaty ýüze çyksa, gurluş gowşak hasaplanýar we töwekgelçilik derejesi ýokary derejä çykarylýar."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 3.2 - Programmadaky pluginleriň wezipesi",
        ["Plugin", "Barlaýan ýagdaýy", "Netije"],
        [
            ["hikvision.py", "80-nji port we web interfeýs alamatlary", "Firmware we standart hasaplary barlamak boýunça maslahat berýär"],
            ["telnet_weak_auth.py", "23-nji port we gowşak Telnet login/parol jübütleri", "Gowşak autentifikasiýa tapylsa ýokary töwekgelçilik belleýär"],
            ["Täze plugin modeli", "Geljekde goşuljak öndüriji ýa-da hyzmat boýunça barlaglar", "Programmany giňeltmäge mümkinçilik berýär"],
        ],
    )

    add_heading(doc, "3.3. Töwekgelçiligi bahalandyrmak we maslahat beriş ulgamynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "Programmada her bir tapylan gurluş üçin töwekgelçilik derejesi başlangyç ýagdaýda low, ýagny "
            "pes baha bilen döredilýär. Soňra portlaryň ýagdaýy, plugin barlaglarynyň netijesi we tapylan "
            "gowşaklyklar boýunça bu dereje üýtgedilip bilýär. Interfeýsde bu bahalar türkmen dilinde pes, "
            "orta we ýokary görnüşinde görkezilýär."
        ),
        (
            "Eger gurluşda açyk port ýok bolsa we anyk gowşaklyk tapylmasa, programma ony pes töwekgelçilik "
            "hökmünde görkezýär. Bu ýagdaýda maslahat hökmünde “Portlar açyk däl, gurluş diňe torda ýüze "
            "çykaryldy” diýen ýazgy berilýär. Bu çözgüt ulanyja gurluşyň torda bardygyny, emma häzirlikçe "
            "açyk hyzmat tapylmandygyny düşündirýär."
        ),
        (
            "HTTP interfeýsi ýa-da kamera bilen bagly alamatlar ýüze çykanda töwekgelçilik orta derejä "
            "çykarylyp bilýär. Şeýle ýagdaýlarda programma firmware täzelenmesini barlamagy, standart "
            "hasaplary öçürmegi we web interfeýsiň daşarky tordan elýeterliligini çäklendirmegi maslahat "
            "berýär. Telnet gowşak autentifikasiýasy tapylanda bolsa töwekgelçilik ýokary derejä çykarylýar."
        ),
        (
            "RemediationEngine synpy düzediş amallarynyň başlangyç modelini üpjün edýär. Bu synp Telnet "
            "parolyny üýtgetmek, FTP ýerine SFTP ulanmak, web interfeýsinde standart paroly çalyşmak we SSH "
            "konfigurasiýasyny berkitmek ýaly mümkinçilikleri modelleşdirýär. Häzirki taslamada esasy üns "
            "awtomatik düzedişe däl, eýsem ulanyja howpsuz we düşnükli maslahat bermäge gönükdirilendir."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 3.3 - Töwekgelçilik we maslahat beriş logikasy",
        ["Ýagdaý", "Töwekgelçilik", "Maslahat"],
        [
            ["Gurluş torda tapyldy, port açyk däl", "Pes", "Gurluş diňe torda ýüze çykaryldy"],
            ["HTTP/web interfeýs açyk", "Orta", "Firmware-i we standart hasaplary barlamak"],
            ["RTSP ýa-da kamera hyzmatlary açyk", "Orta", "Wideo akymynyň rugsatlaryny we parollaryny barlamak"],
            ["Telnet gowşak autentifikasiýasy tapyldy", "Ýokary", "Telnet-i öçürmek ýa-da güýçli parol goýmak"],
            ["Gowşaklyk tapyldy, emma awtomatik düzediş howply bolup biler", "Üýtgeýän", "Administrator tarapyndan el bilen barlag geçirmek"],
        ],
    )

    add_heading(doc, "3.4. Maglumatlar bazasy we skanirleme taryhynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "Skanirleme tamamlanandan soň netijeler database.py modulynyň save_scan funksiýasy arkaly "
            "iot_security.db bazasynda saklanýar. Her tapylan gurluş aýratyn setir hökmünde ýazylýar, emma "
            "bir skanirleme sessiýasyndaky gurluşlar umumy timestamp bahasy bilen birleşdirilýär. Bu gurluş "
            "öňki skanirlemeleri toparlap görkezmäge mümkinçilik berýär."
        ),
        (
            "load_history funksiýasy bazadaky netijeleri okap, timestamp boýunça toparlaýar we interfeýse "
            "taýýar sanaw görnüşinde berýär. Şu maglumatlar Taryh sahypasynda öňki skanirlemeleri görkezmek "
            "üçin ulanylýar. get_scan_stats funksiýasy bolsa umumy skanirleme sanyny, tapylan gurluşlaryň "
            "sanyny, gowşak gurluşlaryň sanyny we soňky skanirleme wagtyny hasaplaýar."
        ),
        (
            "Panel, Netijeler we Taryh sahypalary maglumatlar bazasy bilen göni baglanyşyklydyr. Panel "
            "sahypasy umumy statistikany görkezýär. Netijeler sahypasy iň soňky skanirlemäniň gurluşlaryny "
            "jikme-jik seljerýär. Taryh sahypasy bolsa öňki skanirlemeleri arhiw görnüşinde görkezýär we "
            "zerur bolsa saýlanan netijäni PDF hasabata eksport etmäge mümkinçilik berýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 3.4 - Maglumatlar bazasy bilen işleýän funksiýalar",
        ["Funksiýa", "Ýerleşýän faýly", "Wezipesi"],
        [
            ["init_db", "database.py", "SQLite bazasyny we scans tablisasyny döredýär"],
            ["migrate_scans_table", "database.py", "Täze sütünleri awtomatik goşýar"],
            ["save_scan", "database.py", "Skanirleme netijelerini bazada saklaýar"],
            ["load_history", "database.py", "Öňki skanirlemeleri timestamp boýunça okap berýär"],
            ["get_scan_stats", "database.py", "Panel üçin umumy statistikany hasaplaýar"],
            ["clear_history", "database.py", "Skanirleme taryhyny arassalaýar"],
        ],
    )

    add_heading(doc, "3.5. Hasabat döretmek we duýduryş ulgamynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "Programmada skanirleme netijelerini diňe ekranda görkezmek bilen çäklenilmän, PDF görnüşinde "
            "hasabat döretmek mümkinçiligi hem durmuşa geçirildi. reports.py faýlyndaky generate_pdf "
            "funksiýasy ReportLab kitaphanasynyň kömegi bilen hasabat faýlyny döredýär. Hasabatda gurluşyň "
            "ady, IP salgysy, görnüşi, öndürijisi, tapylan usuly, açyk portlary, gowşaklyk ýagdaýy we maslahat "
            "görkezilýär."
        ),
        (
            "PDF hasabatynda türkmen harplarynyň dogry görkezilmegi üçin programma Windows ulgamyndaky Arial "
            "ýa-da Segoe UI şriftlerini gözleýär we olary hasabata birikdirýär. Şeýle çözgüt hasabatlaryň "
            "okalmagyny gowulandyrýar. Şeýle hem hasabatda howpsuz we gowşak gurluşlaryň paýyny görkezýän "
            "ýönekeý diagramma döredilýär."
        ),
        (
            "Duýduryş ulgamy notifications.py we utils/notifications.py faýllarynda başlangyç görnüşde "
            "durmuşa geçirilendir. notifications.py ýönekeý email duýduryşyny ibermäge mümkinçilik berýär. "
            "utils/notifications.py bolsa giňeldilen NotificationManager synpyny öz içine alýar: ol email, "
            "webhook we SMS ýaly duýduryş ugurlaryny goldamak üçin taslanylandyr. Häzirki programma üçin bu "
            "bölek geljekde howpsuzlyk duýduryşlaryny ösdürmek üçin binýat bolup hyzmat edýär."
        ),
        (
            "Hasabat we duýduryş ulgamlary diplom taslamasynyň amaly ähmiýetini güýçlendirýär. Administrator "
            "skanirleme netijelerini diňe programma içinde görmän, olary PDF faýly hökmünde saklap, ýolbaşça "
            "ýa-da diplom komissiýasyna görkezip bilýär. Bu bolsa programmany diňe synag guraly däl, eýsem "
            "audit netijesini resmileşdirýän kiçi programma toplumy hökmünde görkezýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 3.5 - Hasabat we duýduryş ulgamynyň bölekleri",
        ["Bölek", "Ýerleşýän faýly", "Amaly wezipesi"],
        [
            ["generate_pdf", "reports.py", "Skanirleme netijelerinden PDF hasabat döredýär"],
            ["get_report_font", "reports.py", "Türkmen harplary üçin laýyk şrift saýlaýar"],
            ["Matplotlib diagrammasy", "reports.py", "Howpsuz we gowşak gurluşlaryň paýyny görkezýär"],
            ["send_email", "notifications.py", "Ýönekeý email duýduryşyny ibermek üçin ulanylýar"],
            ["NotificationManager", "utils/notifications.py", "Email, webhook we SMS duýduryşlaryny giňeltmek üçin model berýär"],
        ],
    )

    add_normal(
        doc,
        "Şeýlelikde, üçünji bölümde IoT Security Scanner programma toplumynyň esasy amaly bölekleri "
        "durmuşa geçirildi: skanirleme ýadrosy, plugin ulgamy, töwekgelçiligi bahalandyrmak, maglumatlar "
        "bazasy, PDF hasabaty we duýduryş modeli. Bu bölekleriň bilelikde işlemegi programmany lokal IoT "
        "torlaryny barlamak, netijeleri saklamak we howpsuzlyk boýunça maslahat bermek üçin doly funksional "
        "programma toplumyna öwürýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_styles(main_doc):
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text == "3. Programma toplumynyň işlenip taýýarlanylyşy":
            paragraph.style = "Heading 1"
            set_font(paragraph, 14, True)
        elif text.startswith(("3.1.", "3.2.", "3.3.", "3.4.", "3.5.")):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip() == "3. Programma toplumynyň işlenip taýýarlanylyşy" for p in main_doc.paragraphs):
        print("Section 3 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_styles(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
