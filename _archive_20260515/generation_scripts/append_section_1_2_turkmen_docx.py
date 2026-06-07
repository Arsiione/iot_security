from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOCX_FILE = "bolum_1_1_iot_gurlushlary.docx"
TITLE = "1.2. IoT howpsuzlygynyň wajyplygy"


PARAGRAPHS = [
    (
        "Internet of Things tehnologiýalarynyň giňden ýaýramagy bilen tor howpsuzlygy "
        "meseleleri has hem möhüm ähmiýete eýe boldy. Häzirki wagtda diňe kompýuterler "
        "we serwerler däl, eýsem kameralar, datçikler, akylly öý enjamlary, marşrutizatorlar, "
        "lukmançylyk enjamlary we senagat kontrollerleri hem internete ýa-da lokal tora "
        "birikdirilýär. Bu bolsa toruň düzümindäki gurluşlaryň sanyny artdyrýar we "
        "hüjüm edilip bilinjek nokatlaryň köpelmegine getirýär."
    ),
    (
        "IoT gurluşlarynyň howpsuzlygy aýratyn wajypdyr, sebäbi şeýle enjamlar köplenç "
        "fiziki prosesler bilen baglanyşykly bolýar. Mysal üçin, IP kamera wideo gözegçilik "
        "maglumatlaryny geçirýär, akylly gulplar giriş rugsadyny dolandyrýar, senagat "
        "datçikleri bolsa önümçilik prosesleriniň ýagdaýyny ölçäp berýär. Şeýle gurluşlaryň "
        "goragsyz bolmagy diňe maglumat howpsuzlygyna däl, eýsem fiziki howpsuzlyga hem "
        "täsir edip biler."
    ),
    (
        "IoT gurşawynda esasy meseleleriň biri standart login we parollaryň ulanylmagydyr. "
        "Köp öndürijiler enjamlary başlangyç ýagdaýda admin/admin, root/root ýa-da şuňa "
        "meňzeş ýönekeý giriş maglumatlary bilen goýberýärler. Ulanyjy bu maglumatlary "
        "üýtgetmese, hüjümçi açyk hyzmatlaryň üsti bilen gurluşa aňsatlyk bilen girip biler. "
        "Bu ýagdaý esasan hem Telnet, HTTP we FTP ýaly hyzmatlar açyk bolanda howplydyr."
    ),
    (
        "Ýene bir möhüm mesele programma üpjünçiliginiň wagtynda täzelenmezligidir. IoT "
        "gurluşlarynyň köpüsi birnäçe ýyl dowamynda üznüksiz işleýär we ulanyjy olaryň "
        "firmware täzelenmelerini barlap durmaýar. Netijede, öndüriji tarapyndan öňden "
        "düzedilen gowşaklyklar hem enjamda saklanyp galýar. Hüjümçiler şeýle köne "
        "gowşaklyklary awtomatlaşdyrylan gurallar arkaly tapyp, olardan peýdalanyp bilýär."
    ),
    (
        "Açyk portlar we goragsyz protokollar hem IoT howpsuzlygynyň möhüm meseleleriniň "
        "biridir. Mysal üçin, Telnet protokoly maglumatlary şifrlemezden geçirýär, HTTP "
        "aragatnaşygynda bolsa login we parol ýaly maglumatlaryň goragy pes bolup biler. "
        "RTSP porty açyk bolan kameralar wideo akymynyň rugsatsyz görülmegine sebäp bolup "
        "biler. MQTT hyzmatynyň nädogry sazlanmagy bolsa IoT datçiklerinden gelýän "
        "maglumatlaryň rugsatsyz okalmagyna ýa-da üýtgedilmegine ýol açýar."
    ),
    (
        "IoT gurluşlarynyň howpsuzlygynyň pes bolmagy botnet hüjümleriniň döremegine hem "
        "sebäp bolýar. Taryhda Mirai botneti ýaly hüjümler müňlerçe goragsyz IoT enjamyny "
        "ele geçirip, olary DDoS hüjümlerinde ulanandygyny görkezdi. Bu mysal IoT "
        "gurluşlarynyň diňe bir aýratyn ulanyjy üçin däl, eýsem internet infrastrukturasynyň "
        "umumy durnuklylygy üçin hem howp döredip bilýändigini subut edýär."
    ),
    (
        "IoT gurluşlary şahsy maglumatlaryň goragyna hem gönüden-göni täsir edýär. "
        "Kameralar wideo ýazgylaryny, saglyk datçikleri lukmançylyk maglumatlaryny, "
        "akylly öý ulgamlary bolsa ulanyjynyň gündelik hereketleri barada maglumatlary "
        "ýygnap bilýär. Şeýle maglumatlaryň rugsatsyz ele geçirilmegi şahsy durmuşyň "
        "eldegrilmesizligine zyýan ýetirýär. Şol sebäpli IoT howpsuzlygy diňe tehniki "
        "mesele bolman, eýsem maglumatlaryň gizlinligi bilen baglanyşykly möhüm mesele "
        "hökmünde hem garalmalydyr."
    ),
    (
        "Kärhanalarda we okuw edaralarynda IoT gurluşlarynyň sanynyň artmagy administratorlar "
        "üçin goşmaça jogapkärçilik döredýär. Toruň içinde haýsy gurluşlaryň bardygyny, "
        "olaryň haýsy hyzmatlary işledýändigini we haýsy töwekgelçilik derejesine eýedigini "
        "el bilen barlamak köp wagt talap edýär. Şonuň üçin awtomatlaşdyrylan skanirleme "
        "ulgamlary tor howpsuzlygyny yzygiderli gözegçilikde saklamak üçin zerur guraldyr."
    ),
    (
        "IoT howpsuzlygynyň wajyplygy diňe gowşaklyklary ýüze çykarmak bilen çäklenmeýär. "
        "Howpsuzlyk auditi geçirilenden soň ulanyja anyk maslahatlaryň berilmegi hem möhümdir. "
        "Mysal üçin, açyk Telnet hyzmaty tapylsa, ony öçürmek ýa-da güýçli parol ulanmak "
        "maslahat berilýär. HTTP interfeýsi açyk bolsa, firmware-i täzelemek we standart "
        "hasaplary üýtgetmek zerur bolup biler. Şeýle çemeleşme ýüze çykarylan meseleleri "
        "amaly taýdan azaltmaga mümkinçilik berýär."
    ),
    (
        "Şu sebäplerden ugur alyp, diplom işinde işlenip taýýarlanylýan IoT Security Scanner "
        "programmasynyň ähmiýeti ýokarlanýar. Programma lokal tordaky gurluşlary awtomatik "
        "tapmaga, açyk portlary kesgitlemäge, potensial gowşaklyklary seljermäge we ulanyja "
        "howpsuzlyk boýunça maslahat bermäge gönükdirilendir. Bu bolsa IoT torlarynyň "
        "başlangyç howpsuzlyk auditini geçirmekde peýdaly çözgüt bolup durýar."
    ),
]


TABLE_ROWS = [
    ("Howpsuzlyk meselesi", "Düşündiriş", "Mümkin bolan netijesi"),
    (
        "Standart parollar",
        "Gurluş admin/admin ýa-da root/root ýaly başlangyç parollar bilen galýar.",
        "Rugsatsyz giriş we gurluşyň dolandyrylyp bilinmegi.",
    ),
    (
        "Täzelenmedik firmware",
        "Öndüriji tarapyndan düzedilen gowşaklyklar enjamda saklanyp galýar.",
        "Köne gowşaklyklardan peýdalanmak mümkinçiligi.",
    ),
    (
        "Açyk Telnet/HTTP",
        "Maglumatlar şifrlenmezden geçirilýär ýa-da web-panel goragsyz bolýar.",
        "Parollaryň ele geçirilmegi we konfigurasiýanyň üýtgedilmegi.",
    ),
    (
        "Açyk RTSP",
        "Kameranyň wideo akymy goragsyz elýeterli bolup biler.",
        "Wideo maglumatlarynyň rugsatsyz görülmegi.",
    ),
    (
        "Nädogry MQTT sazlamalary",
        "IoT datçikleriniň habarlary rugsatsyz okalyp ýa-da üýtgedilip bilner.",
        "Maglumatlaryň ýoýulmagy we awtomatlaşdyryş proseslerine täsir etmek.",
    ),
]


CONCLUSION = (
    "Netijede, IoT howpsuzlygy häzirki zaman tor infrastrukturasynyň aýrylmaz bölegi "
    "hasaplanýar. IoT gurluşlaryny wagtynda ýüze çykarmak, olaryň açyk hyzmatlaryny "
    "seljermek we töwekgelçiliklere görä maslahat bermek lokal toruň umumy howpsuzlyk "
    "derejesini ýokarlandyrmaga mümkinçilik berýär."
)


def apply_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)


def add_text_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_section_heading(document, title):
    document.add_paragraph()
    heading = document.add_heading(title, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True


def add_risk_table(document):
    document.add_paragraph()
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, text in enumerate(TABLE_ROWS[0]):
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in TABLE_ROWS[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("Tablisa 1.2 - IoT howpsuzlygynda esasy meseleler we olaryň netijeleri")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True


def section_already_exists(document):
    return any(paragraph.text.strip() == TITLE for paragraph in document.paragraphs)


def append_section():
    document = Document(DOCX_FILE)
    apply_document_defaults(document)

    if section_already_exists(document):
        print(f"{TITLE} eýýäm dokumentde bar")
        return

    add_section_heading(document, TITLE)

    for text in PARAGRAPHS:
        add_text_paragraph(document, text)

    add_risk_table(document)
    add_text_paragraph(document, CONCLUSION)

    document.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == "__main__":
    append_section()
