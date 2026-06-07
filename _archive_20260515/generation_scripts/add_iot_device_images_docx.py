from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


DOCX_FILE = Path("bolum_1_1_iot_gurlushlary.docx")
BACKUP_FILE = Path("bolum_1_1_iot_gurlushlary_before_device_images.docx")
ASSET_DIR = Path("diploma_assets/iot_devices")

SECTION_TITLE = "IoT gurluşlarynyň mysal şekilleri"
CAPTION = "Surat 1.1 - IoT gurluşlarynyň görnüşleri boýunça şekiller"


DEVICES = [
    {
        "title": "IP kamera",
        "description": "Wideo gözegçilik üçin ulanylýar. Esasy howplar: RTSP/HTTP portlarynyň açyk bolmagy we standart parollar.",
        "filename": "ip_camera.png",
        "drawer": "camera",
    },
    {
        "title": "Akylly öý enjamlary",
        "description": "Rozetka, lampa, termostat we datçikler. Esasy howplar: goragsyz API we gowşak awtorizasiýa.",
        "filename": "smart_home.png",
        "drawer": "home",
    },
    {
        "title": "Tor enjamlary",
        "description": "Router, Wi-Fi nokady we switch. Esasy howplar: açyk Telnet/SSH we köne firmware.",
        "filename": "router.png",
        "drawer": "router",
    },
    {
        "title": "Senagat IoT gurluşlary",
        "description": "Kontrollerler we önümçilik datçikleri. Esasy howplar: proseslere rugsatsyz täsir etmek.",
        "filename": "industrial_iot.png",
        "drawer": "industry",
    },
    {
        "title": "Lukmançylyk IoT enjamlary",
        "description": "Saglyk monitorlary we datçikler. Esasy howplar: şahsy maglumatlaryň syzmagy.",
        "filename": "medical_iot.png",
        "drawer": "medical",
    },
]


def font(size=34, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def base_canvas(title):
    image = Image.new("RGB", (900, 420), "#f7fafc")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((18, 18, 882, 402), radius=28, fill="#ffffff", outline="#0f766e", width=5)
    draw.text((40, 28), title, fill="#0f172a", font=font(34, bold=True))
    return image, draw


def draw_network_nodes(draw, points, color="#0f766e"):
    for start, end in zip(points, points[1:]):
        draw.line((start, end), fill=color, width=5)
    for x, y in points:
        draw.ellipse((x - 15, y - 15, x + 15, y + 15), fill=color)


def draw_camera(title):
    image, draw = base_canvas(title)
    draw.rounded_rectangle((150, 145, 470, 285), radius=26, fill="#dff7f3", outline="#0f766e", width=6)
    draw.polygon([(470, 185), (610, 145), (610, 285), (470, 245)], fill="#99f6e4", outline="#0f766e")
    draw.ellipse((245, 175, 365, 295), fill="#0f766e")
    draw.ellipse((285, 210, 330, 255), fill="#e0f2fe")
    draw.rectangle((280, 285, 335, 350), fill="#0f766e")
    draw.rounded_rectangle((210, 340, 405, 365), radius=10, fill="#0f766e")
    draw.arc((635, 150, 760, 275), 300, 60, fill="#14b8a6", width=7)
    draw.arc((660, 178, 735, 253), 300, 60, fill="#14b8a6", width=6)
    return image


def draw_home(title):
    image, draw = base_canvas(title)
    draw.polygon([(225, 240), (390, 115), (555, 240)], fill="#ccfbf1", outline="#0f766e")
    draw.rounded_rectangle((265, 235, 515, 350), radius=16, fill="#ffffff", outline="#0f766e", width=6)
    draw.rectangle((365, 285, 420, 350), fill="#99f6e4", outline="#0f766e", width=4)
    draw.rounded_rectangle((560, 180, 700, 330), radius=28, fill="#dff7f3", outline="#0f766e", width=6)
    draw.ellipse((605, 220, 655, 270), fill="#0f766e")
    draw_network_nodes(draw, [(390, 160), (610, 125), (635, 180), (635, 220)])
    return image


def draw_router(title):
    image, draw = base_canvas(title)
    draw.rounded_rectangle((185, 210, 625, 320), radius=22, fill="#dff7f3", outline="#0f766e", width=6)
    for x in [250, 320, 390]:
        draw.ellipse((x, 250, x + 24, 274), fill="#0f766e")
    draw.rectangle((485, 245, 575, 275), fill="#ffffff", outline="#0f766e", width=4)
    draw.line((245, 210, 185, 120), fill="#0f766e", width=7)
    draw.line((565, 210, 640, 120), fill="#0f766e", width=7)
    for box in [(690, 150, 760, 210), (700, 250, 770, 310), (100, 150, 170, 210)]:
        draw.rounded_rectangle(box, radius=12, fill="#ffffff", outline="#14b8a6", width=5)
    draw_network_nodes(draw, [(625, 260), (700, 180), (725, 180)], "#14b8a6")
    draw_network_nodes(draw, [(625, 285), (710, 280), (735, 280)], "#14b8a6")
    draw_network_nodes(draw, [(185, 260), (145, 180), (135, 180)], "#14b8a6")
    return image


def draw_industry(title):
    image, draw = base_canvas(title)
    draw.rectangle((165, 220, 620, 350), fill="#e0f2fe", outline="#0f766e", width=6)
    draw.rectangle((190, 165, 245, 220), fill="#99f6e4", outline="#0f766e", width=5)
    draw.rectangle((275, 130, 330, 220), fill="#99f6e4", outline="#0f766e", width=5)
    draw.rectangle((360, 185, 415, 220), fill="#99f6e4", outline="#0f766e", width=5)
    draw.rounded_rectangle((655, 150, 760, 350), radius=14, fill="#ffffff", outline="#0f766e", width=6)
    for y in [190, 235, 280]:
        draw.ellipse((690, y, 725, y + 35), fill="#14b8a6")
    draw.line((620, 285, 655, 240), fill="#0f766e", width=5)
    draw.line((620, 245, 655, 285), fill="#0f766e", width=5)
    return image


def draw_medical(title):
    image, draw = base_canvas(title)
    draw.rounded_rectangle((165, 130, 530, 335), radius=24, fill="#ffffff", outline="#0f766e", width=6)
    draw.line((210, 240, 270, 240, 295, 190, 340, 300, 375, 220, 430, 220), fill="#ef4444", width=8)
    draw.rounded_rectangle((585, 150, 735, 335), radius=26, fill="#dff7f3", outline="#0f766e", width=6)
    draw.rectangle((640, 190, 680, 295), fill="#0f766e")
    draw.rectangle((607, 222, 713, 262), fill="#0f766e")
    draw_network_nodes(draw, [(530, 225), (585, 210), (640, 190)], "#14b8a6")
    return image


DRAWERS = {
    "camera": draw_camera,
    "home": draw_home,
    "router": draw_router,
    "industry": draw_industry,
    "medical": draw_medical,
}


def create_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for device in DEVICES:
        image = DRAWERS[device["drawer"]](device["title"])
        image.save(ASSET_DIR / device["filename"])


def set_run_style(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_device_image_block(document):
    heading = document.add_heading(SECTION_TITLE, level=3)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        set_run_style(run, size=14, bold=True)

    intro = document.add_paragraph()
    set_paragraph(intro, first_line=True)
    run = intro.add_run(
        "Aşakdaky şekiller IoT gurluşlarynyň esasy görnüşlerini görkezýär. "
        "Suratlar diplom işindäki seljermäni has düşnükli etmek we skanirlenýän "
        "gurluşlaryň amaly mysallaryny görkezmek üçin goşuldy."
    )
    set_run_style(run, size=14)

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    for cell, text in zip(header, ["IoT gurluşy", "Şekil"]):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_style(run, size=12, bold=True)

    for device in DEVICES:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        title_p = cells[0].paragraphs[0]
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_p.add_run(device["title"])
        set_run_style(title_run, size=12, bold=True)

        desc_p = cells[0].add_paragraph()
        set_paragraph(desc_p)
        desc_run = desc_p.add_run(device["description"])
        set_run_style(desc_run, size=11)

        image_p = cells[1].paragraphs[0]
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_run = image_p.add_run()
        image_run.add_picture(str(ASSET_DIR / device["filename"]), width=Cm(6.2))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run(CAPTION)
    set_run_style(run, size=12, italic=True)

    return [heading._p, intro._p, table._tbl, caption._p]


def find_anchor_paragraph(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Tablisa 1.1"):
            return paragraph
    raise RuntimeError("Tablisa 1.1 tapylmady")


def block_already_exists(document):
    return any(paragraph.text.strip() == SECTION_TITLE for paragraph in document.paragraphs)


def insert_after_anchor(anchor, elements):
    for element in reversed(elements):
        anchor._p.addnext(element)


def add_images_to_docx():
    create_images()
    document = Document(DOCX_FILE)

    if block_already_exists(document):
        print(f"{SECTION_TITLE} eýýäm dokumentde bar")
        return

    if not BACKUP_FILE.exists():
        BACKUP_FILE.write_bytes(DOCX_FILE.read_bytes())

    anchor = find_anchor_paragraph(document)
    elements = add_device_image_block(document)
    insert_after_anchor(anchor, elements)
    document.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == "__main__":
    add_images_to_docx()
