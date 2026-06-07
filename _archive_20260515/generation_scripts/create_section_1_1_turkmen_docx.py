from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_FILE = "bolum_1_1_iot_gurlushlary.docx"


TITLE = "1.1. IoT gurluşlarynyň düşünjesi we aýratynlyklary"


PARAGRAPHS = [
    (
        "Soňky ýyllarda sanly tehnologiýalaryň ösmegi bilen adaty fiziki "
        "obýektleriň internete birikdirilmegi giň gerime eýe boldy. Bu hadysa "
        "Internet of Things ýa-da gysgaça IoT, ýagny zatlaryň interneti diýilýär. "
        "IoT düşünjesi dürli enjamlaryň, datçikleriň, gözegçilik ulgamlarynyň we "
        "dolandyryş modullarynyň tor arkaly özara maglumat alyş-çalyş etmegini "
        "aňladýar. Şeýle gurluşlar diňe bir maglumat ýygnamak bilen çäklenmän, "
        "alnan maglumatlaryň esasynda belli bir hereketleri hem ýerine ýetirip "
        "bilýär."
    ),
    (
        "IoT gurluşlaryna IP kameralar, akylly rozetkalar, akylly lampalar, "
        "temperatura we çyglylyk datçikleri, öý howpsuzlyk ulgamlary, marşrutizatorlar, "
        "giriş gözegçilik terminallary, lukmançylyk datçikleri, önümçilikde ulanylýan "
        "kontrollerler we beýleki tor bilen baglanyşykly enjamlar degişlidir. Bu "
        "gurluşlar gündelik durmuşda, bilim edaralarynda, saglygy goraýyşda, senagatda, "
        "ulag ulgamynda we akylly şäher taslamalarynda giňden ulanylýar."
    ),
    (
        "IoT ulgamynyň esasy aýratynlygy onuň fiziki dünýä bilen sanly ulgamy "
        "birleşdirmegidir. Mysal üçin, IP kamera daşky gurşawy wideo görnüşinde "
        "gözegçilik edýär, temperatura datçigi howanyň ýagdaýyny ölçýär, akylly "
        "rozetka bolsa elektrik üpjünçiligini uzakdan dolandyrmaga mümkinçilik berýär. "
        "Şeýlelikde, IoT gurluşlary diňe maglumat çeşmesi däl, eýsem awtomatlaşdyrylan "
        "dolandyryş ulgamynyň möhüm bölegi bolup çykyş edýär."
    ),
    (
        "Adaty kompýuterlerden tapawutlylykda IoT gurluşlarynyň köpüsi çäkli hasaplaýyş "
        "mümkinçiliklerine eýedir. Olarda prosessor kuwwaty, operatiw ýat we maglumat "
        "saklaýyş göwrümi az bolýar. Şol sebäpli käbir öndürijiler howpsuzlyk "
        "mehanizmlerine ýeterlik üns bermeýärler. Netijede, şeýle enjamlarda ýönekeý "
        "parollar, köne programma üpjünçiligi, şifrlenmedik protokollar we açyk portlar "
        "duş gelip bilýär."
    ),
    (
        "IoT gurluşlarynyň ýene bir möhüm aýratynlygy olaryň uzak wagtlap üznüksiz "
        "işlemegidir. Köp ýagdaýlarda enjam bir gezek gurnalýar we birnäçe ýyl "
        "dowamynda täzelenmän ulanylýar. Ulanyjy ýa-da administrator programma üpjünçiliginiň "
        "täzelenişini, standart parollaryň üýtgedilmegini we açyk hyzmatlaryň "
        "barlagyny hemişe ýerine ýetirmeýär. Bu ýagdaý hüjümçiler üçin amatly şert "
        "döredýär."
    ),
    (
        "Köp IoT gurluşlarynda web-interfeýs, Telnet, SSH, FTP, RTSP, MQTT ýaly "
        "tor hyzmatlary ulanylýar. Bu hyzmatlar enjamy dolandyrmak, maglumat almak "
        "ýa-da wideo akymyny geçirmek üçin zerurdyr. Emma bu hyzmatlaryň nädogry "
        "sazlanmagy howpsuzlyk töwekgelçiligini döredýär. Mysal üçin, Telnet protokoly "
        "maglumatlary açyk görnüşde geçirýär, HTTP bolsa şifrlenmedik web-aragatnaşygy "
        "ulanýar. Şonuň üçin şeýle hyzmatlaryň barlygyny we ýagdaýyny yzygiderli "
        "barlamak zerurdyr."
    ),
    (
        "IoT gurluşlarynyň howpsuzlyk meseleleri diňe bir aýratyn ulanyjynyň şahsy "
        "maglumatlaryna däl, eýsem bütin toruň durnuklylygyna hem täsir edip biler. "
        "Eger bir IoT enjamynda gowşak parol ýa-da açyk howply port bar bolsa, hüjümçi "
        "şol enjam arkaly tora girip, beýleki ulgamlara hem täsir edip biler. Şeýle "
        "ýagdaýda IoT gurluşy hüjümiň esasy maksady bolman, tora giriş nokady hökmünde "
        "ulanylyp bilner."
    ),
    (
        "IoT gurluşlarynyň sanynyň artmagy bilen administratorlar üçin ähli enjamlary "
        "el bilen gözegçilikde saklamak kynlaşýar. Toruň içinde haýsy gurluşlaryň bardygyny, "
        "olaryň haýsy IP salgylary ulanýandygyny, haýsy portlarynyň açykdygyny we haýsy "
        "hyzmatlaryň işleýändigini çalt kesgitlemek zerur bolýar. Şol sebäpli "
        "awtomatlaşdyrylan skanirleme we howpsuzlyk auditi ulgamlary aýratyn ähmiýete "
        "eýe bolýar."
    ),
    (
        "Diplom işinde işlenip taýýarlanylýan IoT Security Scanner programmasy şu "
        "meseleleri çözmäge gönükdirilendir. Programma lokal tordaky gurluşlary tapýar, "
        "olaryň açyk portlaryny kesgitleýär, potensial howpsuzlyk töwekgelçiliklerini "
        "bahalandyrýar we ulanyja degişli maslahatlary hödürleýär. Şeýle çemeleşme "
        "IoT torlarynyň başlangyç howpsuzlyk auditini geçirmek üçin amatly gural "
        "bolup hyzmat edýär."
    ),
]


TABLE_ROWS = [
    ("IoT gurluşynyň görnüşi", "Mysallar", "Mümkin bolan howpsuzlyk töwekgelçiligi"),
    (
        "IP kameralar",
        "Hikvision, Dahua, beýleki wideo gözegçilik kameralary",
        "Açyk RTSP/HTTP portlary, standart parollar, köne firmware",
    ),
    (
        "Akylly öý enjamlary",
        "Rozetkalar, lampalar, termostatlar, datçikler",
        "Gowşak awtorizasiýa, goragsyz API, nädogry sazlamalar",
    ),
    (
        "Tor enjamlary",
        "Marşrutizatorlar, Wi-Fi nokatlary, repeaterler",
        "Açyk Telnet/SSH, köne programma üpjünçiligi, standart giriş maglumatlary",
    ),
    (
        "Senagat IoT gurluşlary",
        "Kontrollerler, ölçeg datçikleri, awtomatika modullary",
        "Önümçilik proseslerine rugsatsyz täsir etmek howpy",
    ),
    (
        "Lukmançylyk IoT enjamlary",
        "Saglyk gözegçilik datçikleri, monitorlar",
        "Şahsy maglumatlaryň syzmagy, maglumatlaryň nädogry geçirilmegi",
    ),
]


CONCLUSION = (
    "Şeýlelikde, IoT gurluşlary häzirki zaman tor infrastrukturasynyň möhüm bölegi "
    "bolup durýar. Emma olaryň köpçülikleýin ulanylmagy täze howpsuzlyk meselelerini "
    "hem döredýär. Şonuň üçin IoT gurluşlaryny tapmak, açyk hyzmatlaryny seljermek "
    "we ýüze çykarylan töwekgelçilikler boýunça maslahat bermek diplom işiniň esasy "
    "ugruny emele getirýär."
)


def set_document_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_table(document):
    document.add_paragraph()
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, text in enumerate(TABLE_ROWS[0]):
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in TABLE_ROWS[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("Tablisa 1.1 - IoT gurluşlarynyň görnüşleri we howpsuzlyk töwekgelçilikleri")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True


def build_document():
    document = Document()
    set_document_styles(document)

    heading = document.add_heading(TITLE, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True

    for text in PARAGRAPHS:
        add_paragraph(document, text)

    add_table(document)
    add_paragraph(document, CONCLUSION)

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_FILE)
