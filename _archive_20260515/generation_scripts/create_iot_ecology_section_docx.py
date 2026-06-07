from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_FILE = "iot_security_zahmeti_goramak_ekologiya.docx"


TITLE = "Zähmeti goramak we ekologik howpsuzlyk boýunça hasaplamalar"


INTRO_PARAGRAPHS = [
    (
        "Bu bölümde IoT Security Scanner programma toplumyny işläp taýýarlamak we ulanmak "
        "bilen baglanyşykly zähmet şertleri, elektrik howpsuzlygy, iş otagynyň mikroklimaty, "
        "ýagtylandyrylyşy, energiýa sarp edilişi we ekologik täsiri seljerilýär. Taslanýan "
        "obýekt önümçilik enjamy däl-de, kompýuter we tor enjamlary bilen işleýän programma "
        "üýtgemesi bolany üçin esasy üns iş ýeriniň ergonomikasyna, elektrik enjamlarynyň "
        "howpsuz ulanylmagyna we energiýanyň tygşytly sarp edilmegine berilýär."
    ),
    (
        "IoT Security Scanner lokal torda ýerleşýän gurluşlary awtomatik tapmak, olaryň açyk "
        "portlaryny kesgitlemek, potensial gowşaklyklary ýüze çykarmak we ulanyja maslahat "
        "bermek üçin niýetlenendir. Programma bilen işleýän hünärmen kompýuteriň öňünde uzak "
        "wagt oturýar, şonuň üçin iş otagynyň ýagtylygy, howa çalyşygy, elektrik goragy we "
        "iş wagtynyň dogry guralmagy zähmeti goramak nukdaýnazaryndan wajyp hasaplanýar."
    ),
]


POWER_TABLE = [
    ("Enjamyň ady", "Sany", "Biriniň kuwwaty, W", "Umumy kuwwat, W"),
    ("Kompýuter ýa-da noutbuk", "1", "180", "180"),
    ("Monitor", "1", "30", "30"),
    ("Wi-Fi router ýa-da test switch", "1", "20", "20"),
    ("LED ýagtylandyryjy panel", "7", "36", "252"),
    ("Jemi", "-", "-", "482"),
]


MICROCLIMATE_TABLE = [
    ("Görkeziji", "Kabul edilen baha", "Düşündiriş"),
    ("Otagyň ini A", "6 m", "Iş otagynyň meýilnamadaky ini"),
    ("Otagyň boýy B", "5 m", "Iş otagynyň meýilnamadaky boýy"),
    ("Otagyň beýikligi H", "3 m", "Potologyň beýikligi"),
    ("Işgärleriň sany n", "2 adam", "Programmany ulanýan we synag geçirýän işgärler"),
    ("Rugsat berilýän temperatura", "22-24 °C", "Kompýuter bilen işlemek üçin amatly aralyk"),
    ("Otnositel çyglylyk", "40-60 %", "Iş otagy üçin maslahat berilýän aralyk"),
]


LIGHT_TABLE = [
    ("Görkeziji", "Belgisi", "Baha"),
    ("Talap edilýän ýagtylandyryş", "E", "300 lx"),
    ("Otagyň meýdany", "S", "30 m²"),
    ("Ätiýaçlyk koeffisiýenti", "k", "1,4"),
    ("Ýagtylygyň deň düşmezlik koeffisiýenti", "z", "1,1"),
    ("Ýagtylyk akymynyň peýdalanma koeffisiýenti", "η", "0,6"),
    ("Bir LED paneliň ýagtylyk akymy", "Φ", "3600 lm"),
]


SCANNER_TABLE = [
    ("Görkeziji", "Baha", "Düşündiriş"),
    ("IP salgylaryň sany", "254", "192.168.0.0/24 lokal tor aralygy"),
    ("Barlanýan portlaryň sany", "15", "IoT gurluşlarynda köp duş gelýän portlar"),
    ("Bir port üçin timeout", "1 s", "TCP birikme synagynyň wagty"),
    ("Akym sany", "10", "Programmada orta güýç derejesi"),
    ("Yzygiderli barlag wagty", "3810 s", "254 · 15 · 1"),
    ("Köp akymly barlag wagty", "381 s", "3810 / 10"),
]


RECOMMENDATIONS = [
    "Kompýuter enjamlary ýer bilen goragly rozetka arkaly birikdirilmelidir.",
    "Elektrik setinde 6 A awtomatik goraýjy we 30 mA UZO ulanmak maslahat berilýär.",
    "Monitor bilen göz aralygy 60-70 sm bolmaly, ekranyň ýokarky bölegi gözüň derejesinden biraz pes ýerleşdirilmelidir.",
    "Her 60 minutlyk işden soň 5-10 minut dynç alyş arakesmesi berilmelidir.",
    "Iş otagynda howa çalyşygy we temperatura yzygiderli gözegçilikde saklanmalydyr.",
    "Skanirleme diňe rugsat berlen lokal torlarda geçirilmelidir.",
    "PDF hasabatlary elektron görnüşde saklamak kagyz sarp edilişini azaltmaga mümkinçilik berýär.",
]


def set_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
    return heading


def add_paragraph(document, text, first_line=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return paragraph


def add_formula(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    return paragraph


def add_table(document, rows, caption):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, text in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(6)
    run = caption_p.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True

    return table


def add_bullet(document, text):
    paragraph = document.add_paragraph(style=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("• " + text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def build_document():
    document = Document()
    set_styles(document)

    add_heading(document, TITLE, level=1)
    for text in INTRO_PARAGRAPHS:
        add_paragraph(document, text)

    add_heading(document, "1. Taslanýan iş ýeriniň umumy häsiýetnamasy", level=2)
    add_paragraph(
        document,
        "Programma üpjünçiligini işläp taýýarlamak we synagdan geçirmek üçin 6 m × 5 m "
        "ölçegli, beýikligi 3 m bolan iş otagy kabul edilýär. Otagda 2 işgär, 1 kompýuter, "
        "1 monitor, router ýa-da test switch, şeýle hem LED ýagtylandyryjy paneller ýerleşýär. "
        "Iş prosesi esasan akyl zähmeti bilen baglanyşykly bolup, fiziki agyrlyk pes derejede "
        "hasaplanýar.",
    )
    add_table(document, MICROCLIMATE_TABLE, "Tablisa 1 - Iş otagynyň başlangyç maglumatlary")

    add_heading(document, "2. Elektrik howpsuzlygynyň hasaplamasy", level=2)
    add_paragraph(
        document,
        "Kompýuter, monitor, tor enjamlary we ýagtylandyryjylar 220 V naprýaženiýeli elektrik "
        "setinden iýmitlenýär. Iş ýeriniň elektrik howpsuzlygyny bahalandyrmak üçin umumy kuwwat "
        "we elektrik togy hasaplanýar.",
    )
    add_table(document, POWER_TABLE, "Tablisa 2 - Iş ýerindäki elektrik enjamlarynyň kuwwaty")
    add_formula(document, "Pjemi = 180 + 30 + 20 + 252 = 482 W")
    add_formula(document, "I = P / U = 482 / 220 = 2,19 A")
    add_formula(document, "Ihasap = 1,25 · I = 1,25 · 2,19 = 2,74 A")
    add_paragraph(
        document,
        "Hasaplama boýunça iş ýeriniň umumy elektrik togy 2,19 A deňdir. Ätiýaçlyk "
        "koeffisiýenti 1,25 kabul edilende hasaplama togy 2,74 A bolýar. Şonuň üçin 6 A "
        "awtomatik goraýjy enjam bu iş ýeri üçin ýeterlikdir. Şeýle hem elektrik howpsuzlygyny "
        "ýokarlandyrmak üçin ýer bilen goragly rozetkalar we 30 mA duýgurlykly UZO ulanmak "
        "maslahat berilýär.",
    )

    add_heading(document, "3. Iş otagynyň howa çalyşygynyň hasaplamasy", level=2)
    add_paragraph(
        document,
        "Kompýuter enjamlary we ýagtylandyryjylar iş wagtynda ýylylyk bölüp çykarýar. Otagyň "
        "mikroklimatyny kadaly saklamak üçin howa çalyşygynyň zerur mukdaryny hasaplamak "
        "zerurdyr. Otagyň göwrümi aşakdaky formula boýunça kesgitlenilýär:",
    )
    add_formula(document, "V = A · B · H = 6 · 5 · 3 = 90 m³")
    add_paragraph(
        document,
        "Ýylylyk çeşmeleri hökmünde kompýuter we tor enjamlary 230 W, ýagtylandyryjylar "
        "252 W, iki işgär bolsa takmynan 2 · 100 W = 200 W ýylylyk bölüp çykarýar.",
    )
    add_formula(document, "Qjemi = 230 + 252 + 200 = 682 W")
    add_paragraph(
        document,
        "Ýylylygy aýyrmak üçin zerur howa mukdary aşakdaky formula boýunça hasaplanýar. "
        "Bu ýerde 0,335 - howanyň ýylylyk sygymy bilen baglanyşykly koeffisiýent, Δt bolsa "
        "içeri we daşarky howanyň temperatura tapawudy bolup, 5 °C kabul edilýär.",
    )
    add_formula(document, "L = Q / (0,335 · Δt) = 682 / (0,335 · 5) = 407,2 m³/sag")
    add_formula(document, "K = L / V = 407,2 / 90 = 4,52 1/sag")
    add_paragraph(
        document,
        "Hasaplama netijesine görä, iş otagynyň howasy bir sagatda takmynan 4,52 gezek "
        "çalyşmalydyr. Bu görkeziji kompýuter enjamlary bilen işleýän kiçi laboratoriýa "
        "otagy üçin amatly hasaplanýar. Şeýle hem işgärleriň sany boýunça sanitariýa talaby "
        "Lsan = 60 · 2 = 120 m³/sag bolýar. L = 407,2 m³/sag bu talapdan uly bolany üçin "
        "howa çalyşygynyň kabul edilen derejesi ýeterlikdir.",
    )

    add_heading(document, "4. Iş otagynyň ýagtylandyrylyşynyň hasaplamasy", level=2)
    add_paragraph(
        document,
        "Kompýuter bilen işlenilýän otaglarda gözüň ýadamagynyň öňüni almak üçin ýagtylandyryş "
        "derejesi kadaly bolmalydyr. Iş otagy üçin talap edilýän ýagtylandyryş E = 300 lx "
        "diýip kabul edilýär.",
    )
    add_table(document, LIGHT_TABLE, "Tablisa 3 - Ýagtylandyrylyş hasaplamasy üçin başlangyç maglumatlar")
    add_formula(document, "N = (E · S · k · z) / (Φ · η)")
    add_formula(document, "N = (300 · 30 · 1,4 · 1,1) / (3600 · 0,6) = 13860 / 2160 = 6,42")
    add_paragraph(
        document,
        "Hasaplama boýunça zerur bolan LED panelleriň sany 6,42 deňdir. Bu baha ýokary tarapa "
        "tegeleklenip, N = 7 sany LED panel kabul edilýär. Her paneliň kuwwaty 36 W bolanda "
        "ýagtylandyryşyň umumy kuwwaty Pýagtylyk = 7 · 36 = 252 W bolýar. Şeýle ýagtylandyryş "
        "iş stolunyň üstünde 300 lx talaby üpjün etmäge mümkinçilik berýär.",
    )

    add_heading(document, "5. Energiýa sarp edilişi we ekologik baha", level=2)
    add_paragraph(
        document,
        "Programma toplumynyň ekologik täsirini bahalandyrmak üçin iş ýeriniň ýyllyk elektrik "
        "energiýa sarp edilişi hasaplanýar. Iş wagty günde 6 sagat, ýylda 250 iş güni diýip "
        "kabul edilýär.",
    )
    add_formula(document, "Eýyl = Pjemi · t · D = 0,482 · 6 · 250 = 723 kW·sag/ýyl")
    add_paragraph(
        document,
        "Elektrik energiýasynyň şertli CO₂ çykaryş koeffisiýenti 0,45 kg/kW·sag diýip kabul "
        "edilende ýyllyk uglerod zyňyndysy aşakdaky ýaly hasaplanýar:",
    )
    add_formula(document, "MCO₂ = Eýyl · 0,45 = 723 · 0,45 = 325,35 kg CO₂/ýyl")
    add_paragraph(
        document,
        "Energiýa tygşytlaýjy režim ulanylanda kompýuteriň ortaça kuwwatyny peseltmek, LED "
        "ýagtylandyryşdan peýdalanmak we enjamlar ulanylmaýan wagtynda olary öçürmek arkaly "
        "ýyllyk sarp edilişi takmynan 480 kW·sag derejesine çenli azaltmak mümkin. Bu ýagdaýda "
        "tygşytlanýan energiýa 723 - 480 = 243 kW·sag/ýyl, CO₂ zyňyndysynyň azalmagy bolsa "
        "243 · 0,45 = 109,35 kg CO₂/ýyl bolar.",
    )
    add_paragraph(
        document,
        "Programmada hasabatlaryň PDF görnüşinde döredilmegi kagyz sarp edilişini hem azaltýar. "
        "Eger bir barlag boýunça 12 sahypalyk hasabat çap edilse we ýylda 30 barlag geçirilse, "
        "onda 12 · 30 = 360 sahypa kagyz sarp edilýär. Elektron hasabat ulanmak bu kagyz "
        "sarp edilişini aradan aýyrýar.",
    )

    add_heading(document, "6. Programma skanirlemesiniň öndürijilik hasaplamasy", level=2)
    add_paragraph(
        document,
        "IoT Security Scanner köp akymly skanirleme usulyny ulanýar. Bu usul lokal tordaky "
        "IP salgylary we portlary yzygiderli däl-de, birnäçe akymda barlamaga mümkinçilik berýär. "
        "Şeýlelikde, skanirleme wagty ep-esli azalýar.",
    )
    add_table(document, SCANNER_TABLE, "Tablisa 4 - Skanirleme wagtynyň hasaplamasy")
    add_formula(document, "Tyzygiderli = 254 · 15 · 1 = 3810 s = 63,5 min")
    add_formula(document, "Tköp akymly = 3810 / 10 = 381 s = 6,35 min")
    add_formula(document, "η = Tyzygiderli / Tköp akymly = 3810 / 381 = 10")
    add_paragraph(
        document,
        "Hasaplama boýunça 10 akym ulanylanda skanirleme wagty takmynan 10 esse azalýar. "
        "Amaly ýagdaýda toruň tizligi, hostlaryň jogap bermek wagty we açyk portlaryň sany "
        "netijä täsir edip biler, emma köp akymly çemeleşme programma toplumynyň netijeliligini "
        "ýokarlandyrýar."
    )

    add_heading(document, "7. Ergonomika we howpsuz iş düzgünleri", level=2)
    add_paragraph(
        document,
        "Programma bilen işlemek kompýuteriň öňünde uzak wagt oturmagy talap edýär. Şonuň üçin "
        "iş ýeriniň ergonomiki taýdan dogry guralmagy işgärleriň saglygyny goramakda möhüm orun "
        "tutýar. Aşakdaky çäreleriň ýerine ýetirilmegi maslahat berilýär:"
    )
    for recommendation in RECOMMENDATIONS:
        add_bullet(document, recommendation)

    add_heading(document, "Netije", level=2)
    add_paragraph(
        document,
        "Geçirilen hasaplamalar IoT Security Scanner programma toplumyny işläp taýýarlamak we "
        "ulanmak üçin niýetlenen iş ýeriniň elektrik howpsuzlygy, mikroklimat, ýagtylandyryş we "
        "energiýa sarp edilişi boýunça kabul edilen talaplara laýyk gelýändigini görkezýär. "
        "Programma lokal toruň howpsuzlyk ýagdaýyny awtomatik seljermek bilen administratoryň "
        "wagtyny tygşytlaýar, PDF görnüşindäki elektron hasabatlar bolsa kagyz sarp edilişini "
        "azaltmaga mümkinçilik berýär."
    )

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_FILE)
