from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOCX_FILE = Path("bolum_1_1_iot_gurlushlary.docx")
BACKUP_FILE = Path("bolum_1_1_iot_gurlushlary_before_giris.docx")


TITLE = "Giriş"


GIRIS_PARAGRAPHS = [
    (
        "Derňewiň obýekti: “Lokal tordaky IoT gurluşlarynyň howpsuzlygyny "
        "awtomatlaşdyrylan skanirleme, açyk portlary ýüze çykarmak we "
        "gowşaklyklar boýunça maslahat bermek ulgamy”."
    ),
    (
        "Işiň maksady: lokal torda ýerleşýän IoT gurluşlaryny awtomatik tapýan, "
        "olaryň açyk hyzmatlaryny we portlaryny kesgitleýän, potensial howpsuzlyk "
        "töwekgelçiliklerini seljerýän hem-de ulanyja hasabat we maslahat berýän "
        "IoT Security Scanner programma toplumyny taslamak we işläp taýýarlamak."
    ),
    (
        "Wajyplygy. Häzirki döwürde sanly tehnologiýalaryň ösmegi bilen internete "
        "birikdirilýän gurluşlaryň sany ýylsaýyn artýar. IP kameralar, akylly öý "
        "enjamlary, datçikler, marşrutizatorlar, giriş gözegçilik ulgamlary we "
        "beýleki IoT gurluşlary gündelik durmuşda, bilim edaralarynda, saglygy "
        "goraýyşda, senagatda we hyzmatlar ulgamynda giňden ulanylýar. Şeýle "
        "gurluşlaryň sanynyň artmagy tor infrastrukturasynyň mümkinçiliklerini "
        "giňeldýär, emma şol bir wagtda täze howpsuzlyk meselelerini hem döredýär."
    ),
    (
        "IoT gurluşlarynyň köpüsinde çäkli hasaplaýyş serişdeleri, ýönekeý "
        "operasion ulgamlar we minimal gorag mehanizmleri bolýar. Käbir ýagdaýlarda "
        "ulanyjylar standart login we parollary üýtgetmeýärler, firmware "
        "täzelenmelerini wagtynda ýerine ýetirmeýärler ýa-da Telnet, HTTP, RTSP, "
        "MQTT ýaly hyzmatlary goragsyz ýagdaýda açyk goýýarlar. Netijede, şeýle "
        "gurluşlar hüjümçiler üçin tora giriş nokady bolup biler."
    ),
    (
        "IoT howpsuzlygynyň pes derejesi diňe aýratyn bir gurluşyň işleýşine däl, "
        "eýsem tutuş lokal toruň durnuklylygyna täsir edip biler. Mysal üçin, "
        "gowşak parolly IP kamera ýa-da açyk Telnet hyzmaty bolan enjam arkaly "
        "hüjümçi tora girip, beýleki ulgamlara täsir edip biler. Şonuň üçin tor "
        "administratorynyň IoT gurluşlaryny wagtynda ýüze çykarmagy, olaryň açyk "
        "portlaryny barlamagy we ýüze çykarylan töwekgelçilikleri seljermegi möhüm "
        "meseleleriň biri bolup durýar."
    ),
    (
        "Amaly gymmatlygy: işlenip taýýarlanylýan programma lokal toruň başlangyç "
        "howpsuzlyk auditini geçirmekde ulanylyp bilner. Programma IP aralygy "
        "boýunça gurluşlary tapýar, olaryň MAC salgysyny, adyny, açyk portlaryny "
        "we hyzmatlaryny görkezýär. Mundan başga-da, plugin ulgamynyň kömegi bilen "
        "Telnet gowşak autentifikasiýasy, açyk web-interfeýsler we IoT gurluşlaryna "
        "mahsus beýleki töwekgelçilikler boýunça maslahat berilýär."
    ),
    (
        "Işiň dowamynda Python programmirleme dili, PyQt6 grafiki interfeýs "
        "kitaphanasy, SQLite maglumat bazasy, Nmap skanirleme guraly, ReportLab "
        "arkaly PDF hasabat döretmek we plugin esasly gowşaklyk barlagy ýaly "
        "tehnologiýalar ulanylýar. Bu tehnologiýalar programma toplumyny modully, "
        "giňeldip bolýan we ulanyjy üçin amatly görnüşde döretmäge mümkinçilik berýär."
    ),
    (
        "Garaşylýan netijeler: lokal tordaky IoT gurluşlaryny tapmak, açyk portlary "
        "we hyzmatlary kesgitlemek, töwekgelçilik derejesini görkezmek, skanirleme "
        "netijelerini maglumat bazasynda saklamak, PDF görnüşinde hasabat döretmek "
        "we ulanyja howpsuzlyk boýunça amaly maslahat bermek mümkinçilikleriniň "
        "döredilmegi göz öňünde tutulýar."
    ),
    (
        "Teklip: IoT Security Scanner programma toplumyny okuw edaralarynyň, kiçi "
        "kärhanalaryň, ofisleriň we öý torlarynyň başlangyç howpsuzlyk barlagynda "
        "ulanmak maksadalaýykdyr. Programma tor administratoryna ýa-da ulanyja "
        "gurluşlaryň ýagdaýy barada tiz maglumat almaga, howply açyk hyzmatlary "
        "ýüze çykarmaga we howpsuzlyk çärelerini meýilleşdirmäge ýardam berer."
    ),
]


def set_document_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True


def add_heading(document, text):
    heading = document.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True


def add_text_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def copy_old_body(new_document, old_document):
    new_body = new_document.element.body
    old_body = old_document.element.body

    for element in old_body:
        if element.tag.endswith("sectPr"):
            continue
        new_body.append(deepcopy(element))


def has_giris(document):
    return any(paragraph.text.strip() == TITLE for paragraph in document.paragraphs)


def prepend_giris():
    old_document = Document(DOCX_FILE)

    if has_giris(old_document):
        print(f"{TITLE} eýýäm dokumentiň başynda bar")
        return

    if not BACKUP_FILE.exists():
        BACKUP_FILE.write_bytes(DOCX_FILE.read_bytes())

    new_document = Document()
    set_document_styles(new_document)

    add_heading(new_document, TITLE)
    for text in GIRIS_PARAGRAPHS:
        add_text_paragraph(new_document, text)

    new_document.add_page_break()
    copy_old_body(new_document, old_document)
    new_document.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == "__main__":
    prepend_giris()
