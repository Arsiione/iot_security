from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOCX_FILE = Path("bolum_1_1_iot_gurlushlary.docx")
BACKUP_FILE = Path("bolum_1_1_iot_gurlushlary_before_sozbasy.docx")


TITLE = "Sözbaşy"


SOZBASY_PARAGRAPHS = [
    (
        "Häzirki döwürde dünýäde sanly tehnologiýalaryň çalt depginler bilen ösmegi "
        "ykdysadyýetiň, bilimiň, saglygy goraýşyň, senagatyň we gündelik ýaşaýşyň "
        "ähli ugurlaryna düýpli täsir edýär. Kompýuter ulgamlary, maglumat torlary, "
        "bulut tehnologiýalary, emeli aň, awtomatlaşdyrylan dolandyryş ulgamlary we "
        "Internet of Things tehnologiýalary häzirki zaman jemgyýetiniň aýrylmaz bölegine "
        "öwrülýär. Şol sebäpli maglumat tehnologiýalarynyň netijeli we howpsuz ulanylmagy "
        "döwrebap ösüşiň möhüm şertleriniň biri bolup durýar."
    ),
    (
        "Sanly ulgamlaryň giňden ornaşdyrylmagy bilen birlikde maglumat howpsuzlygynyň "
        "ähmiýeti hem ýokarlanýar. Lokal torlara diňe kompýuterler däl, eýsem IP kameralar, "
        "akylly öý enjamlary, marşrutizatorlar, datçikler, giriş gözegçilik ulgamlary we "
        "beýleki IoT gurluşlary hem birikdirilýär. Bu gurluşlar maglumatlary ýygnamak, "
        "dolandyrmak we alyş-çalyş etmek üçin amatly mümkinçilik döredýär, emma olaryň "
        "gorag derejesi ýeterlik bolmasa, tutuş tor üçin howp çeşmesine öwrülip biler."
    ),
    (
        "Internet of Things gurluşlarynyň aýratynlygy olaryň köp görnüşli bolmagy, dürli "
        "öndürijiler tarapyndan taýýarlanylmagy we köplenç çäkli hasaplaýyş serişdelerine "
        "eýe bolmagydyr. Şeýle enjamlarda standart parollaryň ulanylmagy, firmware "
        "täzelenmeleriniň wagtynda geçirilmezligi, Telnet, HTTP, RTSP, MQTT ýaly hyzmatlaryň "
        "nädogry sazlanmagy ýa-da açyk portlaryň gözegçiliksiz galmagy ýaly meseleler ýüze "
        "çykyp bilýär. Bu ýagdaý IoT gurluşlarynyň howpsuzlygyny yzygiderli barlamagyň "
        "zerurdygyny görkezýär."
    ),
    (
        "Maglumat howpsuzlygyny üpjün etmek diňe gorag diwarlaryny ýa-da antivirus "
        "serişdelerini ulanmak bilen çäklenmeýär. Häzirki wagtda toruň içinde haýsy "
        "gurluşlaryň bardygyny, olaryň nähili hyzmatlary işledýändigini, haýsy portlarynyň "
        "açykdygyny we bu hyzmatlaryň nähili töwekgelçilik döredip biljekdigini anyklamak "
        "möhüm bolup durýar. Şonuň üçin awtomatlaşdyrylan skanirleme we başlangyç "
        "howpsuzlyk auditi ulgamlary tor administratorynyň işinde aýratyn orun eýeleýär."
    ),
    (
        "Bu diplom işinde lokal tordaky IoT gurluşlarynyň howpsuzlyk ýagdaýyny seljermek "
        "üçin IoT Security Scanner programma toplumyny işläp taýýarlamak meselesi "
        "garalýar. Programma IP aralygy boýunça gurluşlary tapmaga, olaryň açyk portlaryny "
        "we hyzmatlaryny kesgitlemäge, potensial gowşaklyklary ýüze çykarmaga, netijeleri "
        "maglumat bazasynda saklamaga we PDF görnüşinde hasabat taýýarlamaga mümkinçilik "
        "berýär."
    ),
    (
        "Taslamanyň esasy aýratynlygy onuň modully gurluşynda jemlenýär. Programma ýadrosy "
        "tor skanirlemesini ýerine ýetirýär, grafiki interfeýs ulanyjy bilen amatly "
        "aragatnaşygy üpjün edýär, plugin ulgamy bolsa dürli gowşaklyklary barlamak üçin "
        "giňeldilýän mümkinçilik döredýär. Şeýle çemeleşme geljekde täze IoT gurluşlaryny "
        "we täze howpsuzlyk barlaglaryny goşmaga şert döredýär."
    ),
    (
        "Diplom işiniň amaly ähmiýeti, işlenip taýýarlanylýan programma toplumynyň okuw "
        "edaralarynda, kiçi kärhanalarda, ofislerde we öý torlarynda başlangyç howpsuzlyk "
        "barlagyny geçirmek üçin ulanylyp bilinmegindedir. Programma administratora ýa-da "
        "ulanyja toruň ýagdaýy barada çalt maglumat berýär, açyk hyzmatlary görkezýär we "
        "howpsuzlyk çärelerini meýilleşdirmäge ýardam edýär."
    ),
    (
        "Şeýlelikde, IoT Security Scanner programma toplumynyň işlenip taýýarlanylmagy "
        "häzirki zaman maglumat tehnologiýalarynyň, IoT ulgamlarynyň we kiberhowpsuzlygyň "
        "talaplary bilen gönüden-göni baglanyşyklydyr. Bu taslama lokal torlaryň howpsuzlyk "
        "derejesini ýokarlandyrmak, IoT gurluşlarynyň ýagdaýyna gözegçilik etmek we "
        "howpsuzlyk auditi işlerini awtomatlaşdyrmak üçin ähmiýetli çözgüt bolup durýar."
    ),
]


def set_document_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True


def add_heading(document, text):
    heading = document.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True


def add_text_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def copy_old_body(new_document, old_document):
    new_body = new_document.element.body
    old_body = old_document.element.body

    for element in old_body:
        if element.tag.endswith("sectPr"):
            continue
        new_body.append(deepcopy(element))


def has_sozbasy(document):
    return any(paragraph.text.strip() == TITLE for paragraph in document.paragraphs)


def prepend_sozbasy():
    old_document = Document(DOCX_FILE)

    if has_sozbasy(old_document):
        print(f"{TITLE} eýýäm dokumentde bar")
        return

    if not BACKUP_FILE.exists():
        BACKUP_FILE.write_bytes(DOCX_FILE.read_bytes())

    new_document = Document()
    set_document_styles(new_document)

    add_heading(new_document, TITLE)
    for text in SOZBASY_PARAGRAPHS:
        add_text_paragraph(new_document, text)

    new_document.add_page_break()
    copy_old_body(new_document, old_document)
    new_document.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == "__main__":
    prepend_sozbasy()
