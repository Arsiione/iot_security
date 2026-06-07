from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_2_1.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "2. Programma toplumynyň arhitekturasyny taslamak", level=1)
    add_normal(
        doc,
        "Birinji bölümde IoT gurluşlarynyň howpsuzlyk meseleleri, hüjüm wektorlary we bar bolan "
        "skanirleme ulgamlary seljerildi. Şol seljermäniň esasynda ikinji bölümde IoT Security "
        "Scanner programma toplumynyň arhitekturasy, tehnologiýalar toplumyny saýlamagyň sebäpleri, "
        "modul gurluşy, maglumatlar bazasynyň shemasy we skanirleme algoritminiň umumy işleýşi "
        "taslanylýar. Bu bölüm programma döredilmezden öň onuň nähili böleklerden durmalydygyny "
        "we şol bölekleriň özara nähili baglanyşmalydygyny düşündirýär."
    )

    add_heading(doc, "2.1. Tehnologiýalar toplumyny saýlamagyň esaslandyrylyşy", level=2)
    paragraphs = [
        (
            "IoT Security Scanner programma toplumyny işläp taýýarlamakda esasy talaplar hökmünde "
            "lokal torda durnukly işlemek, Windows gurşawynda ulanylyp bilinmek, grafiki interfeýsiň "
            "bolmagy, skanirleme netijelerini maglumatlar bazasynda saklamak, PDF görnüşinde hasabat "
            "döretmek we geljekde täze gowşaklyk barlaglaryny goşmak mümkinçiligi kabul edildi. "
            "Şu talaplardan ugur alnyp, taslama üçin Python dili, PyQt6 kitaphanasy, SQLite maglumatlar "
            "bazasy, Nmap skanirleme guraly we goşmaça kömekçi kitaphanalar saýlanyp alyndy."
        ),
        (
            "Python programmirleme dili bu taslama üçin amatlydyr, sebäbi ol tor bilen işlemek, faýl "
            "ulgamy, maglumatlar bazasy, daşky prosesleri işletmek we köp akymly ýerine ýetiriş ýaly "
            "meseleleri ýönekeý görnüşde çözmäge mümkinçilik berýär. Python dilinde socket, subprocess, "
            "ipaddress, threading ýaly standart mümkinçilikler bar, şeýle hem Nmap, PDF hasabat, grafiki "
            "interfeýs we SSH aragatnaşygy üçin taýýar kitaphanalar giňden elýeterlidir. Bu bolsa diplom "
            "taslamasynyň funksional böleklerini çalt we düşnükli gurmaga şert döredýär."
        ),
        (
            "Grafiki interfeýs üçin PyQt6 saýlanyldy. PyQt6 Windows operasion ulgamynda doly derejeli "
            "desktop programma döretmäge mümkinçilik berýär. Programma web sahypa däl-de, aýratyn EXE "
            "görnüşinde işledilýär, bu bolsa diplom goragy wagtynda internet ýa-da serwer gurşawyna "
            "bagly bolmazdan demonstrasiýa geçirmäge amatlydyr. PyQt6 arkaly sahypalar, düwmeler, "
            "tablisa, progress bar, saýlaw sanawlary we tema sazlamalary döredildi. Şeýlelikde, "
            "skanirleme netijeleri ulanyja diňe tekst görnüşinde däl, eýsem düşnükli interfeýs arkaly "
            "ýetirilýär."
        ),
        (
            "Skanirleme ýadrosy hökmünde Nmap ulanylýar. Nmap tor howpsuzlygynda köp ýyllardan bäri "
            "ulanylýan ygtybarly gural bolup, işjeň hostlary, açyk portlary we käbir hyzmat alamatlaryny "
            "kesgitlemäge mümkinçilik berýär. Diplom taslamasynda Nmap gönüden-göni ulanyjynyň öňünde "
            "komanda setiri hökmünde görkezilmeýär; programma onuň mümkinçiliklerini fon režiminde "
            "işledýär we netijeleri grafiki tablisa, töwekgelçilik derejesi hem-de maslahat görnüşinde "
            "berýär. Şeýlelikde, güýçli professional guralyň netijesi has düşnükli ulanyjy interfeýsi "
            "bilen birleşdirilýär."
        ),
        (
            "Maglumatlary saklamak üçin SQLite saýlanyldy. SQLite aýratyn serwer talap etmeýän ýeňil "
            "maglumatlar bazasy bolup, bir faýlyň içinde skanirleme taryhyny, tapylan gurluşlary we "
            "netijeleri saklamaga mümkinçilik berýär. Diplom taslamasy üçin bu çözgüt has amatlydyr, "
            "sebäbi programma bir kompýuterde işledilýär we goşmaça maglumatlar bazasy serwerini "
            "gurnamak zerurlygy ýok. SQLite arkaly skanirleme netijeleri soňra Panel, Netijeler we Taryh "
            "sahypalarynda gaýtadan görkezilýär."
        ),
        (
            "Hasabat döretmek üçin ReportLab kitaphanasy ulanylýar. Bu kitaphana PDF faýllaryny "
            "programmatiki görnüşde döretmäge mümkinçilik berýär. Programma skanirleme netijelerini "
            "resmileşdirilen hasabat görnüşinde eksport edýär, bu bolsa diplom goragy we amaly audit "
            "üçin möhüm aýratynlykdyr. Matplotlib kitaphanasy bolsa diagramma we statistiki şekilleri "
            "döretmek üçin ulanylyp bilner."
        ),
        (
            "Taslamada Paramiko kitaphanasy uzakdan SSH arkaly düzediş amallaryny ýerine ýetirmek "
            "mümkinçiligi üçin göz öňünde tutulýar. Bu remediation ugrunyň başlangyç binýadydyr. "
            "Requests kitaphanasy web hyzmatlary ýa-da HTTP esasly barlaglar üçin ulanylyp bilner. "
            "PyInstaller bolsa programma toplumyny ulanyjy üçin amatly EXE faýlyna ýygnamaga mümkinçilik "
            "berýär. Şeýlelikde, saýlanan tehnologiýalaryň her biri taslamanyň aýratyn bir talabyny "
            "ýerine ýetirýär."
        ),
    ]
    for text in paragraphs:
        add_normal(doc, text)

    caption = doc.add_paragraph("Tablisa 2.1 - Taslamada ulanylýan tehnologiýalar we olaryň wezipesi")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    headers = ["Tehnologiýa", "Taslamadaky wezipesi", "Saýlanyp alynmagynyň sebäbi"]
    rows = [
        [
            "Python",
            "Programma ýadrosy, skanirleme logikasy, faýl we proses dolandyryşy",
            "Ýönekeý sintaksis, giň kitaphana ekoulgamy we tor bilen işlemek üçin amatlylyk",
        ],
        [
            "PyQt6",
            "Grafiki interfeýs, sahypalar, tablisa, düwmeler we tema",
            "Windows üçin doly desktop programma döretmäge mümkinçilik berýär",
        ],
        [
            "Nmap",
            "Hostlary we portlary skanirlemek",
            "Tor howpsuzlygynda ygtybarly we giňden ulanylýan professional gural",
        ],
        [
            "SQLite",
            "Skanirleme taryhyny we gurluş maglumatlaryny saklamak",
            "Aýratyn serwer talap etmeýär, lokal programma üçin ýeňil we amatly",
        ],
        [
            "ReportLab",
            "PDF hasabatlaryny döretmek",
            "Netijeleri resmileşdirilen görnüşde eksport etmäge mümkinçilik berýär",
        ],
        [
            "Matplotlib",
            "Statistika we diagramma görkezmek",
            "Panel we hasabat sahypalarynda maglumatlary wizual görkezmek üçin amatly",
        ],
        [
            "Paramiko",
            "SSH arkaly uzakdan düzediş mümkinçiligi",
            "Remediation engine üçin howpsuz uzakdan dolandyryş binýady bolup hyzmat edýär",
        ],
        [
            "PyInstaller",
            "Programmany EXE görnüşinde ýygnamak",
            "Diplom goragynda programmany aýratyn faýl hökmünde görkezmäge mümkinçilik berýär",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)

    add_normal(
        doc,
        "Saýlanan tehnologiýalar toplumynyň esasy artykmaçlygy olaryň biri-birini doldurmagyndadyr. "
        "Python programma logikasyny üpjün edýär, PyQt6 ulanyjy bilen amatly aragatnaşyk döredýär, "
        "Nmap skanirleme mümkinçiliklerini berýär, SQLite maglumatlary saklaýar, ReportLab hasabat "
        "taýýarlaýar, Paramiko bolsa geljekde awtomatik düzediş amallaryny giňeltmäge şert döredýär. "
        "Şeýle arhitektura diplom taslamasynyň hem nazary, hem amaly maksatlaryna laýyk gelýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("2.1.") for p in main_doc.paragraphs):
        print("Section 2.1 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
