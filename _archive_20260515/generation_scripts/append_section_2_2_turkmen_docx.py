from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_2_2.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "2.2. Programma arhitekturasy we modul gurluşy")

    paragraphs = [
        (
            "IoT Security Scanner programma toplumy modul arhitektura ýörelgesi boýunça taslanyldy. "
            "Bu çemeleşmede programma bir uly we çylşyrymly faýl hökmünde däl-de, özbaşdak jogapkärçiligi "
            "bolan birnäçe bölege bölünýär. Şeýle gurluş programmany düşünmegi, giňeltmegi, synagdan "
            "geçirmegi we geljekde täze mümkinçilikleri goşmagy ýeňilleşdirýär. Taslamanyň umumy gurluşy "
            "Core, UI, Plugins, Database, Reports we Utils böleklerinden durýar."
        ),
        (
            "Programmany işe girizmek main.py faýlyndan başlanýar. Bu faýl PyQt6 programmasynyň başlangyç "
            "nokady bolup, esasy penjire bolan MainWindow synpyny döredýär. MainWindow interfeýsiň daşky "
            "gurluşyny, çep tarapdaky menýuny, sahypalaryň çalşylmagyny we tema sazlamalaryny dolandyrýar. "
            "Şeýlelikde, main.py programma giriş nokady, ui/main_window.py bolsa interfeýsiň esasy karkasy "
            "hökmünde çykyş edýär."
        ),
        (
            "UI moduly ulanyjy bilen programma ýadrosynyň arasyndaky aragatnaşygy üpjün edýär. ui/pages.py "
            "faýlynda Baş sahypa, Skanirleme, Panel, Netijeler, Taryh we Sazlamalar sahypalarynyň düzülişi "
            "ýerleşýär. Skanirleme sahypasynda ulanyjy tor adapterini saýlaýar, IP aralygyny görýär, "
            "skanirleme derejesini belleýär we netijeleri tablisa görnüşinde alýar. ui/style.qss faýly bolsa "
            "programma interfeýsiniň wizual stilini, reňklerini, düwmelerini we tablisa görnüşini kesgitleýär."
        ),
        (
            "Core moduly programma toplumynyň esasy iş logikasyny öz içine alýar. Bu moduldaky iň möhüm "
            "faýl core/scanner.py bolup, ol tor adapterlerini kesgitleýär, dogry IP aralygyny saýlaýar, "
            "Nmap, ARP, Ping we Windows ARP keşi arkaly gurluşlary tapýar, portlary barlaýar, hyzmatlary "
            "tanadýar, gurluşyň görnüşini kesgitleýär we töwekgelçilik derejesini hasaplaýar. ScanThread "
            "synpynyň aýratyn akymda işlemegi sebäpli skanirleme wagtynda interfeýs doňup galmaýar."
        ),
        (
            "core/remediation_engine.py faýly gowşaklyklary azaltmak boýunça düzediş amallarynyň binýadyny "
            "emele getirýär. Bu modul Telnet, FTP, HTTP ýa-da SSH bilen bagly käbir howpsuzlyk meselelerine "
            "garşy çäreleri ýerine ýetirmek üçin niýetlenendir. Häzirki taslamada remediation logikasy "
            "esasan maslahat we başlangyç düzediş modeli hökmünde ulanylýar, emma arhitektura geljekde "
            "awtomatik SSH arkaly sazlama üýtgetmeleri goşmaga mümkinçilik berýär."
        ),
        (
            "Plugins moduly aýratyn gowşaklyk barlaglaryny programma ýadrosyndan bölüp saklamak üçin "
            "ulanylýar. plugins/hikvision.py faýly IP kamera we Hikvision görnüşli gurluşlar bilen bagly "
            "alamatlary barlamak üçin niýetlenendir. plugins/telnet_weak_auth.py bolsa Telnet hyzmatynda "
            "gowşak autentifikasiýa alamatlaryny barlaýar. Plugin çemeleşmesi programma täze barlaglary "
            "goşmagy aňsatlaşdyrýar: täze gowşaklyk üçin aýratyn plugin döredilýär we scanner.py ony "
            "netijeleri bahalandyrmakda ulanyp bilýär."
        ),
        (
            "Maglumatlar bazasy bilen işlemek database.py faýly arkaly ýerine ýetirilýär. Bu modul SQLite "
            "bazasynda skanirleme netijelerini, gurluşlaryň IP we MAC maglumatlaryny, portlaryny, "
            "töwekgelçilik derejesini, tapylan usulyny we maslahatlaryny saklaýar. load_history, save_scan "
            "we get_scan_stats ýaly funksiýalar Panel, Netijeler we Taryh sahypalarynyň maglumat almak "
            "mümkinçiligini üpjün edýär."
        ),
        (
            "Hasabat ulgamy reports.py faýlynda ýerleşýär. Bu modul ReportLab kitaphanasynyň kömegi bilen "
            "skanirleme netijelerini PDF görnüşinde resmileşdirýär. Hasabatda tapylan gurluşlar, olaryň "
            "portlary, öndürijisi, görnüşi, tapylan usuly, töwekgelçilik derejesi we maslahatlar görkezilýär. "
            "Bu mümkinçilik diplom goragy wagtynda netijeleri resmi dokument görnüşinde görkezmäge amatlydyr."
        ),
        (
            "Kömekçi modul hökmünde utils/notifications.py we notifications.py faýllary ulanylýar. Bu "
            "faýllar email ýa-da beýleki habar beriş mehanizmlerini goşmak üçin niýetlenendir. Şeýle modul "
            "gurluşy programmadaky esasy skanirleme logikasyny habar beriş logikasyndan aýyrýar we geljekde "
            "duýduryş ulgamyny aýratyn ösdürmäge mümkinçilik berýär."
        ),
        (
            "Umumy arhitekturada maglumat akymy şeýle işleýär: ulanyjy interfeýsde skanirlemäni başlaýar, "
            "ui/pages.py ScanThread synpyny döredýär, ScanThread core/scanner.py içindäki algoritmleri "
            "işledýär, tapylan gurluşlar signallar arkaly interfeýse iberilýär, skanirleme tamamlanandan soň "
            "netijeler database.py arkaly SQLite bazasyna ýazylýar, soňra Panel, Netijeler we Taryh sahypalary "
            "şol maglumatlary görkezýär. Ulanyjy zerur bolan ýagdaýynda reports.py arkaly PDF hasabatyny "
            "döredip bilýär."
        ),
    ]
    for text in paragraphs:
        add_normal(doc, text)

    caption = doc.add_paragraph("Tablisa 2.2 - Programma modul gurluşynyň beýany")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    headers = ["Modul ýa-da faýl", "Esasy wezipesi", "Programmadaky ähmiýeti"]
    rows = [
        [
            "main.py",
            "Programmany işe girizmek we esasy penjire döretmek",
            "Programma giriş nokady hökmünde çykyş edýär",
        ],
        [
            "ui/main_window.py",
            "Esasy penjire, çep menýu, sahypa çalşygy we tema",
            "Ulanyjy interfeýsiniň umumy karkasyny döredýär",
        ],
        [
            "ui/pages.py",
            "Baş sahypa, Skanirleme, Panel, Netijeler, Taryh we Sazlamalar",
            "Ulanyjynyň ähli esasy iş hereketlerini üpjün edýär",
        ],
        [
            "ui/style.qss",
            "Interfeýsiň reňkleri, tablisa we düwme stilleri",
            "Programmany bitewi wizual görnüşe getirýär",
        ],
        [
            "core/scanner.py",
            "Adapter saýlamak, host tapmak, port barlamak, töwekgelçiligi kesgitlemek",
            "Programma toplumynyň esasy skanirleme ýadrosy bolup durýar",
        ],
        [
            "core/remediation_engine.py",
            "Gowşaklyklary azaltmak boýunça düzediş logikasy",
            "Remediation ugruny giňeltmek üçin binýat döredýär",
        ],
        [
            "plugins/hikvision.py",
            "IP kamera we Hikvision bilen bagly alamatlary barlamak",
            "IoT kamera howpsuzlygy üçin ýörite plugin bolup durýar",
        ],
        [
            "plugins/telnet_weak_auth.py",
            "Telnet gowşak autentifikasiýasyny barlamak",
            "Goragsyz Telnet hyzmatynyň töwekgelçiligini ýüze çykarmaga kömek edýär",
        ],
        [
            "database.py",
            "SQLite bazasy, skanirleme taryhy we statistika",
            "Netijeleri saklamak we gaýtadan görkezmek üçin ulanylýar",
        ],
        [
            "reports.py",
            "PDF hasabatlaryny döretmek",
            "Audit netijelerini resmi görnüşde eksport edýär",
        ],
        [
            "utils/notifications.py",
            "Habar beriş we duýduryş logikasy",
            "Geljekde email/webhook/SMS duýduryşlaryny giňeltmäge mümkinçilik berýär",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)

    add_normal(
        doc,
        "Bu modul gurluşy programma toplumyny tehniki taýdan durnukly we giňeldip bolýan edýär. "
        "Mysal üçin, geljekde täze IoT gowşaklygy barlanmaly bolsa, scanner.py faýlyny düýpli üýtgetmän, "
        "plugins bukjasyna täze plugin goşmak mümkin. Şeýle hem interfeýs sahypalary aýratynlykda "
        "ösdürilip, maglumatlar bazasy ýa-da hasabat ulgamy bilen baglanyşyk saklap bilýär. Bu bolsa "
        "taslamanyň arhitekturasynyň amaly taýdan dogry gurlandygyny görkezýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_style(main_doc):
    for paragraph in main_doc.paragraphs:
        if paragraph.text.strip().startswith("2.2."):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)
            break


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("2.2.") for p in main_doc.paragraphs):
        print("Section 2.2 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_style(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
