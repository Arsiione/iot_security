from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
MAIN_DOC = ROOT / "kakabalowa.docx"
OUT_DOC = ROOT / "ykdysady_iot_security_yjd_uytgedilen.docx"


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def format_runs(paragraph, size: int = 12, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style_name = f"Heading {level}"
    style_names = {style.name for style in doc.styles}
    paragraph = doc.add_paragraph(text, style=style_name if style_name in style_names else None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    format_runs(paragraph, size=14 if level == 1 else 13, bold=True)


def add_p(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    format_runs(paragraph, size=12)


def add_formula(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    format_runs(paragraph, size=12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                format_runs(paragraph, size=11, bold=(r_idx == 0))


def build_section_doc(template: Path | None = None) -> Document:
    doc = Document(str(template)) if template and template.exists() else Document()
    clear_doc(doc)

    add_heading(doc, "Taslamanyň ykdysady netijeliliginiň hasaplamasy", 1)

    add_heading(doc, "Maýa goýum taslamalarynyň netijeliligini bahalandyrmagyň esasy ýörelgeleri", 2)
    add_p(
        doc,
        "Islendik programma üpjünçiligi taslamasynyň ykdysady taýdan esaslydygyny görkezmek üçin "
        "onuň döredilmegine sarp edilýän serişdeler we taslamadan alynýan peýda deňeşdirilýär. "
        "Bu bölümde IoT Security Scanner programma toplumyny işläp taýýarlamagyň çykdajylary, "
        "zähmet talap edişi, tehniki serişdeler we programma arkaly tygşytlanýan iş wagty hasaplanylýar.",
    )
    add_p(
        doc,
        "Taslamanyň ykdysady netijeliligi diňe göni girdeji bilen ölçenmeýär. Maglumat howpsuzlygy "
        "ulgamlarynda wagty tygşytlamak, tor barlagyny standartlaşdyrmak, hasabat taýýarlamagy "
        "çaltlandyrmak we howpsuzlyk töwekgelçiliklerini ir ýüze çykarmak hem ykdysady peýda hökmünde "
        "bahalandyrylýar.",
    )
    add_p(
        doc,
        "IoT Security Scanner lokal Wi-Fi ýa-da Ethernet torundaky enjamlary tapýar, açyk portlary "
        "kesgitleýär, töwekgelçilik derejesini görkezýär we PDF görnüşli hasabat taýýarlaýar. Şeýlelikde, "
        "programma tor administratorynyň el bilen ýerine ýetirýän gaýtalanýan işlerini azaldýar.",
    )
    add_p(doc, "Netijeliligi bahalandyrmakda aşakdaky ýörelgeler göz öňünde tutulýar:")
    for item in [
        "taslamanyň ähli çykdajylaryny doly hasaba almak;",
        "el bilen audit we programma arkaly audit ýagdaýlaryny deňeşdirmek;",
        "açyk çeşmeli tehnologiýalaryň çykdajylary azaltmakdaky ornuny görkezmek;",
        "wagt faktoryny we işgärleriň zähmet hakyny hasaba almak;",
        "taslamanyň töwekgelçiliklerini we çäklendirmelerini görkezmek.",
    ]:
        add_p(doc, item)

    add_heading(doc, "Taslamanyň netijeliligini hasaplamak üçin başlangyç maglumatlar", 2)
    add_p(
        doc,
        "Hasaplamalar üçin başlangyç maglumatlar hökmünde programma toplumyny işläp taýýarlamaga "
        "gatnaşýan işgärleriň sany, olaryň sagatlaýyn tölegi, iş üçin sarp edilýän wagt, tehniki "
        "serişdeleriň bahasy we programma ulanylanda tygşytlanýan audit wagty kabul edildi.",
    )
    add_p(
        doc,
        "Programma Python dilinde işlenip taýýarlanylýar. Grafiki interfeýs üçin PyQt6, maglumatlary "
        "saklamak üçin SQLite, tor skanirlemesi üçin Nmap, hasabat taýýarlamak üçin ReportLab we "
        "Matplotlib serişdeleri ulanylýar. Bu tehnologiýalaryň esasy bölegi açyk çeşmeli bolany üçin "
        "lisenziýa çykdajylary pes derejede saklanýar.",
    )
    add_p(
        doc,
        "Ykdysady hasaplamalarda bir lokal toruň el bilen barlagy ortaça 5 sagat, programma arkaly "
        "barlagy we hasabat taýýarlamak bolsa 1,25 sagat diýip kabul edildi. Bir ýylda 50 gezek "
        "audit geçirilende programma arkaly tygşytlanýan wagt ýyllyk ykdysady netijäniň esasy çeşmesi bolýar.",
    )

    add_heading(doc, "Kompauting we diskontirleme prosesleriniň mazmuny", 2)
    add_p(
        doc,
        "Maýa goýum taslamalarynda puluň wagt boýunça gymmaty hasaba alynmalydyr. Häzirki wagtda "
        "edilýän çykdajynyň geljekki gymmaty ýa-da geljekde alynjak girdejiniň häzirki gymmaty "
        "kompauting we diskontirleme formulalary arkaly kesgitlenýär.",
    )
    add_formula(doc, "1000 × (1 + 0,3) = 1300 manat")
    add_formula(doc, "1000 × (1 + 0,3)^2 = 1690 manat")
    add_formula(doc, "BC = HC(1 + K)^n")
    add_p(
        doc,
        "Bu ýerde BC – puluň geljekki gymmaty; HC – häzirki gymmat; K – göterim koeffisiýenti; "
        "n – ýyl sany. Bu formula goýlan serişdäniň geljekde näçe möçbere ýetjekdigini görkezýär.",
    )
    add_formula(doc, "HC = BC / (1 + K)^n")
    add_p(
        doc,
        "Bu ýerde HC – geljekki girdejiniň häzirki gymmaty. Mysal üçin, 5 ýyldan 1 000 000 manat "
        "girdeji almak meýilleşdirilse we ortaça göterim derejesi 5% bolsa, häzirki gymmat şeýle "
        "hasaplanylýar:",
    )
    add_formula(doc, "HC = 1 000 000 / (1 + 0,05)^5 = 783 526,2 manat")

    add_heading(doc, "Programma üpjünçiligi taslamasynyň ykdysady-tehniki esaslandyrylyşy", 2)
    add_p(
        doc,
        "IoT Security Scanner taslamasynyň ykdysady-tehniki esaslandyrylyşy programma toplumynyň "
        "döredilmegi, synagdan geçirilmegi we lokal torlarda başlangyç howpsuzlyk auditi üçin "
        "ulanmagy bilen baglanyşykly çykdajylary görkezýär.",
    )
    add_p(
        doc,
        "Taslama boýunça esasy çykdajylar işgärleriň zähmet hakyndan, kompýuter we test tor "
        "serişdelerinden, aragatnaşyk we transport çykdajylaryndan, sosial geçirimlerden hem-de "
        "iş ýeriniň guramaçylyk harajatlaryndan ybaratdyr.",
    )

    add_heading(doc, "Işgärleriň sanynyň we zähmet talap edişiniň hasaplanylyşy", 2)
    add_p(
        doc,
        "Programma toplumyny işläp taýýarlamak üçin programmist, dizaýner-programmist we ulgam "
        "dolandyryjy gatnaşýar. Işgärleriň sany ýerine ýetirilýän işleriň umumy zähmet talap edişine "
        "we her işgär üçin kabul edilen iş wagtyna görä kesgitlenilýär.",
    )
    add_formula(doc, "Nsanaw = Tdb / (Tişçi × Kö.ý.ý.n)")
    add_formula(doc, "Nsanaw = 96 / (32 × 1) = 3 işgär")
    add_p(
        doc,
        "Bu ýerde Tdb – programma toplumyny işläp taýýarlamagyň zähmet talap edişi; Tişçi – bir "
        "işgäriň kabul edilen iş wagty; Kö.ý.ý.n – öndürijiligiň ýerine ýetiriliş normasy.",
    )

    add_table(
        doc,
        [
            ["T№", "Işçileriň kategoriýasy", "Sany", "Razrýad", "Iş şerti", "1 sag. töleg.", "Baýrak %"],
            ["1", "Programmist", "1", "-", "Kadaly", "22", "10"],
            ["2", "Dizaýner-programmist", "1", "-", "Kadaly", "19", "10"],
            ["3", "Ulgam dolandyryjy", "1", "-", "Kadaly", "21", "10"],
            ["4", "", "", "", "", "", ""],
            ["5", "", "", "", "", "", ""],
            ["6", "", "", "", "", "", ""],
            ["7", "", "", "", "", "", ""],
            ["8", "", "", "", "", "", ""],
            ["9", "", "", "", "", "", ""],
            ["10", "", "", "", "", "", ""],
            ["11", "", "", "", "", "", ""],
            ["", "Jemi:", "3", "-", "", "62", "10"],
        ],
    )

    add_heading(doc, "Programma toplumyny işläp taýýarlamagyň bahasyny hasaplap çykarmak", 2)
    add_p(
        doc,
        "Çykdajylary hasaplanda tehniki serişdeler, programma gurşawy, transport harajatlary, "
        "zähmet haky, sosial gorag geçirimleri, bölüm we kärhana boýunça çykdajylar göz öňünde tutulýar.",
    )
    add_table(
        doc,
        [
            ["T№", "Ätiýaçlyk şaýlaryň görnüşi", "Ölçeg birligi", "Sany", "1-niň bahasy", "Umumy gymmaty"],
            ["1", "Kompýuter ýa-da noutbuk", "San", "1", "3200", "3200"],
            ["2", "Test routeri we Wi-Fi enjamy", "San", "1", "420", "420"],
            ["3", "Windows OU", "San", "1", "120", "120"],
            ["4", "Resminama we PDF hasabat taýýarlamak serişdeleri", "Toplum", "1", "160", "160"],
            ["5", "Python, PyQt6, Nmap we SQLite açyk çeşmeli serişdeler", "Toplum", "1", "0", "0"],
            ["6", "USB göteriji we tor kabeli", "Toplum", "1", "90", "90"],
            ["", "Jemi", "", "6", "", "3990"],
        ],
    )

    add_p(doc, "Transport harajatlary materiallaryň we tehniki serişdeleriň bahasynyň 10%-i möçberinde alynýar.")
    add_formula(doc, "Ttr = M × 0,10 = 3990 × 0,10 = 399,00 manat")

    add_p(doc, "Ortaça sagatlaýyn töleg işgärleriň sagatlyk tölegleriniň ortaça bahasy boýunça kesgitlenilýär.")
    add_formula(doc, "Corta = (22 + 19 + 21) / 3 = 20,67 manat/sag.")
    add_formula(doc, "Ztar = Corta × Tdb = 20,67 × 96 = 1984,32 manat")
    add_p(doc, "Baýrak gaznasynyň möçberi tarif boýunça zähmet hakynyň 10%-i hökmünde kabul edilýär.")
    add_formula(doc, "Zbaýr = Ztar × 0,10 = 198,43 manat")
    add_formula(doc, "Zes = Ztar + Zbaýr = 2182,75 manat")
    add_p(doc, "Goşmaça zähmet haky esasy zähmet hakynyň 10%-i möçberinde hasaplanylýar.")
    add_formula(doc, "Zgoşm = Zes × 0,10 = 218,28 manat")
    add_formula(doc, "Zum = Zes + Zgoşm = 2401,03 manat")
    add_p(doc, "Sosial gorag üçin geçirimler umumy zähmet hakynyň 20%-i möçberinde kabul edilýär.")
    add_formula(doc, "Ssos.gor = Zum × 0,20 = 480,21 manat")
    add_p(doc, "Bölüm boýunça we kärhana boýunça çykdajylar esasy zähmet hakyndan kabul edilen normlar boýunça hasaplanylýar.")
    add_formula(doc, "Çseh = Zes × 0,15 = 327,41 manat")
    add_formula(doc, "Çkärh = Zes × 0,30 = 654,83 manat")
    add_formula(doc, "Çenjam = Zes × 0,45 = 982,24 manat")

    add_p(doc, "Netijede IoT Security Scanner programma taslamasy boýunça ähli harajatlar aşakdaky tablisa jemlenilýär.")
    add_table(
        doc,
        [
            ["T№", "Maddanyň görnüşleri", "Gymmaty (Manat)"],
            ["1", "Materiallar, tehniki we programma serişdeleri", "3990,00"],
            ["2", "Transport harajatlary", "399,00"],
            ["3", "Esasy zähmet haky", "2182,75"],
            ["4", "Goşmaça zähmet haky", "218,28"],
            ["5", "Sosial gorag üçin geçirimler", "480,21"],
            ["6", "Seh boýunça çykdajylar", "327,41"],
            ["7", "Kärhana boýunça çykdajylar", "654,83"],
            ["8", "Enjamlara we programma gurşawyna çykdajylar", "982,24"],
            ["9", "Jemi harajatlar", "9234,72"],
            ["10", "Ýyllyk ykdysady netije", "4125,00"],
            ["11", "Özüni ödeýän wagty (ýyl)", "2,24"],
        ],
    )

    add_heading(doc, "Ýyllyk ykdysady netijäniň we özüni ödeýiş möhletiniň kesgitlenilişi", 2)
    add_p(
        doc,
        "Programmany ulanmagyň ykdysady netijesi tor auditine sarp edilýän wagtyň azalmagy bilen "
        "kesgitlenýär. El bilen audit geçirilende bir lokal toruň barlagy ortaça 5 sagat, programma "
        "arkaly barlag we hasabat taýýarlamak bolsa 1,25 sagat dowam edýär.",
    )
    add_formula(doc, "Eýyl = (Tel - Tprog) × Csag × Nýyl")
    add_formula(doc, "Eýyl = (5 - 1,25) × 22 × 50 = 4125,00 manat")
    add_p(
        doc,
        "Bu ýerde Tel – el bilen audit üçin gerek bolan wagt; Tprog – programma arkaly audit üçin "
        "gerek bolan wagt; Csag – hünärmeniň bir sagatlyk tölegi; Nýyl – bir ýylda geçirilýän auditleriň sany.",
    )
    add_formula(doc, "Töz.öd. = K / Eýyl")
    add_formula(doc, "Töz.öd. = 9234,72 / 4125,00 = 2,24 ýyl")
    add_p(
        doc,
        "Hasaplamadan görnüşi ýaly, taslamanyň özüni ödeýiş möhleti 2,24 ýyl, ýagny takmynan "
        "2 ýyl 3 aý bolýar. Programma birnäçe tor segmentinde ýa-da ýygy-ýygydan audit geçirilýän "
        "gurşawda ulanylsa, özüni ödeýiş möhleti has hem gysgalýar.",
    )

    add_heading(doc, "Ykdysady bölüm boýunça netije", 2)
    add_p(
        doc,
        "Geçirilen hasaplamalar IoT Security Scanner programma toplumynyň ykdysady taýdan "
        "maksadalaýykdygyny görkezýär. Taslama açyk çeşmeli tehnologiýalara esaslanýandygy üçin "
        "lisenziýa çykdajylary pesdir, esasy çykdajylar bolsa iş wagtyna we tehniki serişdelere degişlidir.",
    )
    add_p(
        doc,
        "Programma lokal torlaryň başlangyç howpsuzlyk auditini çaltlandyrýar, PDF hasabat taýýarlamagy "
        "awtomatlaşdyrýar we administratorlaryň wagtyny tygşytlaýar. Şonuň üçin bu programma toplumyny "
        "okuw edaralarynda, kiçi kärhanalarda we ofis torlarynda ulanmak ykdysady taýdan esaslydyr.",
    )

    return doc


def backup_main() -> Path:
    backup = MAIN_DOC.with_name(
        f"{MAIN_DOC.stem}_before_ykdysady_clean_{datetime.now().strftime('%Y%m%d_%H%M%S')}{MAIN_DOC.suffix}"
    )
    shutil.copy2(MAIN_DOC, backup)
    return backup


def replace_main_section(section_doc: Document) -> Path:
    backup = backup_main()
    main = Document(str(MAIN_DOC))
    body = main._body._element

    start_idx = None
    end_idx = None
    for paragraph in main.paragraphs:
        text = " ".join(paragraph.text.split())
        idx = body.index(paragraph._element)
        if idx > 300 and text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            start_idx = idx
            break

    if start_idx is None:
        raise RuntimeError("Ykdysady bölümiň başlangyjy tapylmady.")

    for paragraph in main.paragraphs:
        idx = body.index(paragraph._element)
        if idx <= start_idx:
            continue
        text = " ".join(paragraph.text.split())
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if text == "Netije" and style_name.startswith("Heading"):
            end_idx = idx
            break

    if end_idx is None:
        raise RuntimeError("Ykdysady bölümiň soňy tapylmady.")

    for element in list(body)[start_idx:end_idx]:
        body.remove(element)

    netije_element = list(body)[start_idx]
    for child in list(section_doc._body._element):
        if child.tag.endswith("sectPr"):
            continue
        body.insert(body.index(netije_element), copy.deepcopy(child))

    main.save(str(MAIN_DOC))
    return backup


def main() -> None:
    standalone = build_section_doc()
    standalone.save(str(OUT_DOC))

    section_for_main = build_section_doc(MAIN_DOC)
    backup = replace_main_section(section_for_main)
    print(f"Created: {OUT_DOC}")
    print(f"Updated: {MAIN_DOC}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
