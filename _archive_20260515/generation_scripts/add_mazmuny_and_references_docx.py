from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_mazmuny_references.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def make_paragraph(doc, text="", style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.alignment = align
    if style == "Normal":
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)
        if first_line:
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
    set_font(paragraph, 14, style.startswith("Heading"))
    return paragraph


def source_doc_from_blocks(blocks):
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    for block in blocks:
        kind = block[0]
        text = block[1] if len(block) > 1 else ""
        if kind == "h1":
            make_paragraph(doc, text, "Heading 1", WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        elif kind == "h2":
            make_paragraph(doc, text, "Heading 2", WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        elif kind == "toc":
            paragraph = make_paragraph(doc, text, "Normal", WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
            paragraph.paragraph_format.left_indent = Cm(block[2] if len(block) > 2 else 0)
        elif kind == "ref":
            paragraph = make_paragraph(doc, text, "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        else:
            make_paragraph(doc, text)
    return doc


def insert_elements_before(main_doc, target_paragraph, source_doc):
    body = target_paragraph._p.getparent()
    index = body.index(target_paragraph._p)
    for element in list(source_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def append_elements(main_doc, source_doc):
    body = main_doc.element.body
    sect_pr = None
    if len(body) and body[-1].tag.endswith("sectPr"):
        sect_pr = body[-1]
        body.remove(sect_pr)
    for element in list(source_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.append(deepcopy(element))
    if sect_pr is not None:
        body.append(sect_pr)


def add_missing_chapter_1_heading(doc):
    if any(p.text.strip().startswith("1. Predmet oblastynyň") for p in doc.paragraphs):
        return
    target = next((p for p in doc.paragraphs if p.text.strip().startswith("1.1.")), None)
    if not target:
        return
    source = source_doc_from_blocks([
        ("h1", "1. Predmet oblastynyň seljerilişi we meseläniň goýluşy"),
    ])
    insert_elements_before(doc, target, source)


def style_supporting_section_headings(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Zähmeti goramak we ekologik howpsuzlyk boýunça hasaplamalar":
            paragraph.style = "Heading 1"
            set_font(paragraph, 14, True)
        elif text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            paragraph.style = "Heading 1"
            set_font(paragraph, 14, True)


def move_global_netije_to_end(doc):
    ecology = next((p for p in doc.paragraphs if p.text.strip() == "Zähmeti goramak we ekologik howpsuzlyk boýunça hasaplamalar"), None)
    if ecology is None:
        return

    global_netije = None
    for p in doc.paragraphs:
        if p.text.strip() == "Netije" and p.style.name.startswith("Heading"):
            global_netije = p
            break
    if global_netije is None:
        return

    body = doc.element.body
    if body.index(global_netije._p) > body.index(ecology._p):
        return

    start = body.index(global_netije._p)
    end = body.index(ecology._p)
    moved = [deepcopy(element) for element in body[start:end]]
    for element in list(body[start:end]):
        body.remove(element)

    sect_pr = None
    if len(body) and body[-1].tag.endswith("sectPr"):
        sect_pr = body[-1]
        body.remove(sect_pr)
    for element in moved:
        body.append(element)
    if sect_pr is not None:
        body.append(sect_pr)


def add_mazmuny(doc):
    if any(p.text.strip() == "Mazmuny" for p in doc.paragraphs[:10]):
        return

    first = doc.paragraphs[0]
    toc_items = [
        ("toc", "Sözbaşy", 0),
        ("toc", "Giriş", 0),
        ("toc", "1. Predmet oblastynyň seljerilişi we meseläniň goýluşy", 0),
        ("toc", "1.1. IoT gurluşlarynyň düşünjesi we aýratynlyklary", 0.6),
        ("toc", "1.2. IoT howpsuzlygynyň wajyplygy", 0.6),
        ("toc", "1.3. IoT torlarynda howpsuzlyk gowşaklyklary we hüjüm wektorlary", 0.6),
        ("toc", "1.4. Bar bolan skanirleme ulgamlarynyň deňeşdirme seljermesi", 0.6),
        ("toc", "1.5. Diplom işiniň maksady we wezipeleri", 0.6),
        ("toc", "2. Programma toplumynyň arhitekturasyny taslamak", 0),
        ("toc", "2.1. Tehnologiýalar toplumyny saýlamagyň esaslandyrylyşy", 0.6),
        ("toc", "2.2. Programma arhitekturasy we modul gurluşy", 0.6),
        ("toc", "2.3. Maglumatlar bazasynyň taslanylyşy", 0.6),
        ("toc", "2.4. Skanirleme algoritminiň taslanylyşy", 0.6),
        ("toc", "2.5. Töwekgelçiligi bahalandyrmak we remediation modeliniň taslanylyşy", 0.6),
        ("toc", "3. Programma toplumynyň işlenip taýýarlanylyşy", 0),
        ("toc", "3.1. Skanirleme ýadrosynyň durmuşa geçirilmegi", 0.6),
        ("toc", "3.2. Gowşaklyklary ýüze çykarmak üçin plugin ulgamynyň durmuşa geçirilmegi", 0.6),
        ("toc", "3.3. Töwekgelçiligi bahalandyrmak we maslahat beriş ulgamynyň durmuşa geçirilmegi", 0.6),
        ("toc", "3.4. Maglumatlar bazasy we skanirleme taryhynyň durmuşa geçirilmegi", 0.6),
        ("toc", "3.5. Hasabat döretmek we duýduryş ulgamynyň durmuşa geçirilmegi", 0.6),
        ("toc", "4. Ulanyjy interfeýsiniň işlenip taýýarlanylyşy", 0),
        ("toc", "4.1. Interfeýsiň umumy konsepsiýasy we dizaýn çözgüdi", 0.6),
        ("toc", "4.2. Baş sahypa we tor skanirleme sahypasynyň durmuşa geçirilmegi", 0.6),
        ("toc", "4.3. Panel, Netijeler we Taryh sahypalarynyň işleýşi", 0.6),
        ("toc", "4.4. Sazlamalar sahypasy we giňeldilýän interfeýs mümkinçilikleri", 0.6),
        ("toc", "4.5. Interfeýsiň diplom goragynda görkezilişi", 0.6),
        ("toc", "5. Programma toplumyny synagdan geçirmek we netijeliligini bahalandyrmak", 0),
        ("toc", "5.1. Synag metodikasy we test gurşawy", 0.6),
        ("toc", "5.2. Funksional synaglaryň netijeleri", 0.6),
        ("toc", "5.3. Skanirleme netijeleriniň seljerilişi", 0.6),
        ("toc", "5.4. Remediation maslahatlarynyň we hasabat ulgamynyň barlagy", 0.6),
        ("toc", "5.5. Öndürijilik, takyklyk we çäklendirmeler boýunça netije", 0.6),
        ("toc", "Zähmeti goramak we ekologik howpsuzlyk boýunça hasaplamalar", 0),
        ("toc", "Taslamanyň ykdysady netijeliliginiň hasaplamasy", 0),
        ("toc", "Netije", 0),
        ("toc", "Ulanylan edebiýatlar", 0),
        ("p", ""),
    ]
    source = source_doc_from_blocks([("h1", "Mazmuny"), *toc_items])
    insert_elements_before(doc, first, source)


def add_references(doc):
    if any(p.text.strip() == "Ulanylan edebiýatlar" for p in doc.paragraphs):
        return

    references = [
        "1. NIST. SP 800-213: IoT Device Cybersecurity Guidance for the Federal Government: Establishing IoT Device Cybersecurity Requirements. https://csrc.nist.gov/pubs/sp/800/213/final (ýüz tutulan senesi: 13.05.2026).",
        "2. NIST. NISTIR 8259A: IoT Device Cybersecurity Capability Core Baseline. https://csrc.nist.gov/pubs/ir/8259/a/final (ýüz tutulan senesi: 13.05.2026).",
        "3. ENISA. Baseline Security Recommendations for IoT in the context of Critical Information Infrastructures. https://www.enisa.europa.eu/publications/baseline-security-recommendations-for-iot (ýüz tutulan senesi: 13.05.2026).",
        "4. OWASP Foundation. OWASP Internet of Things Project. https://owasp.org/www-project-internet-of-things/ (ýüz tutulan senesi: 13.05.2026).",
        "5. Lyon G. F. Nmap Network Scanning: The Official Nmap Project Guide to Network Discovery and Security Scanning. https://nmap.org/book/ (ýüz tutulan senesi: 13.05.2026).",
        "6. Nmap Project. Nmap Documentation. https://nmap.org/docs.html (ýüz tutulan senesi: 13.05.2026).",
        "7. Tenable. Nessus Documentation. https://docs.tenable.com/Nessus.htm (ýüz tutulan senesi: 13.05.2026).",
        "8. Shodan. Developer API Documentation. https://developer.shodan.io/api (ýüz tutulan senesi: 13.05.2026).",
        "9. Greenbone. Greenbone Community Documentation. https://greenbone.github.io/docs/latest/ (ýüz tutulan senesi: 13.05.2026).",
        "10. Python Software Foundation. Python 3 Documentation. https://docs.python.org/3/ (ýüz tutulan senesi: 13.05.2026).",
        "11. Riverbank Computing. PyQt6 Documentation. https://riverbankcomputing.com/software/pyqt (ýüz tutulan senesi: 13.05.2026).",
        "12. Qt Group. Qt Style Sheets Reference. https://doc.qt.io/qt-6/stylesheet-reference.html (ýüz tutulan senesi: 13.05.2026).",
        "13. SQLite Consortium. SQLite Documentation. https://www.sqlite.org/docs.html (ýüz tutulan senesi: 13.05.2026).",
        "14. ReportLab. ReportLab PDF Library User Guide. https://www.reportlab.com/docs/reportlab-userguide.pdf (ýüz tutulan senesi: 13.05.2026).",
        "15. Matplotlib Development Team. Matplotlib Documentation. https://matplotlib.org/stable/ (ýüz tutulan senesi: 13.05.2026).",
        "16. Paramiko Project. Paramiko Documentation. https://docs.paramiko.org/en/stable/ (ýüz tutulan senesi: 13.05.2026).",
        "17. PyInstaller Development Team. PyInstaller Manual. https://pyinstaller.org/en/stable/ (ýüz tutulan senesi: 13.05.2026).",
        "18. Python Package Index. python-nmap package information. https://pypi.org/project/python-nmap/ (ýüz tutulan senesi: 13.05.2026).",
        "19. Stallings W. Network Security Essentials: Applications and Standards. Pearson Education.",
        "20. Kurose J. F., Ross K. W. Computer Networking: A Top-Down Approach. Pearson Education.",
    ]

    source = source_doc_from_blocks(
        [("h1", "Ulanylan edebiýatlar")]
        + [("ref", item) for item in references]
    )
    append_elements(doc, source)


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    doc = Document(str(MAIN_FILE))

    add_missing_chapter_1_heading(doc)
    style_supporting_section_headings(doc)
    move_global_netije_to_end(doc)
    add_references(doc)
    add_mazmuny(doc)

    doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
