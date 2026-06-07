from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_section_5.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def add_table(doc, caption_text, headers, rows):
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "5. Programma toplumyny synagdan geçirmek we netijeliligini bahalandyrmak", level=1)
    add_normal(
        doc,
        "Programma toplumynyň işlenip taýýarlanylmagy onuň diňe nazary taýdan dogry gurlandygyny däl, "
        "eýsem amaly şertlerde hem işleýändigini görkezmegi talap edýär. Şonuň üçin IoT Security Scanner "
        "programmasy birnäçe ugur boýunça synagdan geçirildi: programma işe girizilişi, tor adapteriniň "
        "saýlanyşy, lokal IP aralygynyň kesgitlenişi, gurluşlaryň tapylyşy, portlaryň barlanyşy, netijeleriň "
        "maglumatlar bazasynda saklanyşy, PDF hasabatynyň döredilişi we interfeýs sahypalarynyň işleýşi "
        "barlanyldy."
    )

    add_heading(doc, "5.1. Synag metodikasy we test gurşawy", level=2)
    for text in [
        (
            "Synag metodikasy programma toplumynyň esasy funksional talaplaryna esaslanyp düzüldi. Ilki "
            "bilen programma Windows gurşawynda EXE faýly hökmünde işe girizildi. Soňra Wi-Fi adapteriniň "
            "dogry saýlanyşy, IP aralygynyň awtomatik doldurylmagy, skanirleme wagtynda gara konsol "
            "penjireleriniň açylmazlygy we netijeleriň interfeýsde görkezilişi barlandy."
        ),
        (
            "Test gurşawy hökmünde Windows operasion ulgamly noutbuk, lokal Wi‑Fi router, şol Wi‑Fi tora "
            "birikdirilen telefon we programma üpjünçiliginiň özi ulanyldy. Häzirki test torunda ýörite "
            "IP kamera ýa-da senagat IoT enjamy bolmadyk ýagdaýynda hem programma routeri, kompýuteri we "
            "torda görünýän beýleki gurluşlary ýüze çykarmagy başarmalydyr. IoT enjamly tora birikdirilende "
            "bolsa programma şol enjamyň açyk portlaryny we töwekgelçilik derejesini görkezmäge niýetlenendir."
        ),
        (
            "Synaglarda skanirleme aralygy hökmünde adaty lokal tor segmenti ulanyldy. Mysal üçin, kompýuteriň "
            "Wi-Fi salgysy 192.168.1.3/24 bolsa, programma 192.168.1.0/24 tor aralygyny saýlamalydyr. Bu "
            "ýagdaý aýratyn barlandy, sebäbi VPN adapteriniň saýlanmagy lokal Wi-Fi gurluşlarynyň netijede "
            "görünmezligine sebäp bolup biler."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 5.1 - Synag gurşawynyň başlangyç maglumatlary",
        ["Görkeziji", "Bahasy ýa-da beýany", "Synagdaky ähmiýeti"],
        [
            ["Operasion ulgam", "Windows", "Programma EXE görnüşinde işe girizilýär"],
            ["Programma görnüşi", "IoT_Security_Scanner.exe", "Ulanyjy üçin taýýar desktop programma"],
            ["Tor görnüşi", "Lokal Wi‑Fi/Ethernet", "IoT gurluşlarynyň ýerleşýän esasy gurşawy"],
            ["Skanirleme aralygy", "192.168.1.0/24 ýa-da şoňa meňzeş lokal segment", "Router, PC we telefon ýaly gurluşlary tapmak"],
            ["Skanirleme guraly", "Nmap, ARP, Ping we ARP cache", "Hostlary birnäçe usul bilen ýüze çykarmak"],
            ["Maglumat bazasy", "iot_security.db", "Netijeleri we taryhy saklamak"],
            ["Hasabat", "PDF", "Audit netijesini resmileşdirmek"],
        ],
    )

    add_heading(doc, "5.2. Funksional synaglaryň netijeleri", level=2)
    for text in [
        (
            "Funksional synaglarda programma aýratyn mümkinçilikler boýunça barlandy. Esasy maksat her bir "
            "modulyň garaşylýan netijäni berýändigini subut etmekdir. Şonuň üçin skanirleme diňe bir “başlaýar "
            "ýa-da başlamýar” görnüşinde däl, eýsem adapter saýlamak, IP aralygyny barlamak, hostlary tapmak, "
            "portlary görkezmek, taryha ýazmak we PDF döretmek ýaly aýratyn ssenariýalar boýunça seljerildi."
        ),
        (
            "Programma işe girizilende gara konsol penjiresi açylmady. Bu PyInstaller ýygnagynyň grafiki "
            "programma hökmünde taýýarlanandygyny we daşky komandalar işledilende CREATE_NO_WINDOW ýörelgesiniň "
            "ulanylýandygyny görkezýär. Şeýle aýratynlyk diplom goragy wagtynda programmanyň professional "
            "görnüşde görkezilmegine kömek edýär."
        ),
        (
            "Skanirleme sahypasynda Wi‑Fi adapteri saýlanylanda IP aralygy awtomatik dolduryldy. Bu ýagdaý "
            "ulanyjynyň nädogry VPN ýa-da başga tor segmentini skanirlemek ähtimallygyny peseldýär. Eger "
            "adapter VPN hökmünde tanalsa, programma ulanyja Wi‑Fi gurluşlaryny görmek üçin Wi‑Fi adapterini "
            "saýlamagyň zerurdygyny düşündirýär."
        ),
        (
            "Netijeler tablisasynyň işleýşi hem synagdan geçirildi. Tapylan gurluş üçin IP, ady, MAC, öndüriji, "
            "görnüşi, tapylan usul, portlar, töwekgelçilik, gowşaklyk we maslahat sütünleri doldurylýar. Eger "
            "gurluş torda tapylyp, emma port açmasa, ol netijelerden aýrylmaýar we pes töwekgelçilik bilen "
            "görkezilýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 5.2 - Funksional synag ssenariýalary",
        ["№", "Synag ssenariýasy", "Garaşylýan netije", "Netije"],
        [
            ["1", "Programmany EXE görnüşinde açmak", "Esasy penjire açylýar, gara konsol çykmaýar", "Üstünlikli"],
            ["2", "Wi‑Fi adapterini saýlamak", "IP aralygy awtomatik doldurylýar", "Üstünlikli"],
            ["3", "Nädogry IP aralygyny girizmek", "Programma duýduryş berýär", "Üstünlikli"],
            ["4", "Lokal tory skanirlemek", "Router, PC we jogap berýän gurluşlar tapylýar", "Üstünlikli"],
            ["5", "Port açmaýan gurluşy görkezmek", "Gurluş tablisa goşulýar, portlar “ýok” bolýar", "Üstünlikli"],
            ["6", "Netijeleri taryha ýazmak", "Skanirleme SQLite bazasynda saklanýar", "Üstünlikli"],
            ["7", "PDF hasabat döretmek", "security_report.pdf faýly döredilýär", "Üstünlikli"],
            ["8", "Panel, Netijeler we Taryh sahypalaryny açmak", "Sahypalar saklanan maglumatlary görkezýär", "Üstünlikli"],
        ],
    )

    add_heading(doc, "5.3. Skanirleme netijeleriniň seljerilişi", level=2)
    for text in [
        (
            "Lokal Wi‑Fi torunda geçirilen synaglaryň esasy maksady programma tarapyndan gurluşlaryň dogry "
            "tapylýandygyny we olaryň netijeler tablisasyna düşýändigini görkezmekdir. Torda ýörite IP kamera "
            "ýa-da IoT datçigi ýok bolsa hem programma routeri, lokal kompýuteri we jogap berýän telefon ýaly "
            "gurluşlary görkezip bilýär. Bu ýagdaý diplom goragy üçin möhüm düşündirişdir: programma diňe "
            "kamerany gözlemeýär, ol tutuş lokal tor kartasyny düzýär."
        ),
        (
            "Eger bir gurluşyň portlary açyk bolmasa, bu onuň howpsuz ýa-da doly barlananlygyny aňlatmaýar, "
            "emma häzirki wagtda umumy TCP port sanawy boýunça açyk hyzmat tapylmandygyny görkezýär. Şonuň "
            "üçin programma şeýle gurluşlary pes töwekgelçilik bilen görkezýär. Eger IoT kamera, router ýa-da "
            "web dolandyryşly enjam torda bolsa, 80, 443, 554, 8080 ýa-da 23 ýaly portlar açyk ýagdaýda tapylyp "
            "bilner we risk derejesi degişli görnüşde ýokarlanar."
        ),
        (
            "Netijeleri seljermekde tapylan usul sütuny aýratyn ähmiýete eýedir. Gurluş Nmap, ARP, Ping, "
            "Gateway ýa-da Local PC usullary arkaly tapylyp bilýär. Bu maglumat administratora gurluşyň "
            "nähili ýagdaýda ýüze çykarylandygyny düşündirýär. Mysal üçin, telefon köplenç port açmasa-da, "
            "ARP ýa-da Ping arkaly torda görünip biler."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 5.3 - Skanirleme netijeleriniň mysal seljermesi",
        ["Gurluş görnüşi", "Tapylýan maglumat", "Programmada görkezilýän netije"],
        [
            ["Router", "Şlýuz IP salgysy, MAC, käbir açyk hyzmatlar", "Görnüşi Router, risk portlara görä kesgitlenýär"],
            ["Lokal kompýuter", "Kompýuteriň IP we host ady", "Görnüşi PC, tapylan usul Local PC/Nmap/ARP bolup biler"],
            ["Telefon", "ARP ýa-da Ping arkaly IP/MAC", "Telefon/Unknown, portlar açyk däl bolsa risk pes"],
            ["IP kamera", "HTTP/RTSP portlary, MAC öndürijisi", "IoT kandidat, 80/554 portlara görä orta risk"],
            ["Telnet açyk enjam", "23-nji port we login sahypasy", "Telnet plugin barlagy, gowşak parol bolsa ýokary risk"],
        ],
    )

    add_heading(doc, "5.4. Remediation maslahatlarynyň we hasabat ulgamynyň barlagy", level=2)
    for text in [
        (
            "Remediation maslahatlarynyň barlagy portlara we plugin netijelerine görä geçirildi. Gurluşda "
            "açyk port ýok bolsa, programma “Portlar açyk däl, gurluş diňe torda ýüze çykaryldy” diýen "
            "maslahaty berýär. HTTP interfeýsi ýüze çykanda firmware-i täzelemek, standart hasaplary öçürmek "
            "we web interfeýsiň elýeterliligini çäklendirmek ýaly maslahatlar görkezilýär."
        ),
        (
            "Telnet gowşak autentifikasiýasyny barlaýan plugin aýratyn synag ssenariýasy hökmünde "
            "bahalandyryldy. Eger 23-nji port açyk bolsa, programma diňe portuň bardygyny görkezmek bilen "
            "çäklenmeýär, login/parol görnüşindäki gowşak kombinasiýalary synap görýär. Gowşak giriş kabul "
            "edilen ýagdaýynda risk ýokary derejä çykarylýar we paroly üýtgetmek ýa-da Telnet hyzmatyny "
            "öçürmek maslahaty berilýär."
        ),
        (
            "PDF hasabat ulgamy hem synagdan geçirildi. Netijeler sahypasyndan ýa-da Taryh sahypasyndan "
            "hasabat eksport edilende programma tapylan gurluşlary, olaryň portlaryny, öndürijisini, görnüşini, "
            "tapylan usulyny, gowşaklyk ýagdaýyny we maslahatlaryny PDF faýlyna ýazýar. Bu mümkinçilik "
            "audit netijesini diplom goragynda ýa-da administrator hasabatynda görkezmek üçin amatlydyr."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 5.4 - Remediation we hasabat ulgamynyň synagy",
        ["Synag obýekti", "Barlanýan ýagdaý", "Netije"],
        [
            ["Port açmaýan gurluş", "Açyk port ýok", "Pes risk we düşündiriş maslahaty görkezilýär"],
            ["HTTP hyzmaty", "80/443 portlary açyk", "Firmware we standart hasaplar barada maslahat berilýär"],
            ["RTSP hyzmaty", "554 porty açyk", "Kamera akymynyň rugsatlaryny barlamak maslahat berilýär"],
            ["Telnet hyzmaty", "23 porty açyk", "Gowşak parol barlagy işledilýär"],
            ["PDF eksport", "Saklanan netije saýlanýar", "Hasabat faýly döredilýär"],
            ["Taryh sahypasy", "Öňki skanirleme saýlanýar", "Gysgaça mazmun we eksport mümkinçiligi görkezilýär"],
        ],
    )

    add_heading(doc, "5.5. Öndürijilik, takyklyk we çäklendirmeler boýunça netije", level=2)
    for text in [
        (
            "Programmany öndürijilik nukdaýnazaryndan bahalandyrmakda onuň köp akymly skanirleme ulanýandygy "
            "möhüm netijedir. ThreadPoolExecutor arkaly birnäçe host bir wagtda barlanýar, bu bolsa /24 "
            "görnüşli lokal torlarda skanirleme wagtyny azaltmaga mümkinçilik berýär. Skanirleme wagty toruň "
            "ölçegine, enjam sanyna, Wi‑Fi ýagdaýyna, Nmap elýeterliligine we hostlaryň jogap beriş tizligine "
            "baglydyr."
        ),
        (
            "Takyklyk boýunça programma birnäçe tapyş usulyny birleşdirýändigi üçin diňe bir Nmap netijesine "
            "garanyňda has maglumatly netije berýär. ARP cache, Ping sweep, gateway we local PC maglumatlarynyň "
            "goşulmagy tordaky gurluşlaryň görünmek mümkinçiligini ýokarlandyrýar. Şeýle-de bolsa, käbir "
            "telefonlar ýa-da IoT gurluşlary batareýany tygşytlamak, privat MAC, firewall ýa-da AP isolation "
            "sebäpli jogap bermän biler. Bu ýagdaý programma kemçiligi däl, lokal toruň howpsuzlyk ýa-da "
            "izolýasiýa sazlamalarynyň netijesi bolup biler."
        ),
        (
            "Synaglaryň netijesi IoT Security Scanner programmasynyň diplom işiniň maksadyna laýyk gelýändigini "
            "görkezdi. Programma lokal tor adapterini saýlap bilýär, IP aralygyny kesgitleýär, işjeň gurluşlary "
            "tapýar, IoT bilen bagly portlary barlaýar, töwekgelçilik derejesini görkezýär, maslahat berýär, "
            "netijeleri maglumatlar bazasynda saklaýar we PDF hasabat döredýär. Bu mümkinçilikler ony lokal IoT "
            "torlarynyň başlangyç howpsuzlyk auditi üçin peýdaly gural edýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 5.5 - Programma netijeliliginiň umumy bahasy",
        ["Bahalandyrma ugry", "Netije", "Düşündiriş"],
        [
            ["Funksionallyk", "Ýerine ýetirildi", "Skanirleme, taryh, hasabat we interfeýs işleýär"],
            ["Ulanyş amatlylygy", "Ýokary", "Grafiki interfeýs we türkmen dili ulanylyşy ýeňilleşdirýär"],
            ["Öndürijilik", "Kanagatlanarly", "Köp akymly barlag /24 lokal tor üçin amatly"],
            ["Takyklyk", "Ýeterlik", "Nmap, ARP, Ping we gateway maglumatlary birleşdirilýär"],
            ["Çäklendirme", "Bar", "AP isolation, firewall ýa-da privat MAC käbir gurluşlary gizläp biler"],
            ["Diplom goragy üçin taýýarlyk", "Ýokary", "Programma işleýär, netijeleri tablisa we PDF görnüşinde görkezýär"],
        ],
    )

    add_normal(
        doc,
        "Şeýlelikde, bäşinji bölümde geçirilen synaglar programma toplumynyň esasy wezipelerini ýerine "
        "ýetirýändigini görkezdi. IoT Security Scanner lokal toruň gurluşlaryny tapmak, portlary barlamak, "
        "töwekgelçiligi bahalandyrmak, maslahat bermek, taryh saklamak we hasabat taýýarlamak mümkinçiliklerini "
        "birleşdirýär. Bu bolsa programmany diplom taslamasynda goýlan maksada laýyk amaly çözgüt hökmünde "
        "bahalandyrmaga mümkinçilik berýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_styles(main_doc):
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text == "5. Programma toplumyny synagdan geçirmek we netijeliligini bahalandyrmak":
            paragraph.style = "Heading 1"
            set_font(paragraph, 14, True)
        elif text.startswith(("5.1.", "5.2.", "5.3.", "5.4.", "5.5.")):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip() == "5. Programma toplumyny synagdan geçirmek we netijeliligini bahalandyrmak" for p in main_doc.paragraphs):
        print("Section 5 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_styles(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
