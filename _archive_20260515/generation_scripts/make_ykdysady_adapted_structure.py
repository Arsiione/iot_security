from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = Path(r"C:\Users\Arslan\Desktop\yjd.docx")
MAIN_DOC = ROOT / "kakabalowa.docx"
OUT_DOC = ROOT / "ykdysady_iot_security_yjd_uytgedilen.docx"


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_heading_text(paragraph, text: str, size: int = 14) -> None:
    set_text(paragraph, text)
    for run in paragraph.runs:
        run.bold = True
        run.font.size = Pt(size)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def build_adapted_doc() -> Document:
    doc = Document(str(SOURCE))

    section_title = doc.paragraphs[0].insert_paragraph_before(
        "Taslamanyň ykdysady netijeliliginiň hasaplamasy"
    )
    try:
        section_title.style = "Heading 1"
    except KeyError:
        pass
    set_heading_text(section_title, "Taslamanyň ykdysady netijeliliginiň hasaplamasy")

    # The original yjd.docx paragraph indexes are shifted by +1 after inserting
    # the section title. The structure, formulas, and tables are kept; only the
    # wording is adapted so the section reads as an independent diploma chapter.
    replacements = {
        0: "Maýa goýum taslamalarynyň netijeliligini bahalandyrmagyň esasy ýörelgeleri.",
        1: (
            "Maýa goýum taslamalarynyň ykdysady netijeliligini kesgitlemek üçin birnäçe "
            "çemeleşme ulanylýar. Olaryň her biri ulanylyş maksady, hasaplama şertleri we "
            "bahalandyrylýan desga boýunça tapawutlanýar, ýöne umumy mazmuny taslamanyň "
            "çykdajylaryny we ondan alynýan peýdany deňeşdirmekden ybaratdyr."
        ),
        2: (
            "Taslamanyň netijeliligi onuň öňde goýlan maksada, gatnaşyjylaryň bähbidine "
            "we ykdysady taýdan kabul ederlikli netijä laýyk gelýänligini görkezýär. Şonuň "
            "üçin diňe umumy taslama däl, eýsem onuň ýerine ýetirilişine gatnaşýan taraplaryň "
            "her biri boýunça hem netijelilik bahalandyrylmalydyr."
        ),
        3: (
            "Bütewi taslamanyň netijeliligi maliýeleşdirmek mümkinçiligini esaslandyrmak "
            "we taslamanyň geljekki özüne çekijiligini kesgitlemek üçin öwrenilýär. Bu "
            "bahalandyrma, adatça, aşakdaky ugurlary öz içine alýar:"
        ),
        4: "taslamanyň durmuş-ykdysady netijesini;",
        5: "taslamanyň täjirçilik ýa-da amaly ykdysady netijesini.",
        6: (
            "Taslama gatnaşmagyň netijeliligi bolsa taslamany durmuşa geçirmek mümkinçiligini "
            "barlamak we oňa gatnaşýan taraplaryň gyzyklanmasyny kesgitlemek üçin ulanylýar. "
            "Bu ugur aşakdaky hasaplamalary öz içine alyp biler:"
        ),
        7: "kärhanalaryň we guramalaryň taslama gatnaşmakdan alýan netijesini kesgitlemek;",
        8: "maýa goýum serişdeleriniň ulanylyşynyň maksadalaýyklygyny kesgitlemek;",
        9: "has ýokary derejeli düzümleriň taslama gatnaşmagynyň netijesini kesgitlemek, şol sanda:",
        14: "Netijeliligi bahalandyrmagyň esasy ýörelgeleri.",
        15: "Taslama netijeliligini kesgitlände şu möhüm ýörelgeler göz öňünde tutulýar:",
        16: "- taslamanyň bütin hereket döwründe onuň ýerine ýetirilişini yzygiderli seljermek;",
        17: "- çykdajy we girdeji pul akymlaryny modelleşdirmek;",
        18: "- dürli taslamalary deňeşdirmek üçin birmeňzeş şertleri üpjün etmek;",
        19: "- kabul ederlikli we ýokary netijelilik derejesini kesgitlemek;",
        20: "- wagt faktorynyň ykdysady netijelilige edýän täsirini hasaba almak;",
        21: "- geljekki çykdajylary we girdejileri aýratynlykda seljermek;",
        22: '- "taslamaly" we "taslamasyz" ýagdaýlaryň tapawudyny görkezmek;',
        23: "- taslamanyň ähli wajyp tehniki, ykdysady we durmuş netijelerini hasaba almak;",
        24: "- taslama gatnaşyjylaryň dürli bähbitleriniň bolup biljekdigini göz öňünde tutmak;",
        25: "- bahalandyrmany birnäçe tapgyrda geçirmek;",
        26: "- dolanyşyk kapitalynyň zerurlygyny we onuň netijelilige täsirini kesgitlemek;",
        27: "- inflýasiýanyň, baha üýtgemeleriniň we walýuta şertleriniň täsirini hasaba almak;",
        28: "- näbellilikleriň we töwekgelçilikleriň täsirini mümkin boldugyça san görnüşinde görkezmek.",
        29: "Taslamanyň netijeliligini hasaplamak üçin başlangyç maglumatlar.",
        30: (
            "Başlangyç maglumatlaryň göwrümi taslamanyň haýsy tapgyrda bahalandyrylýandygyna "
            "baglydyr. Deslapky tapgyrda maglumatlar umumy häsiýete eýe bolup biler, emma "
            "işçi taslama taýýarlananda olar has anyk görnüşde berilmelidir."
        ),
        31: "Ähli tapgyrlarda başlangyç maglumatlara aşakdakylar girizilýär:",
        32: "taslamanyň maksady we çözmeli meselesi;",
        33: "ulanylýan tehnologiýalar, ýerine ýetirilýän işler we hödürlenýän hyzmatlar baradaky maglumatlar;",
        34: "taslamanyň başlanýan we tamamlanýan şertleri, şeýle hem hasabat döwrüniň dowamlylygy;",
        35: "taslamanyň ýerine ýetirilýän ykdysady gurşawy baradaky maglumatlar.",
        36: "Maýa goýumlary teklip etmek tapgyrynda taslama boýunça şu maglumatlar görkezilýär:",
        37: "- ýerine ýetirilmeli işleriň dowamlylygy barada maglumatlar;",
        38: "- başlangyç maýa goýumlaryň möçberi barada maglumatlar;",
        39: "- taslama amala aşyrylanda ýyllar boýunça garaşylýan girdejiler barada maglumatlar;",
        40: "- taslamanyň ýyllar boýunça önümçilik ýa-da hyzmat çykdajylary barada maglumatlar.",
        41: (
            "Maýa goýumlaryny esaslandyrmak tapgyrynda taslama baradaky maglumatlar has giň "
            "görnüşde görkezilýär we esaslandyryjy hasabatlar bilen berkidilýär. Bu maglumatlar "
            "şulary öz içine alýar:"
        ),
        42: "- tehnologiki düzüm we wagt boýunça paýlanan maýa goýumlar;",
        43: "- çykdajylaryň görnüşleri boýunça paýlanyşy we hyzmatlardan alynýan girdeji baradaky maglumatlar.",
        44: (
            "Ykdysady-tehniki esaslandyryş tapgyrynda başlangyç maglumatlar has doly görnüşde "
            "berilýär. Bu tapgyrda aşakdaky maglumatlaryň görkezilmegi maksadalaýykdyr:"
        ),
        45: "taslama we onuň ýerine ýetirijileri barada maglumatlar;",
        46: "taslamanyň ykdysady gurşawy;",
        47: "taslama bilen baglanyşykly pudaklar we hyzmatlar barada maglumatlar;",
        48: "maýa goýum işlerinden emele gelýän pul akymlary;",
        49: "amaly işlerden emele gelýän pul akymlary;",
        50: "maliýe işleri bilen bagly pul akymlary.",
        51: "Taslama barada umumy maglumatlar şu ugurlary öz içine alýar:",
        52: "taslamanyň häsiýeti, ýerine ýetirilýän işleriň ýa-da hyzmatlaryň düzümi baradaky maglumatlar;",
        53: "taslamanyň ýerleşýän gurşawy we ulanylyş meýdany baradaky maglumatlar;",
        54: (
            "tehnologiki prosesleriň aýratynlyklary, ulanylýan serişdeler, ýerine ýetirilen "
            "işleriň netijesini ulanmak ýa-da ýerleşdirmek ulgamy baradaky maglumatlar."
        ),
        55: (
            "Maýa goýum netijeliligine baha bermekde her gatnaşyjynyň wezipesi we onuň "
            "taslamadaky paýy barada goşmaça maglumatlaryň bolmagy zerurdyr."
        ),
        56: (
            "Hasaplamalaryň dogry bolmagy üçin ýerine ýetirijileriň önümçilik mümkinçiligi, "
            "maliýe ýagdaýy we taslamany alyp barmak ukyby baradaky maglumatlar hem göz öňünde tutulýar."
        ),
        58: "Maýa goýum taslamalarynyň netijeliligini bahalandyrmak",
        59: (
            "Taslamanyň durmuş-ykdysady netijesini bahalandyrmak umumy hojalyk derejesindäki "
            "görkezijileri hasaplamakdan ybaratdyr. Bu usulda taslamanyň diňe gönüden-göni "
            "girdejisi däl, eýsem onuň jemgyýetçilik we pudaklaýyn täsiri hem nazara alynýar."
        ),
        60: "- taslamanyň netijeleri beýleki durmuş-ykdysady we ekologik ugurlaryň pul akymlarynda hem şöhlelendirilýär;",
        61: "- dolanyşyk kapitalynyň düzümindäki ätiýaçlar we zerur resurslar göz öňünde tutulýar;",
        62: (
            "- maliýe we amaly işleriň pul akymlary seljerilende karz serişdeleri, göterim "
            "tölegleri, subsidiýalar, salgytlar we beýleki transfert tölegleri aýratynlykda görkezilýär;"
        ),
        63: "- hyzmatlar, işler we sarp edilýän serişdeler ykdysady taýdan esaslandyrylan bahalar boýunça bahalandyrylýar.",
        64: "Maýa goýum işlerinden gelýän pul akymlarynda aşakdakylar hasaba alynýar:",
        65: "hasaplaşyk döwrüniň ähli ädimlerinde esasy serişdelere goýumlar;",
        66: "taslama bes edilende ýüze çykyp biljek çykdajylar;",
        67: "dolanyşyk kapitalynyň artmagyna gönükdirilen goýumlar;",
        68: "taslama tamamlananda emlägi ýa-da maddy däl aktiwleri ýerlemekden alynýan girdeji.",
        69: (
            "Býujet netijeliligi döwlet ýa-da ýerli dolandyryş edaralarynyň talaplaryna görä "
            "kesgitlenýär. Bu ýagdaýda dürli derejeli býujetler üçin pul serişdeleriniň gelşi "
            "we çykdajysy aýratynlykda seljerilýär."
        ),
        70: "Maliýeleşdirmek shemasy taýýarlanylandan soň bahalandyrmagyň indiki tapgyry geçirilýär.",
        71: (
            "Bu tapgyrda gatnaşyjylaryň sany, olaryň borçlary we her biriniň taslamany ýerine "
            "ýetirmek mümkinçiligi hasaba alynýar. Netijede pudaklaýyn, sebitleýin, býujet we "
            "aýratyn kärhana derejesindäki netijelilik görkezilýär."
        ),
        72: (
            "Taslama gatnaşyjylaryň her biri üçin ulanylýan tölegler we girdejiler dürli bolup "
            "bilýär. Şonuň üçin hasaplamalarda gatnaşyjynyň taslamadaky wezipesi, maliýe ýagdaýy "
            "we ýerine ýetirýän işi aýratyn göz öňünde tutulýar."
        ),
        73: (
            "Kärhananyň ýa-da ýerine ýetiriji toparyň mümkinçilikleri onuň tehniki serişdeleri, "
            "işçi personalynyň düzümi, maddy däl aktiwleri we öňki tejribesi boýunça bahalandyrylýar."
        ),
        74: (
            "Täze taslama döredilende gatnaşyjy taraplaryň paýy, başlangyç maýasy we geljekde "
            "öz üstüne aljak borçlary baradaky maglumatlar öňünden takyklanmalydyr."
        ),
        75: "Taslamanyň ykdysady gurşawy baradaky maglumatlar şulary öz içine alýar:",
        76: "Bahalaryň, hyzmatlaryň we serişdeleriň taslama döwründäki üýtgemegi baradaky çaklamalar;",
        77: "Walýuta hümmetiniň we inflýasiýa derejesiniň üýtgemegi baradaky çaklamalar;",
        78: "Salgyt salmak ulgamy baradaky maglumatlar.",
        79: (
            "Bu maglumatlaryň çeşmesi hökmünde ykdysady we maliýe edaralarynyň çaklamalary, "
            "bazar bahalarynyň seljermesi, hyzmatlaryň we serişdeleriň nyrh gurluşy baradaky "
            "maglumatlar ulanylyp bilner."
        ),
        80: "Salgyt salmak boýunça maglumatlarda salgytlaryň, ýygymlaryň we beýleki tölegleriň sanawy görkezilýär.",
        81: "Aýratynlykda ýerli salgytlar we sebitleýin düzgünler göz öňünde tutulmalydyr.",
        82: "Salgyt salmak bazasy;",
        83: "Salgyt stawkasy;",
        84: "Salgytlaryň töleniş döwürleri;",
        85: "Salgyt ýeňillikleri we olaryň ulanylyş şertleri;",
        86: "Salgyt tölegleriniň dürli derejedäki býujetleriň arasynda paýlanylyşy.",
        87: (
            "Dürli maýadarlaryň taslama gatnaşmak şertleri tapawutlanyp biler. Mysal üçin, "
            "karz berýän bank, kärendeçi ýa-da paý goýýan tarap üçin girdejiler we çykdajylar "
            "dürli hasaplanylýar."
        ),
        88: (
            "Şonuň üçin taslama bahalandyrylanda görkezijileriň doly toplumy ulanylýar. "
            "Maliýeleşdirmek boýunça çözgüt kabul edilende netijeliligi görkezýän esasy "
            "ölçegleriň durnukly pul birliginde hasaplanylmagy zerurdyr."
        ),
        89: "Eger taslama ykdysady netijelilik ölçegleriniň esasy talaplaryny kanagatlandyrsa, ony durmuşa geçirmek maksadalaýyk hasaplanylýar.",
        91: "Kompauting we diskontirleme prosesleriniň mazmuny.",
        92: (
            "Maýa goýum taslamalary boýunça çözgüt kabul edilende puluň wagt boýunça gymmaty "
            "hasaba alynýar. Sebäbi şu günki çykdajy bilen birnäçe ýyldan soň alynjak girdeji "
            "birmeňzeş ykdysady ähmiýete eýe däldir. Şonuň üçin kompauting we diskontirleme "
            "usullary ulanylýar."
        ),
        95: "Çylşyrymly göterimiň umumy görnüşi aşakdaky formula arkaly ýazylýar.",
        97: (
            "Bu ýerde BC – puluň geljekki gymmaty; HC – häzirki gymmaty; K – göterim "
            "koeffisiýenti; n – ýyl sany. Bu proses goýlan serişdäniň geljekde nähili "
            "möçbere ýetjekdigini görkezýär."
        ),
        103: (
            "Mysal üçin, 5 ýyldan 1 million manat girdeji almak meýilleşdirilse we depozit "
            "goýumlaryň ortaça derejesi 5% bolsa, häzirki gymmat aşakdaky ýaly kesgitlenýär:"
        ),
        107: "Programma üpjünçiligi taslamasynyň ykdysady-tehniki esaslandyrylyşy.",
        108: (
            "Programma üpjünçiligi taslamasynyň ykdysady-tehniki esaslandyrylyşy (YTE) "
            "IoT Security Scanner programma toplumyny döretmek, synagdan geçirmek we amaly "
            "ulanmak üçin esasy resminama hökmünde çykyş edýär. Bu ýerde tehniki çözgüt, "
            "zerur serişdeler, çykdajylar we ykdysady netije özara baglanyşykda görkezilýär."
        ),
        109: (
            "YTE-de programma arhitekturasy, tor skanirleme algoritmleri, ulanylýan "
            "tehnologiýalar, maglumatlar bazasy, hasabat ulgamy, iş şertleri, zähmet we "
            "ekologik howpsuzlyk meseleleri, şeýle hem taslamanyň ykdysady netijeliligi görkezilýär."
        ),
        110: "Programma üpjünçiligini işläp taýýarlamak üçin maýa goýum çeşmeleri hökmünde şular görkezilip bilner:",
        116: "Programma taslamasynyň ykdysady-tehniki esaslandyrylyşy şu bölümleri öz içine alýar:",
        117: "- umumy düşündiriş ýazgysy;",
        118: "- predmet oblastynyň seljerilişi we tor howpsuzlygy boýunça maglumatlar;",
        119: "- programma arhitekturasy we tehnologiýa çözgütleri;",
        120: "- programma toplumyny dolandyrmak, ulanyjy interfeýsi we iş şertlerini goramak;",
        121: "- maglumatlar bazasy, hasabat we bildiriş ulgamlary;",
        122: "- synag işleri, skanirleme netijeleri we netijelilik bahalandyrmasy;",
        123: "- daşarky gurşawy goramak we kagyz serişdelerini tygşytlamak.",
        124: "- ykdysady netijeliligiň hasaplanylyşy.",
        125: "IoT Security Scanner programma toplumyny işläp taýýarlamakda işgärleriň sanynyň hasaplanylyşy",
        126: (
            "Programma üpjünçiligi taslamasynyň ykdysady netijeliligini kesgitlemek üçin ilki "
            "bilen taslama gatnaşýan hünärmenleriň sany, olaryň iş wagty we programma işleriniň "
            "zähmet talap edişi kesgitlenilýär."
        ),
        127: (
            "Işgärleriň sanyny kesgitlemek üçin iş tertibi, zähmet öndürijiligi we ýerine "
            "ýetirilýän işleriň umumy göwrümi göz öňünde tutulýar."
        ),
        128: (
            "Şu görkezijileriň esasynda taslama üçin zerur işgär sany aşakdaky formula boýunça "
            "kesgitlenilýär:"
        ),
        129: "Tdb – programma toplumyny işläp taýýarlamagyň zähmet talap edişi",
        130: "Tdb – programma üpjünçiligini döretmek işleriniň zähmet talap edişi",
        136: "Programma toplumyny işläp taýýarlamagyň bahasyny hasaplap çykarmak.",
        137: "Çykdajylary hasaba almakda esasan aşakdaky maddalar göz öňünde tutulýar:",
        138: "a) Tehniki serişdeler we programma üpjünçiligi",
        139: "b) Internet, aragatnaşyk we energiýa",
        140: "ç) Işgärleriň esasy we goşmaça zähmet haky",
        141: "d) Sosial gorag üçin geçirimler",
        142: (
            "Enjamlary, kompýuterleri we tor serişdelerini ulanmak we saklamak üçin edilýän "
            "harajatlara amortizasiýa, gündelik hyzmat etmek hem-de beýleki goşmaça çykdajylar degişlidir."
        ),
        143: (
            "Bu harajatlar kompýuterleriň, routeriň, test tor gurşawynyň we iş ýeriniň "
            "ulanylyşy bilen baglanyşykly çykdajylary öz içine alýar."
        ),
        144: (
            "Bölüm boýunça harajatlar programma taslamasyny dolandyrmak, iş ýerini üpjün etmek, "
            "kompýuter enjamlaryny ulanmak we beýleki guramaçylyk çykdajylary bilen baglanyşyklydyr."
        ),
        145: (
            "Taslama smetasynyň esasynda IoT Security Scanner programma toplumyny döretmek üçin "
            "zerur tehniki serişdeleriň, programma gurallarynyň we goşmaça serişdeleriň sanawy "
            "kesgitlenýär."
        ),
        146: (
            "Bu sanawda kompýuter, test routeri, operasion ulgam we programma gurşawyny "
            "taýýarlamak üçin gerek bolan serişdeler görkezilýär."
        ),
        148: "Transport üçin çykdajylar materiallaryň we tehniki serişdeleriň bahasyndan kesgitlenen norma boýunça alynýar.",
        149: "= 399,00",
        150: (
            "Programma toplumyny işläp taýýarlamaga gatnaşýan hünärmenleriň kwalifikasiýasy we "
            "bir sagada tölegi kesgitlenip, zähmet haky hasaplanylýar."
        ),
        151: (
            "Zähmet haky esasy we goşmaça zähmet hakyndan ybaratdyr. Esasy zähmet haky tarif "
            "boýunça tölegden we baýrakdan, goşmaça zähmet haky bolsa rugsat, jemgyýetçilik "
            "işleri we beýleki tölegler bilen bagly böleklerden durýar."
        ),
        152: ", bu ýerde",
        153: "Ztar – tarif boýunça zähmet haky",
        154: "ortaça 1 sagatlyk töleg",
        155: "Td.b – programma üpjünçiligini işläp taýýarlamagyň zähmet talap edişi",
        158: "Ortaça tarif stawkasy aşakdaky formula boýunça kesgitlenilýär.",
        160: "Bu ýerde kesgitli iş üçin bir sagatlyk töleg ulanylýar.",
        161: "ortaça 1 sagatlyk töleg",
        162: "Baýrak gaznasynyň möçberi aşakdaky formula arkaly kesgitlenilýär:",
        157: "Z tar = 20,67*96=1984,32",
        159: "= 20,67",
        164: "=198,43",
        166: "Esasy zähmet haky tarif töleginiň we baýragyň jeminden ybaratdyr.",
        168: "Zes=Ztar+Zbaýr= 2182,75",
        170: "Goşmaça zähmet haky esasy zähmet hakynyň 10 göterimi hökmünde kabul edilýär.",
        172: "Zgoşm=Zesx0,1= 218,28",
        174: "Umumy zähmet haky aşakdaky ýaly kesgitlenilýär:",
        175: "Zum=Zes+ Zgoşm = 2401,03",
        176: "Sosial gorag üçin geçirimler umumy zähmet hakynyň 20 göterimi möçberinde alynýar.",
        177: "Ssos.gor= =480,21",
        178: "Bölüm boýunça harajatlar kärhanada kabul edilen norma esasynda hasaplanylýar.",
        179: "Çseh==327,41",
        180: "Kärhana boýunça çykdajylar esasy zähmet hakyna görä kabul edilen norma bilen kesgitlenýär.",
        182: "= 654,83",
        183: "= 982,24",
        185: "Netijede IoT Security Scanner programma taslamasy boýunça ähli harajatlar aşakdaky tablisa jemlenilýär",
        189: "         T öz.öd.= -------------= 2 ýyl 3 aý",
    }

    for original_index, text in replacements.items():
        shifted = original_index + 1
        if shifted < len(doc.paragraphs):
            if original_index in {0, 14, 29, 58, 91, 107, 125, 136}:
                set_heading_text(doc.paragraphs[shifted], text, size=13)
            else:
                set_text(doc.paragraphs[shifted], text)

    table0 = doc.tables[0]
    set_cell_text(table0.cell(1, 1), "Programmist")
    set_cell_text(table0.cell(1, 5), "22")
    set_cell_text(table0.cell(2, 1), "Dizaýner-programmist")
    set_cell_text(table0.cell(2, 5), "19")
    set_cell_text(table0.cell(3, 1), "Ulgam dolandyryjy")
    set_cell_text(table0.cell(3, 5), "21")
    set_cell_text(table0.cell(12, 5), "62")

    table1 = doc.tables[1]
    material_rows = [
        ("1", "Kompýuter ýa-da noutbuk", "San", "1", "3200", "3200"),
        ("2", "Test routeri we Wi-Fi enjamy", "San", "1", "420", "420"),
        ("3", "Windows OU", "San", "1", "120", "120"),
        ("4", "Resminama we PDF hasabat taýýarlamak serişdeleri", "Toplum", "1", "160", "160"),
        ("5", "Python, PyQt6, Nmap we SQLite açyk çeşmeli serişdeler", "Toplum", "1", "0", "0"),
        ("6", "USB göteriji we tor kabeli", "Toplum", "1", "90", "90"),
    ]
    for row_idx, row_values in enumerate(material_rows, start=1):
        for col_idx, value in enumerate(row_values):
            set_cell_text(table1.cell(row_idx, col_idx), value)
    set_cell_text(table1.cell(7, 1), "Jemi")
    set_cell_text(table1.cell(7, 3), "6")
    set_cell_text(table1.cell(7, 4), "")
    set_cell_text(table1.cell(7, 5), "3990")

    table2 = doc.tables[2]
    set_cell_text(table2.cell(1, 1), "Materiallar, tehniki we programma serişdeleri")
    totals = {
        1: ("Materiallar, tehniki we programma serişdeleri", "3990,00"),
        2: ("Transport harajatlary", "399,00"),
        3: ("Esasy zähmet haky", "2182,75"),
        4: ("Goşmaça zähmet haky", "218,28"),
        5: ("Sosial gorag üçin geçirimler", "480,21"),
        6: ("Seh boýunça çykdajylar", "327,41"),
        7: ("Kärhana boýunça çykdajylar", "654,83"),
        8: ("Enjamlara we programma gurşawyna çykdajylar", "982,24"),
        9: ("Jemi harajatlar", "9234,72"),
        10: ("Ýyllyk ykdysady netije", "4125,00"),
        11: ("Özüni ödeýän wagty (ýyl)", "2,24"),
    }
    for row_idx, (label, value) in totals.items():
        set_cell_text(table2.cell(row_idx, 1), label)
        set_cell_text(table2.cell(row_idx, 2), value)
    set_cell_text(table2.cell(8, 1), "Enjamlara we programma gurşawyna çykdajylar")

    return doc


def backup_main() -> Path:
    backup = MAIN_DOC.with_name(
        f"{MAIN_DOC.stem}_before_ykdysady_yjd_adapted_{datetime.now().strftime('%Y%m%d_%H%M%S')}{MAIN_DOC.suffix}"
    )
    shutil.copy2(MAIN_DOC, backup)
    return backup


def find_section_bounds(doc: Document) -> tuple[int, int]:
    body = doc._body._element
    start_idx = None
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            element_idx = body.index(paragraph._element)
            if element_idx > 300:
                start_idx = element_idx
                break

    if start_idx is None:
        raise RuntimeError("Ykdysady bölümiň başlangyjy tapylmady.")

    for paragraph in doc.paragraphs:
        element_idx = body.index(paragraph._element)
        if element_idx <= start_idx:
            continue
        text = " ".join(paragraph.text.split())
        style_name = paragraph.style.name if paragraph.style else ""
        if text == "Netije" and style_name.startswith("Heading"):
            return start_idx, element_idx

    raise RuntimeError("Ykdysady bölümiň soňy tapylmady.")


def normalize_heading_styles() -> None:
    doc = Document(str(MAIN_DOC))
    start_paragraph_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if idx > 300 and text == "Taslamanyň ykdysady netijeliliginiň hasaplamasy":
            start_paragraph_idx = idx
            break
    if start_paragraph_idx is None:
        return

    offsets = {
        0: "Heading 1",
        1: "Heading 1",
        15: "Heading 2",
        30: "Heading 2",
        59: "Heading 2",
        92: "Heading 2",
        108: "Heading 2",
        126: "Heading 2",
        137: "Heading 2",
    }
    for offset, style_name in offsets.items():
        idx = start_paragraph_idx + offset
        if 0 <= idx < len(doc.paragraphs):
            doc.paragraphs[idx].style = style_name
    doc.save(str(MAIN_DOC))


def replace_main_section(section_doc: Document) -> Path:
    backup = backup_main()
    main = Document(str(MAIN_DOC))
    body = main._body._element
    start_idx, end_idx = find_section_bounds(main)

    for element in list(body)[start_idx:end_idx]:
        body.remove(element)

    netije_element = list(body)[start_idx]
    for child in list(section_doc._body._element):
        if child.tag.endswith("sectPr"):
            continue
        body.insert(body.index(netije_element), copy.deepcopy(child))

    main.save(str(MAIN_DOC))
    normalize_heading_styles()
    return backup


def main() -> None:
    doc = build_adapted_doc()
    doc.save(str(OUT_DOC))
    backup = replace_main_section(doc)
    print(f"Created: {OUT_DOC}")
    print(f"Updated: {MAIN_DOC}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
