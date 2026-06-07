from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = ROOT / "kakabalowa.docx"


EXTRA = {
    "3.6. Tor adapterlerini kesgitlemek we IP aralygyny awtomatik saýlamak": [
        "Adapter maglumatlary alnandan soň programma bu maglumatlary diňe interfeýsde görkezmek bilen çäklenmeýär, olary skanirleme prosesiniň içerki parametrleri hökmünde hem ulanýar. Şeýlelikde, ulanyjy adapteri üýtgedende diňe ýazgy däl, eýsem skanirlenjek tor segmenti, şlýuz maglumatlary we lokal kompýuteriň IP salgysy hem täzelenýär. Bu baglanyşyk programma ýadrosy bilen interfeýsiň arasynda bitewi maglumat akymynyň bardygyny görkezýär.",
        "Daşarky ulgam buýruklarynyň gizlin ýerine ýetirilmegi hem bu bölegiň möhüm tarapy bolup durýar. Windows gurşawynda `ipconfig`, `route print`, `arp -a` ýa-da Nmap ýaly prosesler işledilende gara konsol penjiresiniň açylmagy ulanyjy tejribesini peseldip biler. Şonuň üçin programma prosesleri `CREATE_NO_WINDOW` düzgüni bilen işledýär. Netijede ulanyjy diňe grafiki interfeýsi görýär, tehniki prosesler bolsa fon režiminde ýerine ýetirilýär.",
    ],
    "3.7. Gurluşlary tapmagyň birleşdirilen usullary": [
        "Birleşdirilen usullaryň netijeleri alnan soň programma şol maglumatlary gaýtalanmaz ýaly birleşdirýär. Mysal üçin, şol bir IP salgy Nmap, ARP cache we ping arkaly bir wagtyň özünde tapylyp biler. Şeýle ýagdaýda tablisa birnäçe setir goşulmaýar, tersine, bir gurluşyň `Tapylan usul` meýdanynda birnäçe usul görkezilýär. Bu çözgüt netijeleriň arassa görünmegini üpjün edýär we ulanyjynyň şol bir enjamy birnäçe gezek hasaplamagynyň öňüni alýar.",
        "Gurluş tapmak bilen port barlagynyň aýratyn tapgyr hökmünde bölünmegi hem amaly taýdan peýdalydyr. Eger gurluş torda ýüze çykyp, emma açyk port tapylmasa, ol şonda-da tablisa goşulýar. Bu aýratynlyk telefonlar we käbir IoT datçikleri üçin möhümdir. Şeýle enjamlarda açyk TCP hyzmat bolmasa-da, olaryň torda bardygyny görkezmek administrator üçin peýdaly maglumat bolup durýar.",
    ],
    "3.8. Programma modulynyň giňeldilmegi we goldaw mümkinçilikleri": [
        "Modul giňeldilmegi diňe plugin goşmak bilen çäklenmeýär. Geljekde sazlamalar sahypasy arkaly port sanawyny, timeout bahalaryny, skanirleme akymlarynyň sanyny ýa-da hasabat faýlynyň saklanýan ýerini üýtgetmek mümkinçiligi hem döredilip bilner. Bu ýagdaý programmany dürli tor ölçeglerine uýgunlaşdyrmaga kömek eder. Kiçi öý torunda çalt barlag ýeterlik bolsa, uly ofis segmentinde has uzyn timeout we has giň port sanawy gerek bolup biler.",
        "Goldaw mümkinçilikleriniň ýene bir ugry loglaryň saklanylmagydyr. Häzirki programma skanirleme wagtynda log setirlerini interfeýsde görkezýär, emma geljekde bu loglary aýratyn faýla ýazmak ýa-da hasabat içine goşmak mümkin. Şeýle mümkinçilik näsazlyk ýüze çykanda sebäbini seljermäge, gorag wagtynda bolsa programmanyň nähili tapgyrlar boýunça işländigini görkezmäge kömek eder.",
    ],
    "4.6. Interfeýsde maglumatlaryň okalmagy we ulanyjy hereketleriniň tertibi": [
        "Interfeýs elementleriniň ýerleşişinde işiň ýygylygy hem göz öňünde tutuldy. Skanirlemäni başlatmak, saklamak we eksport etmek ýaly esasy amallar aşaky bölekde ýerleşdirildi, sebäbi ulanyjy parametrleri barlandan soň tebigy ýagdaýda aşakdaky düwmelere geçýär. Progress bar netijeler tablisasynyň ýokarsynda ýerleşýär, bu bolsa skanirlemäniň ýagdaýyny gözden sypdyrmazlyga mümkinçilik berýär.",
        "Log paneli hem ulanyjy üçin düşündiriş gatlagy hökmünde çykyş edýär. Eger tablisa diňe sanlary görkezýän bolsa, log paneli şol sanlaryň nähili emele gelendigini düşündirýär. Mysal üçin, `Adapter saýlandy`, `Gateway goşuldy`, `ARP arkaly tapyldy` ýa-da `Skanirleme tamamlandy` ýaly habarlar ulanyja programma prosesiniň yzygiderliligini görkezýär.",
    ],
    "4.7. Netijeleriň tablisa görnüşinde berilmegi we sahypalaryň tapawutlandyrylmagy": [
        "Tablisada sütünleriň sany köp bolsa-da, olaryň her biri aýratyn maksat üçin ulanylýar. IP salgy gurluşy tanamak üçin, MAC salgy we öndüriji gurluşyň gelip çykyşyny çaklamak üçin, görnüş meýdany bolsa onuň router, PC, telefon ýa-da IoT kandidatdygyny düşündirmek üçin gerekdir. Portlar meýdany açyk hyzmatlary görkezýär, töwekgelçilik we maslahat meýdanlary bolsa ulanyjynyň indiki hereketini kesgitlemäge kömek edýär.",
        "Sahypalaryň tapawutlandyrylmagy maglumatlaryň gaýtalanmagyny azaltýar. Skanirleme sahypasy proses üçin, Netijeler sahypasy jikme-jik analiz üçin, Taryh sahypasy öňki barlaglar üçin, Panel sahypasy bolsa umumy ýagdaý üçin ulanylýar. Şeýle bölünişik programma bilen uzak wagt işlenende hem tertibi saklaýar.",
    ],
    "4.8. Ulanyjy tejribesini ýokarlandyrmak üçin kabul edilen çözgütler": [
        "Ulanyjy tejribesinde ýalňyşlyk ýagdaýlarynyň görkezilişi aýratyn orun tutýar. Programma nädogry IP aralygy girizilende, Nmap tapylmadyk ýagdaýynda ýa-da skanirleme netijesi örän az bolanda ulanyja düşündiriş bermelidir. Şeýle habarlar tehniki ýalňyşlykdan gorkuzman, ony düzetmegiň ýoluny görkezýär. Bu çemeleşme okuw taslamasy üçin hem peýdalydyr, sebäbi ulanyjy diňe netijäni däl, sebäbini hem öwrenýär.",
        "Interfeýsde goýy tema saýlanmagy hem tötänleýin däldir. Tor skanirlemesi wagtynda tablisa setirleri, loglar we risk maglumatlary köp okalýar. Gara fon we açyk ak-gök tekst kontrasty maglumatlary tapawutlandyrýar, aktiw düwmeler bolsa gök reňk bilen bellenýär. Şeýlelikde, dizaýn diňe owadanlyk üçin däl, iş prosesiniň düşnükliligi üçin ulanylýar.",
    ],
    "5.6. Lokal torlarda goşmaça synag ssenariýalary": [
        "Goşmaça ssenariýalarda nädogry IP aralygy hem barlandy. Ulanyjy `192.168.1.0/33` ýa-da tekst görnüşli nädogry aralyk girizse, programma skanirlemäni başlamazdan öň duýduryş bermelidir. Bu barlag programma ýadrosynyň diňe üstünlikli ýagdaýlar üçin däl, eýsem nädogry giriş maglumatlary üçin hem taýýardygyny görkezýär.",
        "Başga bir synagda skanirleme wagtynda tordaky gurluşlaryň käbiri jogap bermeýän ýagdaý kabul edildi. Bu ýagdaý amaly torlarda ýygy duş gelýär: enjam uky režiminde bolup biler, firewall ping-i bloklap biler ýa-da router müşderileri izolirleýär. Programma şeýle ýagdaýda işini tamamlamaly we tapylan maglumatlary ýitirmän görkezmelidir.",
    ],
    "5.7. Tapylan gurluşlaryň dogrulygyny barlamak usuly": [
        "Dogrylyk barlagynda wagt görkezijisi hem hasaba alyndy. Birmeňzeş /24 tor segmenti el bilen barlananda ähli IP salgylary, portlar we maslahatlar aýratynlykda ýazylýar. Programma arkaly bolsa şol maglumatlar awtomatik ýygnalýar we tablisa geçirilýär. Bu deňeşdirme programmanyň diňe takyklyk taýdan däl, wagt tygşytlamak taýdan hem peýdalydygyny görkezýär.",
        "Netijeleri tassyklamakda PDF hasabat hem ulanylyp bilner. Hasabatda tapylan gurluşlar, howpsuz we gowşak ýagdaýlar, portlar we diagrammalar görkezilýär. Şeýle dokument soňra barlagyň subutnamasy hökmünde saklanyp bilner. Bu aýratynlyk programmanyň audit prosesine laýyk gelýändigini görkezýär.",
    ],
    "5.8. Synag netijeleriniň diplom goragynda görkezilişi": [
        "Goragda synag netijelerini görkezmegiň iň amatly tertibi öňünden taýýarlanylýar. Ilki programma açylyp, gara konsol penjiresiniň çykmaýandygy görkezilýär. Soňra skanirleme sahypasynda Wi‑Fi adapteri we IP aralygy düşündirilýär. Ondan soň skanirleme başladylyp, progress bar we log paneli görkezilýär. Bu yzygiderlilik komissiýa üçin programma işiniň sebäpli we düşnükli görünmegine kömek edýär.",
        "Gorag wagtynda internet ýa-da tor ýagdaýy üýtgäp bilýändigi üçin öňki skanirleme taryhynyň bolmagy hem peýdalydyr. Eger häzirki wagtda torda az gurluş görünse, Taryh sahypasy öňki üstünlikli barlagy görkezmäge mümkinçilik berýär. Şeýlelikde, programma diňe häzirki pursatdaky netijä bagly bolman, öňki audit maglumatlaryny hem saklaýar.",
    ],
}


def insert_before(doc: Document, element, text: str) -> None:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def find_next_heading_element(doc: Document, heading: str):
    heading_pos = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if paragraph.style.name.startswith("Heading") and text == heading:
            heading_pos = i
            break
    if heading_pos is None:
        raise ValueError(f"Heading not found: {heading}")
    for i in range(heading_pos + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        if paragraph.style.name.startswith("Heading"):
            return paragraph._p
    raise ValueError(f"Next heading not found after: {heading}")


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC_PATH.with_name(f"kakabalowa_before_expand_sections_3_5_extra_{timestamp}.docx")
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))
    for heading, paragraphs in EXTRA.items():
        anchor = find_next_heading_element(doc, heading)
        for text in paragraphs:
            insert_before(doc, anchor, text)

    doc.save(str(DOC_PATH))
    print(f"Backup: {backup}")
    print(f"Updated: {DOC_PATH}")


if __name__ == "__main__":
    main()
