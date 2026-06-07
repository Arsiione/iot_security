from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_2_4_2_5.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def add_table(doc, caption_text, headers, rows):
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "2.4. Skanirleme algoritminiň taslanylyşy")

    section_24_paragraphs = [
        (
            "IoT Security Scanner programma toplumynyň esasy wezipesi lokal tordaky işjeň gurluşlary "
            "ýüze çykarmak, olaryň portlaryny barlamak we netijäni ulanyja düşnükli görnüşde görkezmekdir. "
            "Bu maksat bilen skanirleme algoritmi birnäçe tapgyrdan ybarat edildi. Algoritm diňe bir Nmap "
            "netijesine daýanman, ARP, Ping, Windows ARP keşi, lokal kompýuter we şlýuz maglumatlaryny hem "
            "birleşdirýär. Şeýle birleşdirilen çemeleşme Wi-Fi torunda port açmaýan telefon ýa-da akylly "
            "gurluş ýaly enjamlary hem netijeler tablisasyna goşmaga mümkinçilik berýär."
        ),
        (
            "Skanirleme prosesi ulanyjynyň tor adapterini saýlamagyndan başlanýar. Programma Windows "
            "ulgamyndan işjeň adapterleri alýar, olaryň IP salgysyny, maskasyny, şlýuzyny we tor aralygyny "
            "kesgitleýär. VPN adapterleri, 169.254.x.x görnüşli awtomatik IP salgylary we loopback görnüşli "
            "adapterler adaty lokal tor hökmünde saýlanmaýar. Eger kompýuterde Wi-Fi ýa-da Ethernet adapteri "
            "bar bolsa, programma şol adapteri ileri tutýar we IP aralygyny awtomatik doldurýar. Mysal üçin, "
            "kompýuteriň salgysy 192.168.1.3/24 bolsa, skanirleme aralygy 192.168.1.0/24 görnüşinde "
            "kesgitlenýär."
        ),
        (
            "Algoritmiň indiki tapgyry işjeň hostlary ýüze çykarmakdyr. Ilki bilen Nmap arkaly ping scan "
            "geçirilýär: nmap -sn -n parametrleri torda jogap berýän hostlary tapmaga kömek edýär. Lokal "
            "hususy torlarda goşmaça ARP scan ulanylýar: nmap -PR -sn -n. Bu usul Wi-Fi ýa-da Ethernet "
            "segmentinde MAC derejesindäki jogaplary görmäge mümkinçilik berýär. Eger Nmap elýeterli bolmasa "
            "ýa-da az host tapsa, programma ping sweep we Windows ARP keşi arkaly goşmaça maglumat ýygnaýar."
        ),
        (
            "Programma diňe daşky skanirleme netijelerine daýanmaýar. Lokal kompýuteriň IP salgysy we esasy "
            "şlýuz, ýagny router salgysy, goşmaça usul hökmünde netijelere girizilýär. Bu aýratynlyk diplom "
            "goragy wagtynda peýdalydyr, sebäbi skanirleme netijesinde iň azyndan kompýuter we router ýaly "
            "esasy gurluşlaryň görkezilmegine kömek edýär. Eger torda diňe bir gurluş tapylsa, programma "
            "ulanyja Wi-Fi aralygyny, Guest Wi-Fi ýa-da AP isolation sazlamasyny we telefonyň torda jogap "
            "berip-bermeýändigini barlamagy maslahat berýär."
        ),
        (
            "Hostlar tapylandan soň her bir IP salgy aýratynlykda barlanýar. ScanThread synpy bu işi aýratyn "
            "akymda ýerine ýetirýär, şonuň üçin skanirleme wagtynda grafiki interfeýs doňup galmaýar. "
            "ThreadPoolExecutor arkaly birnäçe host bir wagtda barlanýar. Şeýle köp akymly çemeleşme "
            "192.168.1.0/24 ýaly adaty lokal tor aralygynda skanirleme wagtyny azaltmaga mümkinçilik berýär."
        ),
        (
            "Her bir tapylan gurluş üçin programma host adyny, MAC salgysyny, MAC boýunça öndürijini, "
            "tapylan usulyny, portlaryny, hyzmatlaryny we gurluş görnüşini kesgitlemäge synanyşýar. "
            "Gurluşyň görnüşi Router, PC, Telefon/Unknown, IoT kandidat ýa-da Unknown bolup biler. Mysal üçin, "
            "gurluşyň IP salgysy saýlanan adapteriň şlýuzy bilen gabat gelse, ol Router hökmünde bellenýär. "
            "Eger portlaryň arasynda HTTP, RTSP, MQTT ýa-da Telnet ýaly IoT bilen bagly hyzmatlar bar bolsa, "
            "gurluş IoT kandidat hökmünde görkezilýär."
        ),
        (
            "Port skanirlemesi ýörite IoT gurşawynda ýygy duş gelýän portlaryň sanawy boýunça geçirilýär. "
            "Programma 21, 22, 23, 80, 443, 554, 8000, 8080, 8888, 1883, 8883, 5683, 1900, 49152 we 5353 "
            "portlaryny barlaýar. Bu portlar FTP, SSH, Telnet, HTTP/HTTPS, RTSP, MQTT, CoAP, SSDP we mDNS "
            "ýaly hyzmatlar bilen baglydyr. Port açyk bolsa, programma ony netijeler tablisasyna ýazýar; "
            "port açyk bolmasa hem gurluş netijelerden aýrylmaýar."
        ),
        (
            "Bu algoritmiň möhüm aýratynlygy gurluş tapmak bilen port barlagynyň aýratyn tapgyrlar hökmünde "
            "seredilmegidir. Adaty port skanirlemede port açmadyk telefon ýa-da akylly enjam görünmän galyp "
            "bilýär. Bu taslamada bolsa enjam ARP ýa-da Ping arkaly tapylan bolsa, onuň portlary açyk bolmasa "
            "hem tablisa goşulýar. Şeýle ýagdaýda portlar sütunynda “ýok”, töwekgelçilik derejesinde “pes” "
            "we maslahat hökmünde “Portlar açyk däl, gurluş diňe torda ýüze çykaryldy” görkezilýär."
        ),
        (
            "Skanirleme tamamlanandan soň tapylan gurluşlar interfeýse signal arkaly iberilýär we şol bir "
            "wagtda maglumatlar bazasynda saklanýar. Bu netijeler soňra Panel sahypasynda umumy statistika, "
            "Netijeler sahypasynda jikme-jik seljerme we Taryh sahypasynda öňki skanirlemeler görnüşinde "
            "görkezilýär. Şeýlelikde, skanirleme algoritmi diňe bir hostlary tapmak bilen çäklenmän, programma "
            "toplumynyň ähli beýleki bölekleri üçin maglumat çeşmesi bolup hyzmat edýär."
        ),
    ]
    for text in section_24_paragraphs:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 2.4 - Skanirleme algoritminiň esasy tapgyrlary",
        ["Tapgyr", "Ýerine ýetirilýän amal", "Netije"],
        [
            ["1", "Tor adapterleri kesgitlenýär we Wi-Fi/Ethernet ileri tutulýar", "Dogry lokal IP aralygy saýlanýar"],
            ["2", "IP aralygy hasaplanýar", "Mysal üçin 192.168.1.0/24 görnüşli tor kesgitlenýär"],
            ["3", "Nmap ping scan ýerine ýetirilýär", "Jogap berýän hostlar tapylýar"],
            ["4", "Nmap ARP scan, ARP keş we Ping sweep ulanylýar", "Port açmaýan gurluşlary hem tapmak mümkinçiligi artýar"],
            ["5", "Lokal kompýuter we şlýuz netijelere goşulýar", "PC we router ýaly esasy gurluşlar görkezilýär"],
            ["6", "Her host üçin MAC, öndüriji, host ady we portlar barlanýar", "Gurluş barada tehniki maglumatlar ýygnalýar"],
            ["7", "Hyzmatlar we gurluş görnüşi kesgitlenýär", "Router, PC, IoT kandidat ýa-da Telefon/Unknown ýaly görnüş bellenýär"],
            ["8", "Plugin barlaglary işledilýär", "Telnet ýa-da kamera bilen bagly gowşaklyklar ýüze çykarylyp bilýär"],
            ["9", "Netijeler interfeýse we SQLite bazasyna geçirilýär", "Panel, Netijeler, Taryh we PDF hasabat üçin maglumat taýýarlanýar"],
        ],
    )

    add_normal(
        doc,
        "Netijede, skanirleme algoritmi diplom taslamasynyň esasy tehniki bölegini emele getirýär. "
        "Ol diňe Nmap komandasynyň daşky görnüşi däl-de, eýsem birnäçe maglumat çeşmesini birleşdirýän, "
        "interfeýs bilen signal arkaly aragatnaşyk saklaýan we netijeleri maglumatlar bazasyna geçirýän "
        "programma logikasydyr. Bu çemeleşme IoT toruny has doly görkezmäge we diplom goragynda algoritmi "
        "aýdyň düşündirmäge mümkinçilik berýär."
    )

    add_heading(doc, "2.5. Töwekgelçiligi bahalandyrmak we remediation modeliniň taslanylyşy")

    section_25_paragraphs = [
        (
            "IoT Security Scanner programmasynda skanirleme netijeleriniň diňe sanaw görnüşinde berilmegi "
            "ýeterlik hasaplanmaýar. Ulanyjy tapylan gurluşyň näderejede howply bolup biljekdigini we "
            "indiki ädimde näme etmelidigini hem görmeli. Şonuň üçin taslamada töwekgelçiligi bahalandyrmak "
            "we remediation, ýagny gowşaklyklary azaltmak boýunça maslahat bermek modeli göz öňünde tutuldy."
        ),
        (
            "Töwekgelçiligi bahalandyrmak modeli üç esasy derejä esaslanýar: pes, orta we ýokary. Pes "
            "dereje gurluş torda ýüze çykarylan, emma açyk port ýa-da anyk gowşaklyk tapylmadyk ýagdaýlarda "
            "ulanylýar. Orta dereje açyk web interfeýs, kamera bilen bagly portlar ýa-da firmware barlagyny "
            "talap edýän hyzmatlar ýüze çykanda bellenip bilýär. Ýokary dereje bolsa gowşak autentifikasiýa, "
            "Telnet arkaly giriş mümkinçiligi ýa-da standart parollar ýaly has howply ýagdaýlar bilen "
            "baglanyşyklydyr."
        ),
        (
            "Programmada töwekgelçiligiň esasy çeşmesi açyk portlar we plugin barlaglarynyň netijeleridir. "
            "Mysal üçin, 80-nji portuň açyk bolmagy web interfeýsiň bardygyny görkezýär. Hikvision plugininiň "
            "logikasy şeýle ýagdaýda kamera ýa-da web dolandyryş interfeýsi bilen bagly töwekgelçiligi orta "
            "derejä çykaryp, firmware täzelenmesini we standart hasaplary barlamagy maslahat berýär. 23-nji "
            "port açyk bolsa, Telnet plugininiň barlagy işledilýär. Eger gowşak login/parol jübüti kabul "
            "edilse, töwekgelçilik ýokary derejä çenli ýokarlanýar."
        ),
        (
            "Remediation modeli iki ugurdan durýar: maslahat beriş we awtomatlaşdyrylan düzediş mümkinçiligi. "
            "Maslahat beriş bölegi häzirki programma üçin esasy amaly netijedir. Her gurluş üçin programma "
            "portlary ýapmak, Telnet hyzmatyny öçürmek, standart parollary üýtgetmek, firmware täzelenmesini "
            "barlamak, HTTP ýerine HTTPS ulanmak ýa-da router izolýasiýa sazlamalaryny gözden geçirmek ýaly "
            "maslahatlary berýär."
        ),
        (
            "Awtomatlaşdyrylan düzediş mümkinçiligi core/remediation_engine.py modulynyň üsti bilen "
            "taslanyldy. Bu modul telnet_weak_auth, ftp_plaintext, http_default_credentials we "
            "weak_ssh_config ýaly düzediş görnüşlerini saklaýar. Mysal üçin, FTP hyzmatyny SFTP bilen "
            "çalşyrmak, web interfeýsinde standart paroly üýtgetmek ýa-da SSH konfigurasiýasyny berkitmek "
            "ýaly hereketler bu modeliň geljekde giňeldilip bilinjek ugurlarydyr. Şeýle-de bolsa, IoT "
            "gurluşlarynda awtomatik düzediş öndüriji, model we firmware wersiýasyna bagly bolýandygy üçin "
            "programma ilki bilen howpsuz maslahat beriş çemeleşmesini öňe çykarýar."
        ),
        (
            "Remediation modeliniň möhüm talaby ulanyja zyýan ýetirmezlikdir. Tor enjamlarynyň sazlamalaryny "
            "awtomatik üýtgetmek käbir ýagdaýlarda enjamyň işini bozup biler. Şonuň üçin diplom taslamasynda "
            "awtomatik düzedişler başlangyç model hökmünde görkezilýär, esasy iş bolsa gowşaklygy ýüze "
            "çykarmak we administratora anyk, ýerine ýetirip bolýan maslahat bermek bilen çäklenýär. Bu "
            "çemeleşme okuw we diplom goragy şertlerinde has ygtybarlydyr."
        ),
        (
            "Töwekgelçilik we maslahat maglumatlary interfeýsde aýratyn sütünler hökmünde görkezilýär. "
            "Ulanyjy her gurluş üçin IP, MAC, portlar, tapylan usul, gurluş görnüşi, töwekgelçilik derejesi "
            "we maslahatlary bir setirde görýär. Şeýle gurluş administratora netijeleri çalt seljermäge "
            "we ilkinji nobatda haýsy gurluşa üns bermelidigini kesgitlemäge kömek edýär."
        ),
    ]
    for text in section_25_paragraphs:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 2.5 - Töwekgelçilik derejeleri we maslahat modeli",
        ["Dereje", "Şert ýa-da alamat", "Programmada berilýän maslahat"],
        [
            [
                "Pes",
                "Gurluş torda tapyldy, emma açyk port ýa-da anyk gowşaklyk ýok",
                "Gurluş diňe torda ýüze çykaryldy, goşmaça howply hyzmat tapylmady",
            ],
            [
                "Orta",
                "HTTP/HTTPS, RTSP, MQTT ýa-da kamera bilen bagly hyzmatlar açyk",
                "Firmware-i barlamak, standart hasaplary öçürmek we elýeterliligi çäklendirmek",
            ],
            [
                "Ýokary",
                "Telnet açyk we gowşak autentifikasiýa alamaty ýüze çykýar",
                "Telnet hyzmatyny öçürmek, güýçli parol goýmak we SSH/HTTPS ulanmak",
            ],
            [
                "Remediation maslahaty",
                "Port ýa-da hyzmat boýunça gowşaklyk ähtimallygy bar",
                "Ulanyja anyk ýerine ýetirilýän howpsuzlyk çäresi görkezilýär",
            ],
            [
                "Awtomatik düzediş modeli",
                "SSH ýa-da web interfeýs arkaly düzediş tehniki taýdan mümkin bolan ýagdaý",
                "RemediationEngine modulynyň üsti bilen geljekde giňeldilip bilinýär",
            ],
        ],
    )

    add_normal(
        doc,
        "Şeýlelikde, töwekgelçiligi bahalandyrmak we remediation modeli programma toplumyny ýönekeý tor "
        "skanerinden tapawutlandyrýar. Programma diňe IP we port sanawyny görkezmän, her netijäniň "
        "manysyny düşündirýär we ulanyja howpsuzlygy ýokarlandyrmak üçin amaly ugur görkezýär. Bu bolsa "
        "diplom taslamasynyň esasy ideýasyna, ýagny IoT torlarynyň awtomatlaşdyrylan auditi we gowşaklyklary "
        "azaltmak maksadyna laýyk gelýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_styles(main_doc):
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("2.4.") or text.startswith("2.5."):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("2.4.") for p in main_doc.paragraphs) or any(
        p.text.strip().startswith("2.5.") for p in main_doc.paragraphs
    ):
        print("Section 2.4 or 2.5 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_styles(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
