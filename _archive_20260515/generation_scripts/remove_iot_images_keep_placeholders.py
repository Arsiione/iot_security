from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_FILE = Path("bolum_1_1_iot_gurlushlary.docx")
BACKUP_FILE = Path("bolum_1_1_iot_gurlushlary_before_image_placeholders.docx")


def has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_placeholder_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")

    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "10")
        border.set(qn("w:color"), "94A3B8")
        borders.append(border)

    p_pr.append(borders)


def set_placeholder_spacing(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "720")
    spacing.set(qn("w:after"), "720")
    spacing.set(qn("w:line"), "720")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)


def replace_images_with_placeholders():
    if not BACKUP_FILE.exists():
        BACKUP_FILE.write_bytes(DOCX_FILE.read_bytes())

    document = Document(DOCX_FILE)
    replaced = 0

    for paragraph in document.paragraphs:
        if not has_drawing(paragraph):
            continue

        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_placeholder_border(paragraph)
        set_placeholder_spacing(paragraph)

        run = paragraph.add_run("Surat üçin ýer")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = True
        replaced += 1

    document.save(DOCX_FILE)
    print(f"placeholders={replaced}")


if __name__ == "__main__":
    replace_images_with_placeholders()
