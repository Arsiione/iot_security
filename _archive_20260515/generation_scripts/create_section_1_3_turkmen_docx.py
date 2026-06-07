from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_FILE = "bolum_1_3_iot_hujum_wektorlary.docx"
TITLE = "1.3. IoT torlarynda howpsuzlyk gowşaklyklary we hüjüm wektorlary"


PARAGRAPHS = [
    (
        "IoT torlarynyň howpsuzlygyny seljermekde diňe gurluşlaryň sanyny ýa-da olaryň "
        "ulanylyş gerimini görkezmek ýeterlik däldir. Esasy mesele şol gurluşlaryň nähili "
        "hyzmatlary işledýändigini, haýsy portlarynyň açykdygyny we bu hyzmatlaryň näderejede "
        "howpsuz sazlanandygyny kesgitlemekden ybaratdyr. Şonuň üçin IoT howpsuzlygynda "
        "gowşaklyklar köplenç tor derejesindäki açyk hyzmatlar, standart giriş maglumatlary "
        "we wagtynda täzelenmedik firmware bilen baglanyşykly bolýar."
    ),
    (
        "IoT gurluşlaryna edilýän hüjümleriň köpüsi awtomatlaşdyrylan gözlegden başlanýar. "
        "Hüjümçi ýa-da zyýanly programma belli bir IP aralygyny skanirläp, açyk portlary "
        "tapýar. Eger gurluşda Telnet, HTTP, FTP, RTSP ýa-da MQTT ýaly hyzmatlar goragsyz "
        "ýagdaýda açyk bolsa, şol hyzmatlar hüjümiň başlangyç nokadyna öwrülip biler. "
        "Şu sebäpli diplom işinde işlenip taýýarlanylýan IoT Security Scanner programmasy "
        "ilkinji nobatda lokal tordaky açyk portlary we hyzmatlary kesgitlemäge gönükdirilýär."
    ),
    (
        "Telnet protokoly IoT torlarynda iň howply hyzmatlaryň biri hasaplanýar. Bu protokol "
        "maglumatlary şifrlemezden geçirýär we köp köne IoT gurluşlarynda henizem ulanylýar. "
        "Eger Telnet hyzmaty açyk bolsa we gurluşda admin/admin, root/root ýa-da şuňa meňzeş "
        "standart parollar saklanyp galan bolsa, hüjümçi gurluşa aňsatlyk bilen girip biler. "
        "Şonuň üçin programma toplumynda Telnet portunyň barlagy we gowşak giriş maglumatlaryny "
        "ýüze çykarmaga niýetlenen plugin aýratyn ähmiýete eýedir."
    ),
    (
        "HTTP we HTTPS web-interfeýsleri IoT gurluşlaryny dolandyrmak üçin giňden ulanylýar. "
        "IP kameralar, marşrutizatorlar, akylly öý enjamlary we beýleki gurluşlar köplenç "
        "web-panel arkaly sazlanýar. Emma web-interfeýsiň açyk bolmagy, standart login we "
        "parollaryň üýtgedilmezligi ýa-da köne firmware ulanylmagy rugsatsyz giriş howpuny "
        "ýokarlandyrýar. HTTP hyzmaty şifrlenmedik ýagdaýda işleýän bolsa, login maglumatlary "
        "torda ele geçirilip bilner."
    ),
    (
        "RTSP protokoly esasan IP kameralarda wideo akymyny geçirmek üçin ulanylýar. Bu hyzmat "
        "nädogry sazlanan ýagdaýynda kamera akymyna rugsatsyz giriş mümkin bolup biler. "
        "Şeýle ýagdaý diňe maglumat howpsuzlygyna däl, eýsem şahsy durmuşyň eldegrilmesizligine "
        "hem howp salýar. Şonuň üçin 554-nji portuň açykdygyny barlamak IoT torunyň wideo "
        "gözegçilik enjamlary bilen baglanyşykly töwekgelçiliklerini kesgitlemekde peýdalydyr."
    ),
    (
        "MQTT protokoly IoT datçikleriniň we dolandyryş ulgamlarynyň arasynda ýeňil habar "
        "alyş-çalyşygy üçin ulanylýar. Bu protokol dogry sazlananda amatly we netijeli "
        "çözgüt bolup durýar. Emma awtorizasiýa goýulmasa ýa-da broker daşarky tora açyk "
        "bolsa, hüjümçi datçiklerden gelýän maglumatlary okap, üýtgedip ýa-da galp buýruklar "
        "iberip biler. Şonuň üçin MQTT hyzmatynyň açyk portlaryny barlamak IoT howpsuzlyk "
        "auditiniň möhüm bölegidir."
    ),
    (
        "FTP we SSH hyzmatlary käbir IoT gurluşlarynda faýl geçirmek ýa-da uzakdan dolandyrmak "
        "üçin ulanylýar. FTP şifrlenmedik protokol bolany sebäpli login we parollar açyk görnüşde "
        "geçip biler. SSH has howpsuz protokol hasaplansa-da, gowşak parol ýa-da nädogry sazlama "
        "bar bolsa, ol hem hüjüm üçin ulanylyp bilner. Şeýle hyzmatlaryň tapylmagy administratora "
        "gurluşyň dolandyryş kanallaryny täzeden barlamaga mümkinçilik berýär."
    ),
    (
        "IoT ulgamyndaky gowşaklyklar köplenç diňe bir enjam bilen çäklenmeýär. Bir gurluş ele "
        "geçirilenden soň, hüjümçi ony lokal tory öwrenmek, beýleki gurluşlara geçmek ýa-da "
        "botnet düzümine goşmak üçin ulanyp biler. Taryhda Mirai botneti ýaly hüjümler goragsyz "
        "IoT enjamlarynyň köpçülikleýin ele geçirilip, DDoS hüjümlerinde ulanylandygyny görkezdi. "
        "Bu ýagdaý IoT gurluşlarynyň global kiberhowpsuzlyga hem täsir edip bilýändigini subut edýär."
    ),
    (
        "IP kameralar bilen baglanyşykly gowşaklyklar aýratyn üns berilmäge mynasypdyr. Hikvision, "
        "Dahua we beýleki öndürijileriň kameralarynda web-interfeýs, RTSP akymy, firmware "
        "täzelenmeleri we giriş maglumatlary bilen baglanyşykly töwekgelçilikler ýüze çykyp "
        "bilýär. Şeýle enjamlaryň açyk portlaryny ýüze çykarmak administratora firmware-i "
        "täzelemäge, standart parollary üýtgetmäge we daşarky elýeterliligi çäklendirmäge "
        "mümkinçilik berýär."
    ),
    (
        "Diplom taslamasyndaky IoT Security Scanner programmasy ýokarda agzalan howplary "
        "başlangyç derejede ýüze çykarmak üçin döredilýär. Programma lokal tordaky IP "
        "salgylary skanirleýär, 21, 22, 23, 80, 443, 554, 8000, 8080, 1883, 8883 we beýleki "
        "IoT bilen baglanyşykly portlary barlaýar, hyzmatlaryň görnüşini kesgitleýär we "
        "pluginleriň kömegi bilen potensial töwekgelçilikleri belleýär."
    ),
    (
        "Programmada ulanylýan plugin çemeleşmesi aýratyn ähmiýete eýedir. Sebäbi IoT "
        "howpsuzlygynda gowşaklyklar öndüriji, model we firmware wersiýasyna baglylykda "
        "üýtgäp durýar. Plugin ulgamy täze barlaglary goşmaga, mysal üçin Telnet gowşak "
        "parollaryny, açyk HTTP interfeýslerini ýa-da belli öndürijilere mahsus alamatlary "
        "aýratynlykda barlamaga mümkinçilik berýär. Bu bolsa programma toplumynyň geljekde "
        "giňeldilip bilinjekdigini görkezýär."
    ),
    (
        "Şeýlelikde, IoT torlarynda esasy hüjüm wektorlary açyk portlar, goragsyz protokollar, "
        "standart parollar, täzelenmedik firmware we nädogry sazlanan hyzmatlar bilen "
        "baglanyşyklydyr. Bu howplary doly aradan aýyrmak üçin diňe bir gurluşy tapmak "
        "ýeterlik däldir; onuň işleýän hyzmatlaryny seljermek, töwekgelçilik derejesini "
        "bahalandyrmak we ulanyja anyk maslahat bermek zerurdyr. IoT Security Scanner "
        "taslamasynyň esasy maksady hem şu amallary awtomatlaşdyrmakdan ybaratdyr."
    ),
]


TABLE_ROWS = [
    ("Hüjüm wektory", "Baglanyşykly port/protokol", "Howpuň mazmuny", "Programmada barlag usuly"),
    (
        "Gowşak Telnet parollary",
        "23/Telnet",
        "Standart login/parol arkaly gurluşa rugsatsyz giriş mümkinçiligi.",
        "telnet_weak_auth plugin arkaly barlag.",
    ),
    (
        "Açyk web-interfeýs",
        "80/HTTP, 443/HTTPS, 8080",
        "Web-paneliň goragsyz bolmagy ýa-da köne firmware sebäpli rugsatsyz dolandyryş howpy.",
        "Açyk porty kesgitlemek we web-interfeýs boýunça maslahat bermek.",
    ),
    (
        "Kamera wideo akymyna giriş",
        "554/RTSP",
        "IP kameranyň wideo akymynyň rugsatsyz görülmegi.",
        "RTSP portunyň açykdygyny görkezmek.",
    ),
    (
        "IoT habarlarynyň goragsyz alyş-çalyşygy",
        "1883/MQTT, 8883/MQTT SSL",
        "Datçik maglumatlarynyň okalmagy ýa-da galp buýruklaryň iberilmegi.",
        "MQTT portlaryny skanirlemek.",
    ),
    (
        "Faýl geçirmek hyzmatynyň goragsyzlygy",
        "21/FTP",
        "Login maglumatlarynyň şifrlenmedik görnüşde geçirilmegi.",
        "FTP portunyň açykdygyny görkezmek.",
    ),
    (
        "Uzakdan dolandyryş kanaly",
        "22/SSH",
        "Gowşak parol ýa-da nädogry sazlama arkaly rugsatsyz giriş.",
        "SSH portuny görkezmek we administratory barlaga ugrukdyrmak.",
    ),
]


CONCLUSION = (
    "Bu bölümde seljerilen hüjüm wektorlary IoT Security Scanner programmasynyň näme üçin "
    "açyk portlary, hyzmatlary we gowşak autentifikasiýa alamatlaryny barlamalydygyny "
    "esaslandyrýar. Şeýlelikde, skanirleme netijeleri diňe tehniki maglumat hökmünde däl, "
    "eýsem howpsuzlyk kararlaryny kabul etmek üçin amaly maglumat çeşmesi hökmünde çykyş edýär."
)


def set_document_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True


def add_heading(document):
    heading = document.add_heading(TITLE, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_table(document):
    document.add_paragraph()
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for index, text in enumerate(TABLE_ROWS[0]):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    for row in TABLE_ROWS[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("Tablisa 1.3 - IoT torlarynda hüjüm wektorlary we programma tarapyndan barlag usullary")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True


def build_document():
    document = Document()
    set_document_styles(document)
    add_heading(document)

    for text in PARAGRAPHS:
        add_paragraph(document, text)

    add_table(document)
    add_paragraph(document, CONCLUSION)

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_FILE)
