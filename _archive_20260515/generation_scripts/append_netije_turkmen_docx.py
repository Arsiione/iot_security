from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_netije.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def build_netije_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 1"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "Netije")

    paragraphs = [
        (
            "Bu diplom taslamasynda ýerine ýetirilen “IoT gurluşlarynyň howpsuzlygyny barlamak üçin "
            "awtomatlaşdyrylan skanirleme programma toplumyny işläp taýýarlamak” atly tema häzirki zaman "
            "lokal torlarynda maglumat howpsuzlygyny ýokarlandyrmak üçin örän möhüm ähmiýete eýedir. "
            "Internet of Things tehnologiýalarynyň giňden ýaýramagy bilen öý, okuw we kärhana torlarynda "
            "IP kameralar, routerler, telefonlar, printerler, akylly öý enjamlary we beýleki tor gurluşlary "
            "köpelýär. Şeýle gurluşlaryň goragsyz ýagdaýda işlemegi bolsa açyk portlar, standart parollar, "
            "köne firmware we rugsatsyz giriş ýaly howplaryň döremegine sebäp bolup biler."
        ),
        (
            "Diplom taslamasynda başlangyç ýumuşa laýyklykda düşündirişli ýazgy, ýagny giriş, predmet "
            "oblastynyň seljermesi, programma arhitekturasynyň taslanylyşy, programma toplumynyň işlenip "
            "taýýarlanylyşy, ulanyjy interfeýsiniň beýan edilişi, synag we netijelilik bahalandyrmasy, "
            "ykdysady bölüm, zähmeti we daşky gurşawy goramak bölümleri ýerine ýetirildi. Şeýle hem "
            "programma toplumynyň işini görkezmek üçin IoT Security Scanner atly amaly programma döredildi."
        ),
        (
            "Düşündirişli ýazgynyň giriş bölüminde häzirki wagtda IoT tehnologiýalarynyň ähmiýeti, lokal "
            "torlarda howpsuzlyk meselesiniň wajyplygy, taslamanyň maksady we ýerine ýetirilmeli wezipeleri "
            "beýan edildi. Bu bölümde işiň näme üçin zerurdygy we işlenip taýýarlanylýan programma çözgüdiniň "
            "haýsy meseläni çözýändigi esaslandyryldy."
        ),
        (
            "Taslamanyň nazary bölüminde IoT gurluşlarynyň görnüşleri we aýratynlyklary öwrenildi. IP "
            "kameralar, akylly öý enjamlary, routerler, senagat IoT gurluşlary we lukmançylyk IoT enjamlary "
            "boýunça aýratyn düşündirişler berildi. Şeýle hem IoT torlaryna edilýän hüjüm wektorlary, şol "
            "sanda Telnet, HTTP/HTTPS, RTSP, MQTT, FTP we SSH hyzmatlary bilen bagly töwekgelçilikler "
            "seljerildi. Nmap, Shodan, Nessus we OpenVAS ýaly bar bolan gurallar bilen deňeşdirme geçirilip, "
            "IoT Security Scanner programmasynyň lokal tor üçin ýöriteleşdirilen, türkmen dilindäki we "
            "maslahat beriş mümkinçilikli çözgüt hökmünde tapawutlanýandygy görkezildi."
        ),
        (
            "Taslamanyň arhitektura bölüminde programma toplumyny döretmek üçin saýlanan tehnologiýalar "
            "esaslandyryldy. Python dili programma logikasyny ýazmak üçin, PyQt6 grafiki interfeýs döretmek "
            "üçin, SQLite maglumatlary saklamak üçin, Nmap tor skanirlemesi üçin, ReportLab PDF hasabat "
            "taýýarlamak üçin, Paramiko bolsa remediation mümkinçiliklerini giňeltmek üçin saýlanyp alyndy. "
            "Programma Core, UI, Plugins, Database, Reports we Utils ýaly modullara bölünip taslanyldy."
        ),
        (
            "Programma toplumynyň işlenip taýýarlanylyş bölüminde skanirleme ýadrosynyň, plugin ulgamynyň, "
            "töwekgelçiligi bahalandyrmak we maslahat beriş logikasynyň, maglumatlar bazasynyň, taryh "
            "ulgamynyň we PDF hasabat döretmek mümkinçilikleriniň durmuşa geçirilişi görkezildi. Programma "
            "Wi‑Fi ýa-da Ethernet adapterini saýlap, dogry lokal IP aralygyny kesgitleýär, Nmap, ARP, Ping "
            "we Windows ARP keşi arkaly gurluşlary tapýar, olaryň portlaryny barlaýar we netijeleri tablisa "
            "görnüşinde görkezýär."
        ),
        (
            "Ulanyjy interfeýsi bölüminde programma sahypalarynyň işleýşi beýan edildi. Skanirleme sahypasynda "
            "adapter saýlamak, IP aralygyny görmek, skanirlemäni başlatmak, netijeleri real wagtda almak we "
            "PDF hasabat döretmek mümkinçilikleri bar. Panel sahypasy umumy statistikany görkezýär, Netijeler "
            "sahypasy soňky skanirlemäni jikme-jik seljerýär, Taryh sahypasy bolsa öňki skanirlemeleri "
            "maglumatlar bazasyndan görkezýär. Interfeýsiň türkmen dilinde taýýarlanmagy diplom goragynda "
            "programmany düşündirmegi has aňsatlaşdyrýar."
        ),
        (
            "Synag bölüminde programma toplumynyň esasy mümkinçilikleri barlandy. Programma EXE görnüşinde "
            "gara konsol penjiresi açylmazdan işledildi, Wi‑Fi adapteriniň dogry saýlanmagy, lokal IP "
            "aralygynyň awtomatik kesgitlenmegi, gurluşlaryň tapylmagy, portlaryň görkezilmegi, netijeleriň "
            "SQLite bazasynda saklanmagy we PDF hasabatynyň döredilmegi synagdan geçirildi. Synaglaryň "
            "netijesinde programma lokal IoT torlarynyň başlangyç howpsuzlyk auditi üçin niýetlenen esasy "
            "wezipeleri ýerine ýetirýändigi anyklanyldy."
        ),
        (
            "Ykdysady bölümde programma toplumynyň işlenip taýýarlanylmagynyň we ulanylmagynyň ykdysady "
            "tarapdan maksadalaýykdygy hasaplanyldy. Programma el bilen geçirilýän tor auditine sarp edilýän "
            "wagty azaltmaga, tölegli skaner lisenziýalaryna baglylygy peseltmäge we PDF görnüşli elektron "
            "hasabatlary döretmek arkaly iş prosesini ýeňilleşdirmäge mümkinçilik berýär."
        ),
        (
            "Zähmeti we daşky gurşawy goramak bölüminde programma üpjünçiligini işläp taýýarlamak we ulanmak "
            "bilen baglanyşykly iş ýeriniň howpsuzlygy, elektrik howpsuzlygy, iş otagynyň mikroklimaty, "
            "ýagtylandyrylyşy, energiýa sarp edilişi we ekologik täsirleri boýunça degişli maglumatlar we "
            "hasaplamalar getirildi. Bu bölüm programma taslamasynyň diňe tehniki däl, eýsem iş gurşawy we "
            "ekologik nukdaýnazardan hem esaslandyrylmagyna hyzmat edýär."
        ),
        (
            "Diplom taslamasynyň netijesinde döredilen IoT Security Scanner programma toplumy lokal torda "
            "ýerleşýän gurluşlary tapmaga, olaryň IP we MAC maglumatlaryny, öndürijisini, görnüşini, tapylan "
            "usulyny, açyk portlaryny, töwekgelçilik derejesini we howpsuzlyk maslahatlaryny görkezmäge "
            "mümkinçilik berýär. Şeýle hem programma skanirleme taryhyny saklaýar we netijeleri PDF hasabat "
            "görnüşinde eksport edýär."
        ),
        (
            "Diplom taslamasynda ýerine ýetirilen programma çözgüdi okuw edaralarynda, kiçi kärhanalarda we "
            "öý lokal torlarynda başlangyç IoT howpsuzlyk auditini geçirmek üçin peýdaly bolup biler. Bu "
            "programma administratora torda haýsy gurluşlaryň bardygyny görmek, açyk hyzmatlara üns bermek "
            "we gowşaklyklary azaltmak boýunça ilkinji çäreleri kesgitlemekde ýardam berer."
        ),
        (
            "Şeýlelikde, diplom taslamasynda goýlan maksat ýerine ýetirildi. IoT gurluşlarynyň howpsuzlygyny "
            "barlamak üçin awtomatlaşdyrylan, grafiki interfeýsli, maglumatlar bazasy we hasabat ulgamy bolan "
            "programma toplumy işlenip taýýarlanyldy. Taslamanyň netijeleri häzirki zaman lokal torlarynda "
            "IoT howpsuzlygyny ýokarlandyrmak ugrunda amaly ähmiýete eýedir."
        ),
    ]

    for text in paragraphs:
        add_normal(doc, text)

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))

    # If a global Netije already exists before ecology, avoid duplicating it.
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Netije" and paragraph.style.name.startswith("Heading"):
            print("Netije already exists; no changes made.")
            return
        if text.startswith("Z") and "ekologik" in text:
            break

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_netije_doc()
    insert_before_ecology(main_doc, section_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
