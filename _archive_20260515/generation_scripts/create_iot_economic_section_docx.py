from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_FILE = "iot_security_ykdysady_bolum.docx"


TITLE = "Taslamanyň ykdysady netijeliliginiň hasaplamasy"


LABOR_TABLE = [
    ("T№", "Iş görnüşi / işgär kategoriýasy", "Sany", "Wagt, sag.", "1 sag. töleg, manat", "Tarif zähmet haky, manat", "Baýrak 10%, manat"),
    ("1", "Programmist", "1", "160", "20", "3200,00", "320,00"),
    ("2", "UI/UX dizaýner", "1", "40", "18", "720,00", "72,00"),
    ("3", "Tor howpsuzlygy boýunça testçi", "1", "50", "17", "850,00", "85,00"),
    ("", "Jemi", "-", "250", "-", "4770,00", "477,00"),
]


EQUIPMENT_TABLE = [
    ("T№", "Enjamyň ady", "Bahasy, manat", "Amortizasiýa normasy, %/ýyl", "Ulanylan möhlet, aý", "Amortizasiýa, manat"),
    ("1", "Kompýuter ýa-da noutbuk", "3000", "20", "2", "100,00"),
    ("2", "Router ýa-da test switch", "350", "20", "2", "11,67"),
    ("", "Jemi", "3350", "-", "-", "111,67"),
]


COST_TABLE = [
    ("T№", "Harajat maddalary", "Gymmaty, manat"),
    ("1", "Materiallar we resminama taýýarlamak", "80,00"),
    ("2", "Internet we aragatnaşyk harajatlary", "150,00"),
    ("3", "Enjamlaryň amortizasiýasy", "111,67"),
    ("4", "Elektrik energiýasy", "5,75"),
    ("5", "Esasy zähmet haky", "5247,00"),
    ("6", "Goşmaça zähmet haky", "524,70"),
    ("7", "Sosial gorag üçin geçirimler", "1154,34"),
    ("8", "Bölüm boýunça çykdajylar", "787,05"),
    ("9", "Kärhana boýunça çykdajylar", "1311,75"),
    ("10", "Enjamlary ulanmak we hyzmat etmek çykdajylary", "524,70"),
    ("", "Jemi taslama gymmaty", "9896,96"),
]


EFFECT_TABLE = [
    ("Görkeziji", "El bilen audit", "Programma arkaly audit", "Netije"),
    ("Bir lokal tory barlamagyň wagty", "6 sagat", "1 sagat", "5 sagat tygşytlanýar"),
    ("Bir sagatlyk hünärmen tölegi", "20 manat", "20 manat", "-"),
    ("Bir audit boýunça tygşyt", "120 manat", "20 manat", "100 manat"),
    ("Ýyllyk audit sany", "48", "48", "-"),
    ("Ýyllyk zähmet tygşydy", "-", "-", "4800 manat"),
    ("Tölegli skaner lisenziýasyny ulanmazlyk", "-", "-", "3000 manat/ýyl"),
    ("Umumy ýyllyk ykdysady netije", "-", "-", "7800 manat/ýyl"),
]


INTRO_PARAGRAPHS = [
    (
        "Bu bölümde IoT Security Scanner programma toplumyny işläp taýýarlamagyň ykdysady "
        "tarapy seljerilýär. Hasaplamanyň maksady programma üpjünçiligini döretmek üçin "
        "zerur bolan harajatlary kesgitlemek, iş wagtynyň tygşytlanyşyny bahalandyrmak "
        "we taslamanyň özüni ödeýiş möhletini hasaplamakdan ybaratdyr."
    ),
    (
        "IoT Security Scanner lokal tordaky gurluşlary awtomatik tapmak, açyk portlary "
        "barlamak, potensial gowşaklyklary ýüze çykarmak, netijeleri maglumat bazasynda "
        "saklamak we PDF görnüşinde hasabat taýýarlamak üçin niýetlenendir. Şeýle ulgam "
        "tor administratorynyň el bilen ýerine ýetirýän işini azaldýar we howpsuzlyk "
        "auditiniň tizligini ýokarlandyrýar."
    ),
    (
        "Hasaplamalarda ähli bahalar şertli görnüşde manat hasabynda kabul edilýär. "
        "Bu sanlar taslamanyň ykdysady netijeliligini görkezmek üçin ulanylýar we "
        "hakyky kärhana şertlerinde bazar nyrhlaryna görä üýtgäp biler."
    ),
]


RECOMMENDATIONS = [
    "Programmany açyk çeşmeli tehnologiýalar bilen döretmek lisenziýa çykdajylaryny azaldýar.",
    "PDF hasabatlarynyň elektron görnüşde saklanmagy kagyz we çap çykdajylaryny peseldýär.",
    "Köp akymly skanirleme administrator wagtyny tygşytlaýar.",
    "Programmany birnäçe lokal tor segmentinde ulanmak ýyllyk ykdysady netijäni ýokarlandyrýar.",
    "Nmap, SQLite, Python we PyQt6 ýaly gurallar taslamanyň başlangyç çykdajysyny peseldýär.",
]


def set_styles(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def add_heading(document, text, level=1, center=False):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
    return heading


def add_paragraph(document, text, first_line=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return paragraph


def add_formula(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    return paragraph


def add_table(document, rows, caption):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for index, text in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(6)
    run = caption_p.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    return table


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("• " + text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def build_document():
    document = Document()
    set_styles(document)

    add_heading(document, TITLE, level=1, center=True)

    for text in INTRO_PARAGRAPHS:
        add_paragraph(document, text)

    add_heading(document, "1. Taslamany işläp taýýarlamagyň zähmet talap edişi", level=2)
    add_paragraph(
        document,
        "Programma toplumyny döretmek üçin esasy işler talaplary seljermekden, programma "
        "ýadrosyny ýazmakdan, grafiki interfeýsi taýýarlamakdan, plugin ulgamyny gurmakdan, "
        "maglumat bazasy bilen işlemekden we synag geçirmekden ybaratdyr. Umumy zähmet "
        "talap edişi 250 adam-sagat diýip kabul edilýär.",
    )
    add_table(document, LABOR_TABLE, "Tablisa 1 - Taslama boýunça zähmet talap edişi we tarif zähmet haky")
    add_formula(document, "Tjemi = 160 + 40 + 50 = 250 adam-sagat")
    add_formula(document, "Ztar = 160 · 20 + 40 · 18 + 50 · 17 = 4770 manat")
    add_formula(document, "Zbaýr = Ztar · 10% = 4770 · 0,10 = 477 manat")
    add_formula(document, "Zes = Ztar + Zbaýr = 4770 + 477 = 5247 manat")
    add_paragraph(
        document,
        "Hasaplamanyň netijesinde taslamany işläp taýýarlamak üçin tarif boýunça zähmet haky "
        "4770 manat, baýrak gaznasy 477 manat we esasy zähmet haky 5247 manat bolýar."
    )

    add_heading(document, "2. Goşmaça zähmet haky we sosial geçirimler", level=2)
    add_paragraph(
        document,
        "Goşmaça zähmet haky esasy zähmet hakynyň 10%-i möçberinde kabul edilýär. Sosial "
        "gorag üçin geçirimler umumy zähmet hakynyň 20%-i möçberinde hasaplanýar."
    )
    add_formula(document, "Zgoş = Zes · 10% = 5247 · 0,10 = 524,70 manat")
    add_formula(document, "Zum = Zes + Zgoş = 5247 + 524,70 = 5771,70 manat")
    add_formula(document, "Ssos = Zum · 20% = 5771,70 · 0,20 = 1154,34 manat")
    add_paragraph(
        document,
        "Şeýlelikde, taslama boýunça umumy zähmet haky 5771,70 manat, sosial gorag üçin "
        "geçirimler bolsa 1154,34 manat bolýar."
    )

    add_heading(document, "3. Enjamlaryň amortizasiýasy we energiýa harajatlary", level=2)
    add_paragraph(
        document,
        "Programmany işläp taýýarlamak üçin kompýuter ýa-da noutbuk, monitor, router ýa-da "
        "test switch ýaly enjamlardan peýdalanylýar. Enjamlaryň ýyllyk amortizasiýa normasy "
        "20% diýip kabul edilýär, taslama boýunça ulanylan möhlet 2 aý hasaplanýar."
    )
    add_table(document, EQUIPMENT_TABLE, "Tablisa 2 - Enjamlaryň amortizasiýa hasaplamasy")
    add_formula(document, "A = C · Ha · m / 12")
    add_formula(document, "Akomp = 3000 · 0,20 · 2 / 12 = 100 manat")
    add_formula(document, "Arout = 350 · 0,20 · 2 / 12 = 11,67 manat")
    add_formula(document, "Ajemi = 100 + 11,67 = 111,67 manat")
    add_paragraph(
        document,
        "Elektrik energiýasy boýunça hasaplamada kompýuter, monitor we tor enjamlarynyň ortaça "
        "kuwwaty 0,23 kW, iş wagty 250 sagat we 1 kW·sag energiýanyň bahasy 0,10 manat "
        "diýip kabul edilýär."
    )
    add_formula(document, "E = P · t = 0,23 · 250 = 57,5 kW·sag")
    add_formula(document, "Çel = E · B = 57,5 · 0,10 = 5,75 manat")

    add_heading(document, "4. Taslamanyň umumy harajatlarynyň hasaplamasy", level=2)
    add_paragraph(
        document,
        "Taslamanyň umumy gymmaty zähmet haky, sosial geçirimler, enjamlaryň amortizasiýasy, "
        "internet we aragatnaşyk harajatlary, elektrik energiýasy hem-de kärhana boýunça "
        "goşmaça çykdajylar boýunça jemlenýär. Bölüm boýunça çykdajylar esasy zähmet hakynyň "
        "15%-i, kärhana boýunça çykdajylar 25%-i, enjamlary ulanmak we hyzmat etmek çykdajylary "
        "bolsa 10%-i möçberinde kabul edilýär."
    )
    add_formula(document, "Çbölüm = Zes · 15% = 5247 · 0,15 = 787,05 manat")
    add_formula(document, "Çkärhana = Zes · 25% = 5247 · 0,25 = 1311,75 manat")
    add_formula(document, "Çenjam = Zes · 10% = 5247 · 0,10 = 524,70 manat")
    add_table(document, COST_TABLE, "Tablisa 3 - IoT Security Scanner taslamasynyň umumy harajatlary")
    add_paragraph(
        document,
        "Hasaplamalaryň netijesinde IoT Security Scanner programma toplumyny işläp taýýarlamagyň "
        "umumy gymmaty 9896,96 manat boldy."
    )

    add_heading(document, "5. Ykdysady netijeliligiň hasaplamasy", level=2)
    add_paragraph(
        document,
        "Programmany ulanmagyň ykdysady netijesi tor howpsuzlyk auditini geçirmek üçin sarp "
        "edilýän wagtyň azalmagy bilen baglanyşyklydyr. El bilen barlag geçirilende lokal toruň "
        "gurluşlaryny tapmak, portlary barlamak we hasabat taýýarlamak ortaça 6 sagat wagt "
        "talap edýär. IoT Security Scanner ulanylanda bu iş ortaça 1 sagatda ýerine ýetirilýär."
    )
    add_table(document, EFFECT_TABLE, "Tablisa 4 - Programma ulanmagyň ykdysady netijesi")
    add_formula(document, "Eaudit = (Tel - Tprog) · Bsag = (6 - 1) · 20 = 100 manat")
    add_formula(document, "Ezähmet = Eaudit · Nýyl = 100 · 48 = 4800 manat/ýyl")
    add_formula(document, "Eýyl = Ezähmet + Elisenziýa = 4800 + 3000 = 7800 manat/ýyl")
    add_paragraph(
        document,
        "Bu ýerde 48 sany ýyllyk audit kärhananyň dürli lokal tor segmentleri boýunça geçirilýän "
        "yzygiderli barlaglary aňladýar. Mundan başga-da, açyk çeşmeli çözgüt ulanylandygy üçin "
        "tölegli skaner lisenziýasyna harajat edilmeýär. Bu tygşyt 3000 manat/ýyl diýip kabul edilýär."
    )

    add_heading(document, "6. Taslamanyň özüni ödeýiş möhleti", level=2)
    add_paragraph(
        document,
        "Taslamanyň özüni ödeýiş möhleti umumy taslama gymmatynyň ýyllyk ykdysady netijä bolan "
        "gatnaşygy arkaly hasaplanýar."
    )
    add_formula(document, "Töz.öd. = K / Eýyl")
    add_formula(document, "Töz.öd. = 9896,96 / 7800 = 1,27 ýyl")
    add_paragraph(
        document,
        "Hasaplama boýunça programma toplumynyň özüni ödeýiş möhleti 1,27 ýyl bolýar. Bu görkeziji "
        "programma çözgüdiniň ykdysady taýdan maksadalaýykdygyny görkezýär. Eger programma has köp "
        "tor segmentinde ulanylsa, özüni ödeýiş möhleti has hem gysgalýar."
    )

    add_heading(document, "7. Ykdysady netijeliligi ýokarlandyrmagyň ýollary", level=2)
    for recommendation in RECOMMENDATIONS:
        add_bullet(document, recommendation)

    add_heading(document, "Netije", level=2)
    add_paragraph(
        document,
        "Geçirilen ykdysady hasaplamalar IoT Security Scanner programma toplumynyň işlenip "
        "taýýarlanylmagynyň we ulanylmagynyň ykdysady taýdan esaslydygyny görkezýär. Taslamanyň "
        "umumy gymmaty 9896,96 manat, ýyllyk ykdysady netijesi 7800 manat we özüni ödeýiş möhleti "
        "1,27 ýyl boldy. Programma tor administratorynyň iş wagtyny tygşytlaýar, tölegli skaner "
        "lisenziýalaryna bolan zerurlygy azaldýar we howpsuzlyk auditiniň netijeliligini ýokarlandyrýar."
    )

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_FILE)
