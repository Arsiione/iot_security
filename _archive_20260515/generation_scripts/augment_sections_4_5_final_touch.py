from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = ROOT / "kakabalowa.docx"


EXTRA = {
    "4.8. Ulanyjy tejribesini ýokarlandyrmak üçin kabul edilen çözgütler": [
        "Ulanyjy tejribesiniň hilini ýokarlandyrýan ýene bir çözgüt maglumatlaryň yzygiderli terminologiýa bilen berilmegidir. Bir sahypada `töwekgelçilik`, beýleki sahypada bolsa başga manydaky söz ulanylmaýar; risk, port, maslahat, tapylan usul ýaly düşünjeler ähli interfeýsde birmeňzeş saklanýar. Bu ýagdaý ulanyjynyň her sahypany täzeden öwrenmeginiň öňüni alýar we programma bilen işleşmegi çaltlaşdyrýar.",
        "Interfeýsde geljekde giňeldiljek mümkinçilikler üçin hem ýer goýuldy. Sazlamalar sahypasy, eksport düwmeleri we netijeleri aýratyn sahypalarda görkezmek modeli soňra täze filtrleri, hasabat görnüşlerini ýa-da goşmaça howpsuzlyk barlaglaryny goşmaga mümkinçilik berýär. Şeýlelikde, häzirki interfeýs diňe diplom goragy üçin däl, programma toplumynyň geljekki ösüşi üçin hem esas bolup hyzmat edýär.",
    ],
    "5.8. Synag netijeleriniň diplom goragynda görkezilişi": [
        "Diplom goragynda netijeleri görkezmek üçin skrinşotlar we öňünden taýýarlanan PDF hasabat hem ulanylyp bilner. Bu usul torda häzirki wagtda IoT enjam ýok bolsa-da, programmanyň nähili netije berýändigini düşündirmäge kömek edýär. Şeýle ýagdaýda janly skanirleme bilen bilelikde öňki skanirleme taryhy we hasabat faýly görkezilse, taslamanyň funksional mümkinçilikleri doly açylýar.",
        "Synaglaryň jemleýji netijesi programma toplumynyň goýlan maksada laýyk işleýändigini görkezýär. Programma tor segmentini saýlaýar, gurluşlary tapýar, portlary barlaýar, töwekgelçiligi kesgitleýär, netijeleri bazada saklaýar we ulanyja hasabat görnüşinde berýär. Bu yzygiderlilik diplom işinde nazary seljerme bilen amaly programma çözgüdiniň özara baglanyşygyny subut edýär.",
    ],
}


def insert_before(doc: Document, element, text: str) -> None:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def next_heading_element(doc: Document, heading: str):
    start = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name.startswith("Heading") and " ".join(paragraph.text.split()) == heading:
            start = i
            break
    if start is None:
        raise ValueError(f"Heading not found: {heading}")
    for i in range(start + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith("Heading"):
            return doc.paragraphs[i]._p
    raise ValueError(f"Next heading not found after {heading}")


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC_PATH.with_name(f"kakabalowa_before_expand_sections_4_5_final_{timestamp}.docx")
    shutil.copy2(DOC_PATH, backup)
    doc = Document(str(DOC_PATH))
    for heading, paragraphs in EXTRA.items():
        anchor = next_heading_element(doc, heading)
        for text in paragraphs:
            insert_before(doc, anchor, text)
    doc.save(str(DOC_PATH))
    print(f"Backup: {backup}")
    print(f"Updated: {DOC_PATH}")


if __name__ == "__main__":
    main()
