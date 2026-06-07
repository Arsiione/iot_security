from html import unescape
from pathlib import Path
import re

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image


DOCX_FILE = Path("bolum_1_1_iot_gurlushlary.docx")
BACKUP_FILE = Path("bolum_1_1_iot_gurlushlary_before_real_images.docx")
ASSET_DIR = Path("diploma_assets/iot_real_devices")

OLD_SECTION_TITLE = "IoT gurluşlarynyň mysal şekilleri"
OLD_CAPTION = "Surat 1.1 - IoT gurluşlarynyň görnüşleri boýunça şekiller"
NEW_SECTION_TITLE = "IoT gurluşlarynyň görnüşleri we amaly mysallar"

COMMONS_API = "https://commons.wikimedia.org/w/api.php"


DEVICES = [
    {
        "title": "IP kameralar",
        "file": "File:Outdoor wireless security IP camera at Nuthurst, Sussex 2.jpg",
        "filename": "ip_camera.png",
        "paragraph": (
            "IP kamera IoT torlarynda iň köp duş gelýän gurluşlaryň biridir. Şeýle enjamlar "
            "wideo gözegçilik üçin ulanylýar we köplenç HTTP web-interfeýsi, RTSP wideo akymy "
            "we käbir ýagdaýlarda Telnet ýa-da SSH ýaly dolandyryş hyzmatlary bilen işleýär. "
            "Eger kameranyň standart paroly üýtgedilmese ýa-da firmware wagtynda täzelenmese, "
            "hüjümçi kamera arkaly wideo akymyna ýa-da dolandyryş paneline rugsatsyz girip biler. "
            "IoT Security Scanner programmasy şeýle gurluşlarda 80, 443, 554, 8000 we 8080 ýaly "
            "portlaryň açykdygyny görkezmek arkaly administratora howpsuzlyk ýagdaýyny seljermäge "
            "kömek edýär."
        ),
    },
    {
        "title": "Akylly öý enjamlary",
        "file": "File:Smart-plug.jpg",
        "filename": "smart_plug.png",
        "paragraph": (
            "Akylly öý enjamlaryna akylly rozetkalar, lampalar, termostatlar, datçikler we "
            "dolandyryş modullary degişlidir. Bu gurluşlar adatça mobil programma, bulut hyzmaty "
            "ýa-da lokal web-interfeýs arkaly dolandyrylýar. Olaryň esasy howpsuzlyk meseleleri "
            "gowşak awtorizasiýa, goragsyz API, köne programma üpjünçiligi we nädogry sazlanan "
            "tor hyzmatlary bilen baglanyşyklydyr. Şeýle enjamlaryň lokal torda görünmegi we "
            "açyk portlarynyň kesgitlenmegi administratora olaryň howpsuzlygyny wagtynda barlamaga "
            "mümkinçilik berýär."
        ),
    },
    {
        "title": "Tor enjamlary",
        "file": "File:Cisco Linksys WRT54G2 V1.jpg",
        "filename": "router.png",
        "paragraph": (
            "Tor enjamlaryna marşrutizatorlar, Wi-Fi nokatlary, switchler we repeaterler degişlidir. "
            "Bu gurluşlar lokal toruň merkezi bölegi bolup durýar, sebäbi beýleki IoT enjamlarynyň "
            "köpüsi şolaryň üsti bilen tora birikýär. Routerde açyk Telnet, SSH ýa-da web-panel "
            "bar bolsa, onuň goragynyň pesligi tutuş toruň howpsuzlygyna täsir edip biler. "
            "Şonuň üçin skanirleme wagtynda 22, 23, 80, 443 we 8080 portlarynyň ýagdaýyny görmek "
            "möhüm ähmiýete eýedir."
        ),
    },
    {
        "title": "Senagat IoT gurluşlary",
        "file": "File:Control-panel-plc.jpg",
        "filename": "industrial_iot.png",
        "paragraph": (
            "Senagat IoT gurluşlaryna programmalaşdyrylýan logiki kontrollerler, önümçilik "
            "datçikleri, ölçeg modullary we awtomatlaşdyryş panelleri girýär. Şeýle gurluşlar "
            "önümçilik proseslerini dolandyrmakda ulanylýandygy sebäpli olaryň howpsuzlygy aýratyn "
            "wajypdyr. Açyk dolandyryş portlary ýa-da goragsyz web-interfeýsler önümçilik prosesine "
            "rugsatsyz täsir etmek howpuny döredip biler. IoT Security Scanner başlangyç auditde "
            "şeýle gurluşlaryň tor hyzmatlaryny ýüze çykarmaga mümkinçilik berýär."
        ),
    },
    {
        "title": "Lukmançylyk IoT enjamlary",
        "file": "File:Yokota pulse oximeter.jpg",
        "filename": "medical_iot.png",
        "paragraph": (
            "Lukmançylyk IoT enjamlaryna saglyk ýagdaýyny ölçeýän datçikler, pulse oksimetrler, "
            "monitorlar we uzakdan gözegçilik ulgamlary degişlidir. Bu gurluşlar şahsy we saglyk "
            "maglumatlary bilen işleýändigi üçin olaryň goragy aýratyn ähmiýete eýedir. "
            "Nädogry sazlanan tor hyzmatlary maglumatlaryň syzmagyna ýa-da nädogry maglumatlaryň "
            "ulgama girizilmegine sebäp bolup biler. Şonuň üçin şeýle enjamlaryň açyk portlaryny "
            "we hyzmatlaryny gözegçilikde saklamak zerurdyr."
        ),
    },
]


def clean_html(value):
    value = re.sub(r"<[^>]+>", "", value or "")
    return unescape(value).strip()


def get_image_info(file_title):
    params = {
        "action": "query",
        "format": "json",
        "prop": "imageinfo",
        "titles": file_title,
        "iiprop": "url|extmetadata",
        "iiurlwidth": 1000,
    }
    response = requests.get(
        COMMONS_API,
        params=params,
        timeout=30,
        headers={"User-Agent": "IoTSecurityDiploma/1.0"},
    )
    response.raise_for_status()
    page = next(iter(response.json()["query"]["pages"].values()))
    if "missing" in page:
        raise RuntimeError(f"Commons file tapylmady: {file_title}")

    info = page["imageinfo"][0]
    metadata = info.get("extmetadata", {})
    return {
        "url": info.get("thumburl") or info["url"],
        "page_url": f"https://commons.wikimedia.org/wiki/{file_title.replace(' ', '_')}",
        "license": clean_html(metadata.get("LicenseShortName", {}).get("value", "")),
        "artist": clean_html(metadata.get("Artist", {}).get("value", "")),
    }


def download_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for device in DEVICES:
        info = get_image_info(device["file"])
        image_path = ASSET_DIR / device["filename"]
        response = requests.get(info["url"], timeout=60, headers={"User-Agent": "IoTSecurityDiploma/1.0"})
        response.raise_for_status()
        temp_path = image_path.with_suffix(".download")
        temp_path.write_bytes(response.content)
        with Image.open(temp_path) as image:
            image.convert("RGB").save(image_path, format="PNG")
        temp_path.unlink(missing_ok=True)
        device["path"] = image_path
        device["source"] = info


def paragraph_text(element):
    return "".join(element.itertext()).strip()


def remove_section(document, start_titles, end_captions):
    body = document.element.body
    children = list(body)
    start = None
    end = None

    for index, element in enumerate(children):
        text = paragraph_text(element)
        if start is None and text in start_titles:
            start = index
        if start is not None and text in end_captions:
            end = index
            break

    if start is not None:
        if end is None:
            end = start
        for element in children[start : end + 1]:
            body.remove(element)


def remove_old_image_table(document):
    for table in list(document.tables):
        first_cell = table.cell(0, 0).text.strip() if table.rows and table.columns else ""
        if first_cell == "IoT gurluşy":
            table._element.getparent().remove(table._element)


def set_run_style(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_text(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_run_style(run, size=14)
    return paragraph._p


def add_centered_text(document, text, size=12, italic=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_style(run, size=size, italic=italic)
    return paragraph._p


def add_device_blocks(document):
    elements = []

    heading = document.add_heading(NEW_SECTION_TITLE, level=3)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        set_run_style(run, size=14, bold=True)
    elements.append(heading._p)

    elements.append(add_text(
        document,
        "Aşakdaky bölümlerde IoT gurluşlarynyň esasy görnüşleri aýratynlykda görkezilýär. "
        "Her mysalda gurluşyň lokal tordaky orny, bolup biljek howpsuzlyk töwekgelçiligi "
        "we IoT Security Scanner programmasy bilen nähili baglanyşýandygy düşündirilýär."
    ))

    for number, device in enumerate(DEVICES, start=1):
        title = document.add_heading(f"1.1.{number}. {device['title']}", level=3)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            set_run_style(run, size=14, bold=True)
        elements.append(title._p)

        elements.append(add_text(document, device["paragraph"]))

        image_p = document.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_run = image_p.add_run()
        image_run.add_picture(str(device["path"]), width=Cm(11))
        elements.append(image_p._p)

        source = device["source"]
        caption = (
            f"Surat 1.{number} - {device['title']}. "
            f"Çeşme: Wikimedia Commons, {source['artist'] or 'awtor görkezilmedik'}, "
            f"{source['license'] or 'lisenziýa görkezilmedik'}."
        )
        elements.append(add_centered_text(document, caption, size=12, italic=True))

    return elements


def find_anchor(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Tablisa 1.1"):
            return paragraph._p
    raise RuntimeError("Tablisa 1.1 tapylmady")


def insert_after(anchor, elements):
    for element in reversed(elements):
        anchor.addnext(element)


def update_docx():
    download_images()

    if not BACKUP_FILE.exists():
        BACKUP_FILE.write_bytes(DOCX_FILE.read_bytes())

    document = Document(DOCX_FILE)
    remove_section(document, {OLD_SECTION_TITLE, NEW_SECTION_TITLE}, {OLD_CAPTION})
    remove_old_image_table(document)

    # If a previous version of the new block exists, remove it up to the last known picture caption.
    remove_section(
        document,
        {NEW_SECTION_TITLE},
        {"Surat 1.5 - Lukmançylyk IoT enjamlary. Çeşme:"},
    )

    anchor = find_anchor(document)
    elements = add_device_blocks(document)
    insert_after(anchor, elements)
    document.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == "__main__":
    update_docx()
