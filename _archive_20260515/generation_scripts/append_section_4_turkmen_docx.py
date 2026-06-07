from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_section_4.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text, level=2):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def add_table(doc, caption_text, headers, rows):
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "4. Ulanyjy interfeýsiniň işlenip taýýarlanylyşy", level=1)
    add_normal(
        doc,
        "IoT Security Scanner programma toplumynyň netijeli ulanylmagy diňe skanirleme algoritmlerine däl, "
        "eýsem ulanyjy interfeýsiniň düşnükliligine hem baglydyr. Tor howpsuzlygy bilen işleýän ulanyjy "
        "tapylan IP salgylary, portlary, töwekgelçilik derejesini we maslahatlary çalt görüp bilmeli. "
        "Şonuň üçin programma PyQt6 esasynda grafiki desktop interfeýs görnüşinde işlenip taýýarlanyldy. "
        "Interfeýs lokal tor auditi üçin zerur bolan esasy amallary bir penjirede jemleýär."
    )

    add_heading(doc, "4.1. Interfeýsiň umumy konsepsiýasy we dizaýn çözgüdi", level=2)
    for text in [
        (
            "Programma interfeýsiniň esasy konsepsiýasy administrator üçin düşnükli, tertipli we çalt "
            "ulanylýan iş gurşawyny döretmekden ybaratdyr. Ulanyjy komanda setirinde Nmap parametrlerini "
            "ýat tutman, grafiki penjire arkaly adapter saýlap, IP aralygyny görüp, skanirlemäni başladyp "
            "bilýär. Şeýlelikde, programma professional tor gurallarynyň çylşyrymlylygyny azaltmaga we "
            "diplom goragy wagtynda netijeleri görkezmegi ýeňilleşdirmäge hyzmat edýär."
        ),
        (
            "Esasy penjire ui/main_window.py faýlynda ýerleşýän MainWindow synpy arkaly döredilýär. Penjire "
            "çep tarapdaky navigasiýa panelinden we sag tarapdaky esasy mazmun meýdanyndan durýar. Çep "
            "panelde Baş sahypa, Skanirleme, Panel, Netijeler, Taryh we Sazlamalar sahypalary ýerleşdirildi. "
            "Sahypalar QStackedWidget arkaly çalşyrylýar, bu bolsa programma içinde birnäçe ekrany bir "
            "penjirede dolandyrmaga mümkinçilik berýär."
        ),
        (
            "Interfeýsde gara tema başlangyç görnüş hökmünde saýlanyldy. Gara tema tor howpsuzlygy bilen "
            "işleýän programma üçin amatlydyr, sebäbi tablisa, log we netijeler uzak wagtlap okalanda göz "
            "ýadawlygyny azaldýar. Şeýle hem programma ýagty tema geçmek mümkinçiligini saklaýar. Tema "
            "çalşygy MainWindow içindäki toggle_theme funksiýasy arkaly ýerine ýetirilýär."
        ),
        (
            "Dizaýn çözgüdinde esasy üns maglumatlaryň okalmagyna gönükdirildi. Tablisalaryň sütünleri "
            "IP, ady, MAC, öndüriji, görnüşi, tapylan usul, portlar, töwekgelçilik, gowşaklyk we maslahat "
            "ýaly maglumatlary görkezýär. Uzyn tekstler üçin tooltip ulanylýar, netijede tablisa gysga "
            "görnüşde saklanýar, emma ulanyjy zerur bolan maglumatlary doly görüp bilýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 4.1 - Interfeýs dizaýnynyň esasy ýörelgeleri",
        ["Ýörelge", "Durmuşa geçirilişi", "Ulanyjy üçin peýdasy"],
        [
            ["Ýönekeý navigasiýa", "Çep panelde esasy sahypalar ýerleşdirildi", "Ulanyjy zerur bölüme çalt geçýär"],
            ["Gara tema", "Başlangyç tema hökmünde dark palette ulanylýar", "Uzyn skanirleme netijelerini okamak ýeňilleşýär"],
            ["Türkmen dili", "Düwmeler, sütünler we habarlar türkmençe berilýär", "Diplom goragynda düşündiriş aňsatlaşýar"],
            ["Tertipli tablisa", "Netijeler aýratyn sütünlerde görkezilýär", "IP, port, risk we maslahat çalt tapylýar"],
            ["Real wagt logy", "Skanirleme wagtynda ýagdaý ýazgylary görkezilýär", "Ulanyjy programma näme edýändigini görýär"],
        ],
    )

    add_heading(doc, "4.2. Baş sahypa we tor skanirleme sahypasynyň durmuşa geçirilmegi", level=2)
    for text in [
        (
            "Baş sahypa programma girilende ulanyja umumy maglumat we çalt hereketler görkezmek üçin "
            "niýetlenendir. Bu sahypada ähli gurluşlar, gowşak gurluşlar, goralan gurluşlar we soňky "
            "skanirleme ýaly görkezijiler ýerleşdirilýär. Şeýle çemeleşme programma açylan badyna umumy "
            "ýagdaý barada düşünje almaga kömek edýär."
        ),
        (
            "Skanirleme sahypasy programma toplumynyň iň möhüm interfeýs bölegidir. Bu sahypada ulanyjy "
            "tor adapterini saýlaýar, IP aralygyny görýär, güýç derejesini we skanirleme görnüşini belleýär. "
            "Adapter saýlananda IP aralygy awtomatik täzelenýär. Mysal üçin, Wi-Fi adapteriniň salgysy "
            "192.168.1.3/24 bolsa, IP aralygy 192.168.1.0/24 görnüşinde görkezilýär."
        ),
        (
            "Skanirleme başlamazdan öň IP aralygynyň dogrylygy barlanýar. Eger ulanyjy daşarky ýa-da nädogry "
            "aralyk girizse, programma duýduryş görkezýär. Bu gorag mehanizmi tötänleýin daşarky tory "
            "skanirlemegiň öňüni alýar we programmanyň diňe lokal tor auditi üçin ulanylmagyna kömek edýär."
        ),
        (
            "Skanirleme wagtynda progress bar we log meýdany işjeňleşýär. Her tapylan gurluş tablisa goşulýar "
            "we log bölüminde onuň IP salgysy, görnüşi, tapylan usuly, portlary we töwekgelçilik derejesi "
            "ýazylýar. Ulanyjy skanirleme tamamlanmagyna garaşman, netijeleriň peýda bolşuny görüp bilýär."
        ),
        (
            "Sahypada üç esasy düwme bar: skanirlemäni başlatmak, skanirlemäni saklamak we netijeleri eksport "
            "etmek. Skanirleme tamamlanandan soň netijeler database.py arkaly saklanýar we gerek bolan "
            "ýagdaýynda PDF hasabata eksport edilýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 4.2 - Skanirleme sahypasynyň interfeýs elementleri",
        ["Element", "Wezipesi", "Baglanyşykly programma logikasy"],
        [
            ["Tor adapteri", "Wi-Fi/Ethernet adapterini saýlamak", "get_network_adapters we adapter maglumatlary"],
            ["IP aralygy", "Skanirlenjek lokal tory görkezmek", "ipaddress arkaly validasiýa"],
            ["Güýç derejesi", "Akym sanyny saýlamak", "max_workers bahasyny kesgitlemek"],
            ["Netijeler tablisasy", "Tapylan gurluşlary görkezmek", "device_found signaly we fill table logikasy"],
            ["Progress bar", "Skanirleme ýagdaýyny görkezmek", "progress signaly"],
            ["Log meýdany", "Amallaryň gidişini ýazmak", "append_log funksiýasy"],
            ["PDF eksport", "Netijeleri hasabata geçirmek", "generate_pdf funksiýasy"],
        ],
    )

    add_heading(doc, "4.3. Panel, Netijeler we Taryh sahypalarynyň işleýşi", level=2)
    for text in [
        (
            "Programmada Panel, Netijeler we Taryh sahypalary biri-birine meňzeş görünmez ýaly aýratyn "
            "wezipeler bilen taslanyldy. Panel sahypasy umumy ýagdaýy görkezýär, Netijeler sahypasy iň "
            "soňky skanirlemäni jikme-jik seljerýär, Taryh sahypasy bolsa öňki skanirlemeleri arhiw görnüşinde "
            "saklap görkezýär."
        ),
        (
            "Panel sahypasynda ähli skanirlemeler, soňky tapylan gurluşlar, gowşak gurluşlar, portly gurluşlar, "
            "howpsuz gurluşlar we soňky skanirleme wagty ýaly görkezijiler kart görnüşinde görkezilýär. "
            "Mundan başga-da, soňky skanirleme boýunça gysgaça bellikler, töwekgelçilik paýlanyşy we gurluş "
            "görnüşleriniň paýlanyşy aýratyn tablisalarda görkezilýär."
        ),
        (
            "Netijeler sahypasy soňky skanirlemäniň gurluşlaryny jikme-jik öwrenmek üçin niýetlenendir. "
            "Bu sahypada filtr ulgamy bar: ähli gurluşlar, gowşak gurluşlar, portly gurluşlar, Router, PC, "
            "IoT kandidat we Telefon/Unknown ýaly görnüşler boýunça saýlap bolýar. Ulanyjy tablisadan bir "
            "gurluşy saýlanda onuň IP, MAC, öndüriji, portlar, risk we maslahat maglumatlary aýratyn tekst "
            "meýdanynda görkezilýär."
        ),
        (
            "Taryh sahypasy öňki skanirleme sessiýalaryny görkezýär. Tablisada skanirleme wagty, tor aralygy, "
            "adapter, gurluş sany, gowşak gurluş sany, portly gurluş sany we iň ýokary töwekgelçilik derejesi "
            "berilýär. Ulanyjy islendik skanirlemäni saýlap, onuň gysgaça mazmunyny görüp ýa-da şol skanirleme "
            "boýunça PDF hasabat döredip bilýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 4.3 - Panel, Netijeler we Taryh sahypalarynyň tapawudy",
        ["Sahypa", "Esasy wezipesi", "Ulanyjy üçin peýdasy"],
        [
            ["Panel", "Umumy statistika we soňky skanirleme boýunça ýagdaý", "Administrator umumy howpsuzlyk ýagdaýyny çalt görýär"],
            ["Netijeler", "Soňky skanirlemäniň gurluşlaryny jikme-jik seljermek", "Filtr we saýlanan gurluşyň analizi arkaly netije düşündirilýär"],
            ["Taryh", "Öňki skanirlemeleri saklamak we görkezmek", "Netijeleri deňeşdirmek we öňki hasabatlary eksport etmek mümkin bolýar"],
        ],
    )

    add_heading(doc, "4.4. Sazlamalar sahypasy we giňeldilýän interfeýs mümkinçilikleri", level=2)
    for text in [
        (
            "Sazlamalar sahypasy programma arhitekturasynda geljekde giňeldiljek bölüm hökmünde göz öňünde "
            "tutuldy. Häzirki görnüşinde bu sahypa ulgam sazlamalary üçin aýratyn ýer hökmünde görkezilýär. "
            "Esasy tema çalşygy bolsa çep paneldäki “Temany çalyş” düwmesi arkaly ýerine ýetirilýär."
        ),
        (
            "Geljekde Sazlamalar sahypasynda Nmap ýoluny görkezmek, skanirleme timeout bahalaryny üýtgetmek, "
            "PDF hasabat bukjasyny saýlamak, email duýduryşlaryny sazlamak we taryhy arassalamak ýaly "
            "mümkinçilikleri ýerleşdirmek meýilleşdirilip bilner. Şeýle giňeldilýän gurluş programma "
            "arhitekturasynyň geljekki ösüşe açykdygyny görkezýär."
        ),
        (
            "Interfeýsde ulanyjy tejribesini gowulandyrmak üçin tablisalarda horizontal scroll, sütünleriň "
            "gysgalma düzgüni, tooltip, saýlanan setiriň aýratyn görkezilmegi we refresh düwmeleri ulanylýar. "
            "Bu ownuk görünýän elementler programma bilen işlemekde möhüm ähmiýete eýedir, sebäbi tor "
            "skanirleme netijeleri köplenç köp sütünli we uzyn tekstli bolýar."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 4.4 - Interfeýsi giňeltmek üçin meýilleşdirilýän mümkinçilikler",
        ["Mümkinçilik", "Mazmuny", "Amaly ähmiýeti"],
        [
            ["Nmap sazlamasy", "Nmap ýoluny we parametrlerini görkezmek", "Dürli kompýuterlerde sazlamagy ýeňilleşdirýär"],
            ["Skanirleme parametrleri", "Timeout, port sanawy we akym sanyny sazlamak", "Tor ölçegine görä skanirlemäni uýgunlaşdyrýar"],
            ["Hasabat sazlamalary", "PDF faýlynyň ady we saklanýan ýeri", "Hasabat bilen işlemegi tertipleşdirýär"],
            ["Duýduryş sazlamalary", "Email ýa-da webhook maglumatlaryny girizmek", "Gowşaklyk tapylanda habar bermäge mümkinçilik berýär"],
            ["Taryh dolandyryşy", "Köne ýazgylary arassalamak ýa-da saklama möhleti", "Maglumat bazasynyň tertipli saklanmagyna kömek edýär"],
        ],
    )

    add_heading(doc, "4.5. Interfeýsiň diplom goragynda görkezilişi", level=2)
    for text in [
        (
            "Diplom goragy wagtynda programma interfeýsiniň görkezilişi aýratyn ähmiýete eýedir. Komissiýa "
            "programma diňe kod derejesinde däl, eýsem işleýän amaly çözgüt hökmünde görkezilmelidir. Şonuň "
            "üçin demonstrasiýa tor adapterini saýlamakdan, skanirlemäni başlatmakdan, tapylan gurluşlary "
            "düşündirmekden we PDF hasabat döretmekden ybarat yzygiderli ssenariýa boýunça geçirilip bilner."
        ),
        (
            "Ilki bilen Skanirleme sahypasynda Wi-Fi adapteri saýlanýar we IP aralygynyň 192.168.1.0/24 ýaly "
            "lokal tor görnüşinde dogry görkezilýändigi düşündirilýär. Soňra skanirleme başladylyp, progress "
            "bar we log ýazgylary görkezilýär. Tapylan gurluşlaryň IP, MAC, öndüriji, portlar, tapylan usul "
            "we töwekgelçilik derejesi tablisa boýunça düşündirilýär."
        ),
        (
            "Eger torda IoT kamera ýa-da açyk hyzmatly enjam ýok bolsa, programma şonda-da routeri, kompýuteri "
            "we torda görünýän beýleki gurluşlary görkezýär. Bu ýagdaýda portlar “ýok” diýlip görkezilýär we "
            "programma gurluşyň diňe torda ýüze çykarylandygyny düşündirýär. Eger IoT enjamly tora birikdirilse, "
            "programma onuň açyk portlaryny we degişli töwekgelçilik maslahatlaryny görkezmäge ukyplydyr."
        ),
        (
            "Demonstrasiýanyň indiki tapgyrynda Panel sahypasy açylyp, umumy statistika görkezilýär. Netijeler "
            "sahypasynda iň soňky skanirleme filtrlenýär we bir gurluşyň jikme-jik seljermesi görkezilýär. "
            "Taryh sahypasynda öňki skanirlemeleriň saklanýandygy görkezilýär. Ahyrynda PDF hasabat döredilip, "
            "programma netijeleriniň resmileşdirilip bilinýändigi subut edilýär."
        ),
    ]:
        add_normal(doc, text)

    add_table(
        doc,
        "Tablisa 4.5 - Diplom goragynda interfeýsi görkezmegiň ssenariýasy",
        ["Ädim", "Görkezilýän amal", "Komissiýa düşündirilýän pikir"],
        [
            ["1", "Programma açylýar we Baş sahypa görkezilýär", "Programma desktop görnüşinde özbaşdak işleýär"],
            ["2", "Skanirleme sahypasynda Wi-Fi adapter saýlanýar", "Programma dogry lokal tory saýlap bilýär"],
            ["3", "Skanirleme başladylyp, netijeler tablisa düşýär", "IP, MAC, port we risk maglumatlary real wagtda görkezilýär"],
            ["4", "Panel sahypasynda statistika görkezilýär", "Netijeler umumy howpsuzlyk ýagdaýyna öwrülýär"],
            ["5", "Netijeler sahypasynda filtr we gurluş seljermesi görkezilýär", "Her gurluş boýunça maslahat almak bolýar"],
            ["6", "Taryh sahypasynda öňki skanirlemeler görkezilýär", "Netijeler maglumatlar bazasynda saklanýar"],
            ["7", "PDF hasabat eksport edilýär", "Audit netijesi resmi faýl görnüşinde berilýär"],
        ],
    )

    add_normal(
        doc,
        "Şeýlelikde, dördünji bölümde IoT Security Scanner programma toplumynyň ulanyjy interfeýsi "
        "seljerildi. Interfeýs tor howpsuzlyk auditini tehniki taýdan çylşyrymly komandalar toplumyndan "
        "düşnükli grafiki iş prosesine öwürýär. Sahypalaryň wezipeler boýunça bölünmegi, türkmen dilindäki "
        "ýazgylar, real wagt netijeleri, taryh we PDF eksport mümkinçilikleri programmany diplom goragynda "
        "görkezmek üçin amatly edýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_styles(main_doc):
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text == "4. Ulanyjy interfeýsiniň işlenip taýýarlanylyşy":
            paragraph.style = "Heading 1"
            set_font(paragraph, 14, True)
        elif text.startswith(("4.1.", "4.2.", "4.3.", "4.4.", "4.5.")):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip() == "4. Ulanyjy interfeýsiniň işlenip taýýarlanylyşy" for p in main_doc.paragraphs):
        print("Section 4 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_styles(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
