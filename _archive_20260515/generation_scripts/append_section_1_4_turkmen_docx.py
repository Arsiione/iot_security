from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_1_4.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "1.4. Bar bolan skanirleme ulgamlarynyň deňeşdirme seljermesi")

    paragraphs = [
        (
            "IoT torlarynyň howpsuzlygyny üpjün etmek üçin dünýä tejribesinde dürli skanirleme, "
            "monitoring we gowşaklyk seljeriş gurallary ulanylýar. Şeýle gurallar tor gurluşlaryny "
            "tapmaga, açyk portlary kesgitlemäge, hyzmatlaryň wersiýalaryny barlamaga we belli "
            "gowşaklyklary ýüze çykarmaga kömek edýär. Emma her bir çözgüdiň maksady, ulanyş "
            "çylşyrymlylygy, bahasy we lokal IoT gurşawyna laýyklygy tapawutlanýar. Şonuň üçin "
            "diplom taslamasynda işlenip taýýarlanylýan IoT Security Scanner programmasynyň "
            "ornuny kesgitlemek üçin bar bolan analoglar bilen deňeşdirme seljermesi geçirilýär."
        ),
        (
            "Nmap tor skanirlemesi üçin iň giňden ulanylýan açyk çeşmeli gurallaryň biridir. "
            "Ol IP aralyklaryny barlap, işjeň hostlary, açyk portlary we hyzmatlary kesgitläp "
            "bilýär. Nmap güýçli we çeýe gural bolmagyna garamazdan, onuň netijeleri köplenç "
            "komanda setiri görnüşinde berilýär we täze ulanyjy üçin düşündiriş talap edýär. "
            "Mundan başga-da, Nmap adaty ýagdaýda tapylan portuň näme üçin howply bolup "
            "biljekdigini ýa-da ulanyjynyň nähili düzediş etmelidigini doly düşündirmeýär. "
            "Şu sebäpli diplom taslamasynda Nmap skanirleme ýadrosy hökmünde peýdalanylyp, onuň "
            "netijeleri has düşnükli grafiki interfeýs, töwekgelçilik derejesi we maslahatlar "
            "bilen baýlaşdyrylýar."
        ),
        (
            "Shodan internetde açyk duran gurluşlary gözlemek üçin niýetlenen hyzmatdyr. Ol "
            "dünýä boýunça internetden görünýän kameralar, routerler, serwerler we beýleki "
            "gurluşlar barada maglumat ýygnap bilýär. Şeýle mümkinçilik global howpsuzlyk "
            "seljermesi üçin peýdaly bolsa-da, diplom taslamasynyň meselesi başga häsiýete "
            "eýedir: programma lokal Wi‑Fi ýa-da Ethernet torundaky gurluşlary barlamaly. "
            "Shodan köplenç içerki lokal segmentdäki telefon, kamera ýa-da akylly öý enjamyny "
            "görkezip bilmeýär, sebäbi olar internetden gönüden-göni elýeterli däl. Şonuň üçin "
            "lokal tor auditi üçin aýratyn programma zerurdyr."
        ),
        (
            "Nessus we OpenVAS ýaly gowşaklyk skanerleri has giň mümkinçiliklere eýedir. Olar "
            "hyzmatlaryň wersiýalaryny, belli CVE gowşaklyklaryny we konfigurasiýa meselelerini "
            "has çuňňur seljerip bilýär. Şeýle-de bolsa, bu ulgamlar köplenç hünärmenler üçin "
            "niýetlenendir, sazlamasy çylşyrymly bolup biler we käbir ýagdaýlarda lisenziýa ýa-da "
            "serwer infrastrukturasyny talap edýär. Diplom taslamasynda bolsa esasy maksat okuw "
            "we kiçi lokal torlar üçin ýeňil, düşnükli we çalt ulanarlykly çözgüt döretmekdir. "
            "Şu nukdaýnazardan IoT Security Scanner ýöriteleşdirilen, grafiki interfeýsli we "
            "ulanyja gönükdirilen programma hökmünde tapawutlanýar."
        ),
        (
            "IoT Security Scanner programmasynyň esasy aýratynlygy onuň lokal tor we IoT "
            "gurluşlary üçin ýörite uýgunlaşdyrylmagydyr. Programma Wi‑Fi/Ethernet adapterini "
            "saýlap, dogry IP aralygyny kesgitlemäge kömek edýär, Nmap, ARP we Ping ýaly usullary "
            "birleşdirip gurluşlary tapýar, portlar açyk bolmasa hem enjamy netijeler sanawyna "
            "goşýar. Bu aýratynlyk telefon, router ýa-da IoT gurluşy diňe torda görünýän ýagdaýynda "
            "hem administratora doly tor kartasyny almaga mümkinçilik berýär."
        ),
        (
            "Mundan başga-da, programma diňe tehniki maglumatlary görkezmek bilen çäklenmeýär. "
            "Her bir tapylan gurluş üçin IP salgysy, ady, MAC salgysy, öndüriji, görnüşi, tapylan "
            "usuly, açyk portlary, töwekgelçilik derejesi, gowşaklyk ýagdaýy we maslahat görkezilýär. "
            "Şeýle çemeleşme diplom taslamasynyň remediation, ýagny gowşaklyklary azaltmaga "
            "gönükdirilen ugruny güýçlendirýär. Ulanyjy diňe “port açyk” diýen netijäni däl, eýsem "
            "näme etmeli diýen amaly maslahatlary hem alýar."
        ),
    ]

    for text in paragraphs:
        add_normal(doc, text)

    caption = doc.add_paragraph("Tablisa 1.4 - Bar bolan analoglaryň deňeşdirme seljermesi")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    headers = ["Gural", "Esasy maksady", "Artykmaçlygy", "Çäklendirmesi", "Diplom taslamasy bilen tapawudy"]
    rows = [
        [
            "Nmap",
            "Tor hostlaryny we portlary skanirlemek",
            "Güýçli, açyk çeşmeli, çeýe parametrli",
            "Netijeleri täze ulanyjy üçin çylşyrymly bolup biler",
            "IoT Security Scanner Nmap netijelerini grafiki görnüşde düşündirýär we maslahat berýär",
        ],
        [
            "Shodan",
            "Internetden görünýän gurluşlary gözlemek",
            "Global gözleg we köp maglumat bazasy",
            "Lokal Wi‑Fi torundaky içerki enjamlary görkezmeýär",
            "Taslama lokal toruň içindäki gurluşlary barlamaga gönükdirilýär",
        ],
        [
            "Nessus",
            "Çuňňur gowşaklyk skanirlemesi",
            "Köp CVE we professional hasabat mümkinçilikleri",
            "Sazlamasy çylşyrymly, lisenziýa talap edip biler",
            "Taslama ýeňil, okuw we kiçi torlar üçin düşnükli çözgüt berýär",
        ],
        [
            "OpenVAS",
            "Açyk çeşmeli gowşaklyk auditi",
            "Giň plugin bazasy we çuňňur seljeriş",
            "Serwer sazlamasy we dolandyryşy has kyn",
            "Taslama bir programma içinde çalt skanirleme we PDF hasabat hödürleýär",
        ],
        [
            "IoT Security Scanner",
            "Lokal IoT toruny tapmak, portlary barlamak we maslahat bermek",
            "Türkmen interfeýsi, adapter saýlamak, ARP/Nmap/Ping birleşmesi, PDF hasabat",
            "Professional enterprise skanerleriň ähli çuňňur CVE bazasyny çalyşmaýar",
            "Diplom taslamasy üçin ýörite işlenip, IoT gurluşlaryna we remediation maslahatlaryna gönükdirilýär",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)

    add_normal(
        doc,
        "Deňeşdirme seljermesinden görnüşi ýaly, IoT Security Scanner programmasy Nmap ýa-da "
        "Nessus ýaly professional gurallary doly çalyşmak üçin däl, eýsem lokal IoT toruny "
        "düşnükli görnüşde barlamak we ulanyja amaly maslahat bermek üçin döredilýär. Bu onuň "
        "diplom taslamasyndaky ylmy-amaly ähmiýetini ýokarlandyrýar, sebäbi programma hakyky "
        "ulanyjy meselesini çözýär: torda haýsy gurluşlaryň bardygyny görmek, olaryň portlaryny "
        "barlamak, töwekgelçiligi bahalandyrmak we hasabat taýýarlamak."
    )
    add_normal(
        doc,
        "Şeýlelikde, bar bolan analoglar bilen deňeşdirilende taslamanyň esasy artykmaçlygy "
        "lokal tor üçin ýöriteleşdirilenligi, türkmen dilindäki grafiki interfeýsi, netijeleriň "
        "ýönekeý düşündirilmegi, PDF hasabatynyň döredilmegi we gowşaklyklary azaltmak boýunça "
        "maslahatlaryň berilmegidir. Bu aýratynlyklar IoT Security Scanner programmasyny okuw, "
        "kiçi kärhana we diplom goragy şertlerinde görkezmek üçin amatly çözgüt edýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("1.4.") for p in main_doc.paragraphs):
        print("Section 1.4 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
