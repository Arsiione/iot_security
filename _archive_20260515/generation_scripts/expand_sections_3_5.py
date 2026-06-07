from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = ROOT / "kakabalowa.docx"


SECTION_3 = [
    (
        "3.6. Tor adapterlerini kesgitlemek we IP aralygyny awtomatik saýlamak",
        [
            "Hasabat döretmek we duýduryş ulgamynyň durmuşa geçirilmegi bilen programma toplumy ulanyja diňe netijäni görkezmek bilen çäklenmän, skanirleme üçin dogry başlangyç şertleri hem döretmelidir. Şol sebäpli IoT Security Scanner programmasynda tor adapterlerini kesgitlemek aýratyn amaly mesele hökmünde ýerine ýetirildi. Windows gurşawynda bir kompýuterde birnäçe adapter bolup biler: Wi‑Fi, Ethernet, VPN, wirtual adapter, Npcap loopback ýa-da wagtlaýyn 169.254.x.x görnüşli awtomatik salgy. Eger programma nädogry adapteri saýlasa, ulanyjy özüniň hakyky lokal toruny däl-de, VPN ýa-da wirtual segmenti skanirläp biler. Bu ýagdaýda telefona, routere ýa-da IoT gurluşlaryna degişli maglumatlar tablisa düşmez.",
            "Adapterleri kesgitlemekde esasy maksat ulanyjynyň el bilen çylşyrymly `ipconfig` ýa-da `route print` buýruklaryny seljermeli bolmazlygyny üpjün etmekdir. Programma işjeň tor interfeýslerini okaýar, olaryň IP salgysyny, maskasyny we mümkin bolan şlýuzyny kesgitleýär. Soňra Wi‑Fi we Ethernet adapterleri ileri tutulýar, sebäbi diplom taslamasynda esasy barlag gurşawy lokal simsiz ýa-da simli tor hökmünde kabul edilýär. VPN adapterleriniň doly gadagan edilmezligi hem möhümdir: käbir kärhana torlarynda administrator hakykatdan hem VPN segmentini barlamaly bolup biler. Emma adaty demonstrasiýa ýagdaýynda programma VPN-i awtomatik esasy adapter hökmünde saýlamazdan, ulanyja hakyky lokal tory hödürleýär.",
            "IP aralygy adapteriň IP salgysy we maskasy boýunça hasaplanylýar. Mysal üçin, kompýuter `192.168.1.3/24` salgysynda bolsa, programma skanirleme meýdanyna `192.168.1.0/24` aralygyny ýerleşdirýär. Şeýlelikde, ulanyjy tordaky 254 mümkin hosty barlamak üçin salgy aralygyny el bilen ýazyp oturmaýar. Bu çözgüt diplom goragynda hem amatlydyr, sebäbi komissiýa programmasynyň diňe taýýar sanawy görkezmeýändigini, eýsem kompýuteriň häzirki tor ýagdaýyndan ugur alyp dogry başlangyç parametrleri döredýändigini görýär.",
            "Adapter saýlanylanda nädogry ýa-da peýdasyz salgylary süzmek hem ýerine ýetirilýär. Loopback, 127.0.0.1 görnüşli ýerli salgylar, 169.254.x.x görnüşli awtomatik link-local salgylar we diňe wirtual gurşaw üçin döredilen interfeýsler esasy saýlawdan çykarylýar. Bu süzgüçler programma netijeleriniň hilini ýokarlandyrýar. Sebäbi IoT skanirlemesinde iň möhüm zat dogry tor segmentini saýlamakdyr. Nädogry segmentde skanirleme üstünlikli tamamlanan ýaly görünse-de, hakykatda ulanyjy gerek bolan gurluşlary tapyp bilmez. Şonuň üçin adapter saýlamak programmadaky skanirleme algoritminiň ilkinji gorag gatlagy hökmünde çykyş edýär.",
            "Bu bölekde ulanyjy interfeýsi bilen programma ýadrosynyň arasyndaky baglanyşyk hem aýdyň görünýär. Interfeýsde adapterleriň atlary düşnükli görnüşde görkezilýär: mysal üçin, `Wi‑Fi 192.168.1.3/24` ýa-da `Ethernet 192.168.0.15/24`. Ulanyjy adapteri üýtgeden badyna IP aralygy hem täzelenýär. Şeýle mehanizm programmanyň amaly ulanylyşyny ýeňilleşdirýär we diplom işiniň esasy pikirini, ýagny howpsuzlyk barlagynyň ulanyjy üçin düşnükli görnüşe geçirilmegini görkezýär.",
        ],
    ),
    (
        "3.7. Gurluşlary tapmagyň birleşdirilen usullary",
        [
            "Adapter we IP aralygy dogry saýlanylandan soň programma lokal torda ýerleşýän gurluşlary tapmaga girişýär. Bu ýerde diňe bir usula daýanmak ýeterlik däldir, sebäbi dürli gurluşlar tor soraglaryna dürli görnüşde jogap berýär. Käbir telefonlar ICMP ping soragyna jogap bermeýär, käbir IoT enjamlarynda açyk port ýok, käbir routerler bolsa ARP soragynda görünýär, emma TCP port barlagynda çäkli maglumat berýär. Şol sebäpli IoT Security Scanner programmasynda gurluş tapmak birleşdirilen usulda ýerine ýetirildi.",
            "Birinji usul hökmünde Nmap ping scan ulanylýar. Bu usul `nmap -sn` görnüşindäki barlag bilen IP aralygyndaky işjeň hostlary kesgitlemäge kömek edýär. Onuň artykmaçlygy Nmap-yň tor howpsuzlygynda giňden ulanylýan ygtybarly gural bolmagydyr. Emma Nmap-yň netijesi hem hemişe doly bolman biler. Şonuň üçin programma diňe Nmap netijesine daýanman, ARP cache, ARP scan, ICMP ping sweep we gateway maglumatlaryny hem birleşdirýär.",
            "ARP usuly lokal tor üçin aýratyn ähmiýetlidir. Eger kompýuter şol bir Wi‑Fi ýa-da Ethernet segmentinde ýerleşýän gurluş bilen aragatnaşyk saklan bolsa, onuň MAC salgysy Windows ARP cache içinde görünip biler. Programma `arp -a` netijesini gizlin proses hökmünde okaýar we şol ýerden IP-MAC jübütlerini alýar. Bu usul esasan telefonlary we port açmaýan gurluşlary görkezmekde peýdalydyr. Şeýle ýagdaýda tablisa portlar hökmünde `ýok` ýazylsa-da, gurluşyň torda ýüze çykarylandygy görkezilýär.",
            "Ping sweep fallback usuly hökmünde ulanylýar. Bu usul ähli IP salgylara gysga wagtlaýyn ICMP sorag iberýär. Käbir gurluşlar ping-e jogap berýär, käbirleri jogap bermeýär. Şonuň üçin ping aýratynlykda doly netije bermeýär, emma beýleki usullar bilen birleşdirilende umumy tapyş takyklygyny ýokarlandyrýar. Şeýle hem programma router ýa-da default gateway salgysyny aýratynlykda netijelere goşýar. Sebäbi lokal toruň howpsuzlyk seljermesinde şlýuz merkezi element bolup durýar: ähli gurluşlar internet ýa-da beýleki tor segmentleri bilen aragatnaşygy köplenç şol enjam arkaly alyp barýar.",
        ],
        [
            ["Usul", "Programmadaky wezipesi", "Aýratyn peýdasy"],
            ["Nmap -sn", "IP aralygyndaky işjeň hostlary tapmak", "Tor skanirlemesinde ygtybarly başlangyç netije berýär"],
            ["ARP cache", "Windows ulgamyndaky IP-MAC maglumatlaryny almak", "Port açmaýan telefon we IoT kandidatlaryny görkezmäge kömek edýär"],
            ["Ping sweep", "Jogap berýän IP salgylary fallback görnüşde barlamak", "Nmap ýa-da ARP bilen tapylmadyk käbir hostlary ýüze çykaryp biler"],
            ["Gateway", "Routeri netijelere goşmak", "Lokal toruň merkezi gurluşyny hökmany görkezýär"],
        ],
            "Tablisa 3.6 - Gurluşlary tapmakda ulanylýan birleşdirilen usullar",
    ),
    (
        "3.8. Programma modulynyň giňeldilmegi we goldaw mümkinçilikleri",
        [
            "Gurluşlary tapmak we portlary barlamak mehanizmi diňe häzirki diplom taslamasy üçin däl, eýsem geljekki giňeltmeler üçin hem esas bolup hyzmat edýär. Programma modul görnüşinde gurlandygy sebäpli täze barlag usullaryny goşmak, täze IoT öndürijileri boýunça plugin ýazmak ýa-da täze hasabat görnüşini taýýarlamak mümkinçiligi saklanýar. Bu aýratynlyk taslamanyň diňe bir gezeklik demonstrasiýa däl-de, ösdürilip bilinýän programma çözgüdidigini görkezýär.",
            "Plugin ulgamynyň esasy pikiri belli bir hyzmat ýa-da öndüriji bilen baglanyşykly gowşaklygy aýratyn modulda barlamakdan ybaratdyr. Mysal üçin, Telnet hyzmaty açyk bolsa, telnet_weak_auth plugin gowşak login-parol jübütlerini barlap biler. Web interfeýs ýa-da kamera hyzmatlary boýunça bolsa Hikvision ýaly öndürijiler bilen baglanyşykly barlag logikasy aýratyn modulda ýerleşdirilip bilner. Şeýle çemeleşme esasy skanirleme ýadrosyny aşa çylşyrymlaşdyrmazdan, programma mümkinçiliklerini giňeltmäge ýol açýar.",
            "Goldaw mümkinçilikleriniň ýene biri maglumatlar bazasynyň giňeldilmegidir. Häzirki ýagdaýda SQLite bazasynda IP, MAC, host ady, portlar, töwekgelçilik derejesi, tapylan usul we maslahat ýaly maglumatlar saklanýar. Geljekde bu tablisa hyzmat wersiýalary, CVE belgileri, firmware maglumatlary ýa-da öndüriji boýunça kategoriýa ýaly sütünler goşulyp bilner. Şeýle maglumatlar programma tarapyndan döredilýän PDF hasabatlaryň hem mazmunyny baýlaşdyrar.",
            "Programmany goldamakda ýalňyş ýagdaýlary düşnükli görkezmek hem möhümdir. Mysal üçin, Nmap tapylmasa, adapter ýok bolsa, IP aralygy nädogry girizilse ýa-da skanirleme wagtynda tor aragatnaşygy kesilse, programma ulanyja düşnükli habar bermelidir. Bu talap diplom taslamasynda hem göz öňünde tutuldy. Ulanyjy diňe boş tablisa görmän, näme sebäpli az gurluş tapylandygyny ýa-da haýsy ýagdaýy barlamalydygyny düşündirýän log setirlerini görýär.",
            "Şeýlelikde, üçünji bölümde beýan edilen programma durmuşa geçirilişi skanirleme ýadrosyndan başlap, plugin, maglumatlar bazasy we hasabat ulgamyna çenli bitewi mehanizm hökmünde çykyş edýär. Täze goşulan adapter saýlamak, birleşdirilen tapyş usullary we modul giňeldilmegi baradaky düşündirişler programma toplumynyň amaly taýdan has durnukly we geljekde ösdürmäge taýýardygyny görkezýär.",
        ],
    ),
]


SECTION_4 = [
    (
        "4.6. Interfeýsde maglumatlaryň okalmagy we ulanyjy hereketleriniň tertibi",
        [
            "Interfeýsiň diplom goragynda görkezilişi beýan edilenden soň, ulanyjynyň programma bilen işleşýän wagty maglumatlary nähili kabul edýändigini aýratyn görkezmek zerurdyr. IoT Security Scanner programmasynda interfeýs diňe bezeg elementi hökmünde däl, eýsem skanirleme prosesini dolandyrýan esasy iş gurşawy hökmünde taslanyldy. Ulanyjy programma açylanda ilki umumy ýagdaýy görýär, soňra skanirleme sahypasyna geçýär, adapteri we IP aralygyny barlaýar, skanirlemäni başlaýar we netijeleri aýratyn sahypalarda seljerýär. Bu hereketleriň tertibi programma bilen ilkinji gezek işleýän adam üçin hem düşnükli bolmalydyr.",
            "Baş sahypada ulanyja iň möhüm maglumatlar gysga görnüşde berilýär. Bu ýerde programma ady, umumy ýagdaý görkezijileri we soňky hereketler ýerleşdirilýär. Maksat ulanyjyny derrew çylşyrymly tablisa bilen ýüzbe-ýüz goýmak däl-de, onuň haýsy bölüme geçmelidigini görkezmekdir. Çep tarapdaky menýu ähli sahypalara birmeňzeş geçiş berýär. Bu çözgüt desktop programma üçin amatlydyr, sebäbi ulanyjy bir gezek öwrenen navigasiýa modelini ähli sahypalarda ulanýar.",
            "Skanirleme sahypasynda maglumatlaryň okalmagy has işjeň häsiýete eýedir. Ýokarky bölekde adapter, IP aralygy, skanirleme derejesi we görnüşi ýerleşýär. Orta bölekde progress bar we netijeler tablisası bar. Aşaky bölekde bolsa log paneli we esasy düwmeler ýerleşdirilýär. Şeýle tertip skanirleme prosesiniň tapgyrlaryny ýokardan aşaklygyna okalmaga mümkinçilik berýär: ilki parametrler, soňra proses, soňra netije we ahyrynda amallar. Bu ýagdaý ulanyjynyň ünsüni tertipli ugrukdyrýar.",
            "Ulanyjy hereketleriniň tertibi hem çylşyrymly däl: adapter saýlanýar, IP aralygy barlanýar, skanirleme başlat düwmesi basylýar, netijeler tablisa düşýär, soňra saklamak ýa-da eksport etmek ýerine ýetirilýär. Bu tertip diplom goragynda düşündirilende hem amatlydyr. Komissiýa agzalary programmanyň nähili işleýändigini diňe söz bilen däl, eýsem interfeýsdäki hereketleriň yzygiderliligi arkaly hem görüp bilýär.",
            "Maglumatlaryň okalmagynda reňkleriň we sütünleriň hem ähmiýeti bar. Gara tema uzak tablisa bilen işlenende gözüň ýadamagyny azaldýar, açyk gök reňk işjeň elementleri görkezýär, ýaşyl ýa-da pes töwekgelçilikli ýagdaýlar howpsuzlygy aňladýar. Şeýlelikde, interfeýsde maglumatlar diňe tekst görnüşinde däl, eýsem wizual alamatlar arkaly hem berilýär. Bu bolsa programma netijelerini has çalt kabul etmäge kömek edýär.",
        ],
    ),
    (
        "4.7. Netijeleriň tablisa görnüşinde berilmegi we sahypalaryň tapawutlandyrylmagy",
        [
            "IoT Security Scanner programmasynda netijeler tablisa görnüşinde berilýär, sebäbi tor skanirlemesinde maglumatlaryň köp bölegi deňeşdirilýän häsiýete eýedir. Bir gurluşyň IP salgysy, MAC salgysy, öndürijisi, görnüşi, tapylan usuly, portlary, töwekgelçiligi we maslahatlary beýleki gurluşlar bilen deň hatarda görkezilýär. Şeýle gurluş ulanyja bir seredişde haýsy enjamyň has üns talap edýändigini görmäge mümkinçilik berýär.",
            "Skanirleme sahypasyndaky tablisa real wagt netijesini görkezmäge niýetlenendir. Skanirleme dowam edýärkä täze tapylan gurluşlar setir-setir goşulýar. Bu ýagdaý ulanyja programmanyň işleýändigini görkezýär we garaşmak döwründe näme bolup geçýändigini düşündirýär. Netijeler sahypasy bolsa soňky skanirlemäniň seljeriş görnüşini berýär. Bu ýerde ulanyjy tablisa maglumatlaryna has giňişleýin seredip, saýlanan gurluşyň loglaryny ýa-da maslahatlaryny okap biler.",
            "Panel sahypasy bilen Netijeler sahypasynyň tapawudy hem möhüm. Panel umumy ýagdaýy görkezýär: näçe gurluş tapyldy, näçe port açyk, näçe gowşaklyk bar, soňky skanirleme haçan geçirildi. Netijeler sahypasy bolsa her gurluş boýunça jikme-jik maglumat berýär. Taryh sahypasy öňki skanirlemeleri saklamak we deňeşdirmek üçin ulanylýar. Şeýlelikde, üç sahypa birmeňzeş maglumatlary gaýtalamaýar, tersine, şol maglumatlara dürli nukdaýnazardan seretmäge mümkinçilik berýär.",
        ],
        [
            ["Sahypa", "Esasy maglumat", "Ulanyjy üçin ähmiýeti"],
            ["Panel", "Umumy statistika we soňky ýagdaý", "Howpsuzlyk ýagdaýyny çalt görmek"],
            ["Netijeler", "Soňky skanirlemäniň jikme-jik tablisası", "Her gurluş boýunça IP, MAC, port we risk seljermek"],
            ["Taryh", "Öňki barlaglaryň sanawy we loglary", "Netijeleri deňeşdirmek we hasabat taýýarlamak"],
            ["Skanirleme", "Real wagt progress we tapylan gurluşlar", "Barlag prosesini göni dolandyrmak"],
        ],
            "Tablisa 4.1 - Interfeýs sahypalarynyň maglumat taýdan tapawudy",
    ),
    (
        "4.8. Ulanyjy tejribesini ýokarlandyrmak üçin kabul edilen çözgütler",
        [
            "Ulanyjy tejribesini ýokarlandyrmak üçin programma interfeýsinde birnäçe amaly çözgüt kabul edildi. Birinjiden, programmadaky ähli esasy ýazgylar türkmen dilinde berilýär. Bu diplom goragynda we okuw gurşawynda möhüm artykmaçlykdyr, sebäbi ulanyjy tehniki terminleri özüne düşnükli dilde görýär. Şeýle-de bolsa, IP, MAC, Nmap, TCP, ARP ýaly halkara tehniki düşünjeler öz ady bilen saklanýar, sebäbi olar tor howpsuzlygynda standart terminlerdir.",
            "Ikinjiden, programma gara konsol penjiresi açylmazdan işleýär. Tor skanirlemesi ýaly daşarky prosesler ulanyjy üçin gizlin ýerine ýetirilýär, netijeler bolsa grafiki interfeýsde görkezilýär. Bu çözgüt programmasyny diplom goragynda görkezmek üçin aýratyn amatlydyr. Komissiýa agzalary birnäçe gara penjire ýa-da çylşyrymly komanda setirleri bilen ýüzbe-ýüz bolman, taýýarlanan desktop programma bilen işleýşi görýär.",
            "Üçünjiden, programmada boş netije ýagdaýlary hem göz öňünde tutulýar. Eger torda diňe bir kompýuter tapylsa ýa-da telefon görünmese, programma munuň sebäbini düşündirýän log maglumatlaryny berip bilýär. Mysal üçin, Wi‑Fi diapazony, guest network, AP isolation, telefon ekranynyň ýagdaýy ýa-da privat MAC ulanylyşy ýaly ýagdaýlar ulanyja düşündirilip bilner. Şeýle habarlar programma ynamly görünmegine kömek edýär, sebäbi ol boş netijäni ýalňyşlyk hökmünde däl, anyklanmaly tor ýagdaýy hökmünde düşündirýär.",
            "Dördünjiden, eksport düwmesi we PDF hasabat taýýarlamak mümkinçilikleri interfeýsiň amaly ähmiýetini ýokarlandyrýar. Skanirleme netijeleri diňe ekranda galmaýar, eýsem dokumentleşdirilip bilýär. Bu aýratynlyk diplom işiniň hasabat bölümine hem laýyk gelýär: programma diňe barlag geçirýän gural däl, eýsem netijeleri resmileşdirýän programma toplumy hökmünde çykyş edýär.",
            "Umuman, dördünji bölümde beýan edilen interfeýs çözgütleri programma toplumynyň ulanyjy üçin düşnükli, yzygiderli we goragda görkezmäge amatly bolmagyny üpjün edýär. Baş sahypa, skanirleme, panel, netijeler, taryh we sazlamalar sahypalary bir bitewi iş akymyna birleşdirildi. Bu bolsa IoT howpsuzlyk auditini tehniki buýruklardan grafiki we düşündirilen programma tejribesine geçirýär.",
        ],
    ),
]


SECTION_5 = [
    (
        "5.6. Lokal torlarda goşmaça synag ssenariýalary",
        [
            "Öndürijilik, takyklyk we çäklendirmeler boýunça netije çykarylandan soň, programma toplumynyň dürli lokal tor ýagdaýlarynda özüni nähili alyp barýandygyny görkezmek hem möhümdir. Sebäbi IoT skanirlemesi diňe bir ideal Wi‑Fi torunda däl, dürli sazlamaly routerlerde, telefonlarda, kompýuterlerde we IoT kandidat enjamlarynda ulanylyp bilner. Şonuň üçin goşmaça synag ssenariýalary programma netijeleriniň durnuklylygyny bahalandyrmak üçin ulanyldy.",
            "Birinji ssenariýada adaty öý Wi‑Fi tory kabul edilýär. Bu ýagdaýda kompýuter, router we telefon bir segmentde ýerleşýär. Programma Wi‑Fi adapterini saýlap, `192.168.x.0/24` görnüşli IP aralygyny kesgitleýär. Garaşylýan netije hökmünde kompýuter we router hökmany tapylmalydyr, telefon bolsa onuň tor sazlamasyna baglylykda ARP ýa-da ping arkaly ýüze çykyp biler. Eger telefon privat MAC ulanýan bolsa, öndüriji kesgitlenmän `Näbelli` bolup galmagy mümkin. Bu ýagdaý ýalňyşlyk däl, häzirki mobil ulgamlaryň howpsuzlyk aýratynlygydyr.",
            "Ikinji ssenariýada VPN işjeň bolan ýagdaý barlanýar. Käbir kompýuterlerde VPN adapteri esasy IP salgysy ýaly görünip biler. Programma şeýle ýagdaýda Wi‑Fi ýa-da Ethernet adapterini ileri tutmaly, VPN-i bolsa diňe ulanyjy saýlasa skanirlemäge rugsat bermelidir. Bu ssenariýa diplom taslamasynda aýratyn ähmiýete eýedir, sebäbi öňki synaglarda nädogry tor segmentiniň saýlanmagy netijeleriň doly bolmazlygyna sebäp bolup bilýärdi.",
            "Üçünji ssenariýada port açmaýan gurluşlaryň görkezilmegi barlanýar. Telefon ýa-da käbir IoT enjamlary TCP port açmazdan hem torda işjeň bolup biler. Programma şeýle gurluşlary tablisa goşup, `Portlar: ýok`, `Töwekgelçilik: pes` we `Gurluş diňe torda ýüze çykaryldy` görnüşli maslahat bermelidir. Bu ýagdaý programma diňe port skaneri däl, eýsem host discovery ulgamydygyny görkezýär.",
        ],
        [
            ["Ssenariýa", "Barlanýan ýagdaý", "Garaşylýan netije"],
            ["Adaty Wi‑Fi tor", "Router, kompýuter we telefon bir segmentde", "Gateway we lokal kompýuter tapylýar, telefon ARP/Ping bilen görünip biler"],
            ["VPN işjeň", "Wi‑Fi bilen birlikde VPN adapter bar", "Programma Wi‑Fi/Ethernet-i ileri tutýar"],
            ["Port açmaýan gurluş", "Telefon ýa-da IoT kandidat port açmaýar", "Gurluş tablisa düşýär, risk pes bolýar"],
            ["Nädogry IP aralygy", "Ulanyjy nädogry CIDR girizýär", "Programma duýduryş berýär we skanirlemäni başlamaz"],
        ],
            "Tablisa 5.1 - Goşmaça synag ssenariýalary",
    ),
    (
        "5.7. Tapylan gurluşlaryň dogrulygyny barlamak usuly",
        [
            "Skanirleme netijeleriniň dogrulygyny barlamak üçin diňe programma tablisasyna seretmek ýeterlik däldir. Netijeleri tassyklamak maksady bilen birnäçe kömekçi çeşme ulanylyp bilner. Birinjiden, Windows ulgamynda `ipconfig` arkaly kompýuteriň hakyky IP salgysy we default gateway salgysy barlanýar. Ikinjiden, `arp -a` netijesi bilen programma tarapyndan görkezilen IP-MAC jübütleri deňeşdirilýär. Üçünjiden, routeriň administratiw panelindäki DHCP client list bilen tapylan gurluşlar deňeşdirilip bilner.",
            "Bu barlaglaryň maksady programma ýalňyş tapýar ýa-da tapmaýar diýen netijä howlukmazdan, lokal toruň aýratynlyklaryny hasaba almakdyr. Mysal üçin, käbir routerlerde client isolation işjeň bolsa, bir Wi‑Fi müşderisi beýleki müşderini görüp bilmez. Şeýle ýagdaýda telefon şol bir Wi‑Fi adyna birikdirilen hem bolsa, kompýuter tarapyndan skanirlenende görünmezligi mümkin. Programma muny doly aýyryp bilmeýär, sebäbi bu toruň howpsuzlyk sazlamasy bilen baglydyr. Emma loglarda şeýle ýagdaýy barlamagy maslahat bermek arkaly ulanyja dogry düşündiriş berýär.",
            "Dogrylyk barlagynda açyk portlaryň netijesi hem aýratyn seljerilýär. Eger routerde 80, 443, 22 ýa-da 53 portlary açyk görünse, bu adaty ýagdaý bolup biler. Emma Telnet 23 porty açyk bolsa ýa-da kamera hyzmatlary bilen baglanyşykly 554 RTSP porty görünse, programma töwekgelçiligi orta ýa-da ýokary derejede görkezýär. Şeýle ýagdaýlarda netijäni başga gural, mysal üçin Nmap komandasy bilen deňeşdirmek hem mümkin. Diplom goragynda bu çemeleşme programmasynyň professional tor gurallary bilen baglanyşyklydygyny görkezýär.",
            "Tapylan gurluşlaryň görnüşini kesgitlemekde MAC öndüriji maglumatlary hem ulanylýar. Eger MAC salgysy Intel, TP-Link, Xiaomi ýa-da başga öndüriji bilen baglanyşykly bolsa, programma gurluşyň görnüşini has düşnükli görkezmäge synanyşýar. Emma privat MAC ulanylanda öndüriji näbelli bolup biler. Şeýle ýagdaýda programma gurluşy `Telefon/Unknown` ýa-da `Näbelli` hökmünde görkezýär. Bu hem ýalňyşlyk däl, sebäbi häzirki mobil enjamlaryň gizlinlik mehanizmleri MAC öndürijini bilkastlaýyn gizläp bilýär.",
        ],
    ),
    (
        "5.8. Synag netijeleriniň diplom goragynda görkezilişi",
        [
            "Programma toplumyny synagdan geçirmek diňe içerki barlag hökmünde däl, eýsem diplom goragynda görkeziljek amaly subutnama hökmünde hem taýýarlanyldy. Gorag wagtynda programma EXE görnüşinde açylyp, çep menýudaky sahypalar yzygiderli görkezilýär. Ilki baş sahypa arkaly programmanyň umumy maksady düşündirilýär. Soňra skanirleme sahypasynda adapteriň awtomatik saýlanmagy we IP aralygynyň doldurylmagy görkezilýär.",
            "Skanirleme başladylandan soň progress bar, log paneli we netijeler tablisasynyň işleýşi görkezilýär. Bu ýerde komissiýa programmasynyň diňe statik ekran däldigini, hakyky tor barlagyny geçirýändigini görýär. Eger torda IoT gurluşy ýok bolsa, programma router, kompýuter ýa-da telefon ýaly bar bolan gurluşlary görkezýär. Şeýle ýagdaýda düşündiriş şeýle berilýär: programma lokal torda işjeň gurluşlary tapýar; eger IP kamera ýa-da beýleki IoT enjam birikdirilse, onuň IP, MAC, port we risk maglumatlary hem şol tablisa düşer.",
            "Netijeler sahypasynda tapylan gurluşlaryň jikme-jik maglumatlary düşündirilýär. IP we MAC salgylary, öndüriji, görnüş, tapylan usul, açyk portlar we maslahatlar aýratyn görkezilýär. Panel sahypasynda umumy statistika görkezilip, administrator üçin ýagdaýyň gysga görnüşde berilýändigi aýdylýar. Taryh sahypasynda öňki skanirlemeleriň saklanýandygy, bu bolsa audit netijelerini soňra deňeşdirmäge mümkinçilik berýändigi görkezilýär.",
            "Goragda aýratyn nygtalmaly zat programmasynyň çäklendirmelerini hem dogry düşündirmegidir. Telefon ýa-da IoT enjam torda görünmese, munuň sebäbi programma hökmany suratda işlemeýär diýen many bermeýär. AP isolation, guest network, firewall, privat MAC, ICMP jogaplarynyň öçürilmegi ýaly sebäpler netijä täsir edip biler. Programma bu ýagdaýlary loglar we maslahatlar arkaly düşündirýär. Şeýle çemeleşme taslamanyň ylmy we amaly taýdan dogruçyl görkezilmegine kömek edýär.",
            "Şeýlelikde, bäşinji bölümde beýan edilen synaglar programma toplumynyň esasy wezipelerini ýerine ýetirýändigini tassyklaýar: dogry adapter saýlamak, lokal torda gurluşlary tapmak, portlary barlamak, töwekgelçiligi bahalandyrmak, netijeleri saklamak we hasabat taýýarlamak. Goşmaça ssenariýalar we dogrylyk barlaglary bolsa programmasynyň diňe bir görkezme üçin däl, hakyky lokal tor şertlerinde hem ulanylyp bilinjekdigini görkezýär.",
        ],
    ),
]


CONTENT_LINES = {
    "3.5. Hasabat döretmek we duýduryş ulgamynyň durmuşa geçirilmegi": [
        "3.6. Tor adapterlerini kesgitlemek we IP aralygyny awtomatik saýlamak",
        "3.7. Gurluşlary tapmagyň birleşdirilen usullary",
        "3.8. Programma modulynyň giňeldilmegi we goldaw mümkinçilikleri",
    ],
    "4.5. Interfeýsiň diplom goragynda görkezilişi": [
        "4.6. Interfeýsde maglumatlaryň okalmagy we ulanyjy hereketleriniň tertibi",
        "4.7. Netijeleriň tablisa görnüşinde berilmegi we sahypalaryň tapawutlandyrylmagy",
        "4.8. Ulanyjy tejribesini ýokarlandyrmak üçin kabul edilen çözgütler",
    ],
    "5.5. Öndürijilik, takyklyk we çäklendirmeler boýunça netije": [
        "5.6. Lokal torlarda goşmaça synag ssenariýalary",
        "5.7. Tapylan gurluşlaryň dogrulygyny barlamak usuly",
        "5.8. Synag netijeleriniň diplom goragynda görkezilişi",
    ],
}


def word_count(text: str) -> int:
    return len(re.findall(r"\w+", text, flags=re.UNICODE))


def insert_paragraph_after(doc: Document, element, text: str = "", style: str | None = None, align: WD_ALIGN_PARAGRAPH | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    if align is not None:
        paragraph.alignment = align
    if style and style.startswith("Heading"):
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(13 if style == "Heading 2" else 14)
    else:
        paragraph.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(28)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def insert_paragraph_before(doc: Document, element, text: str = "", style: str | None = None, align: WD_ALIGN_PARAGRAPH | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    if align is not None:
        paragraph.alignment = align
    if style and style.startswith("Heading"):
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(13 if style == "Heading 2" else 14)
        paragraph.paragraph_format.first_line_indent = Pt(0)
    else:
        paragraph.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(28)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def insert_table_before(doc: Document, element, rows: list[list[str]]) -> tuple[Table, object]:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.bold = r_idx == 0
    element.addprevious(table._tbl)
    return table, table._tbl


def find_paragraph(doc: Document, text: str, after_index: int = 0) -> Paragraph:
    for i, paragraph in enumerate(doc.paragraphs):
        if i < after_index:
            continue
        if " ".join(paragraph.text.split()) == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def update_mazmuny(doc: Document) -> None:
    existing = {" ".join(p.text.split()) for p in doc.paragraphs[:60]}
    for anchor_text, lines in CONTENT_LINES.items():
        anchor = None
        for i, paragraph in enumerate(doc.paragraphs[:60]):
            if " ".join(paragraph.text.split()) == anchor_text:
                anchor = paragraph
                break
        if anchor is None:
            raise ValueError(f"Mazmuny anchor not found: {anchor_text}")
        current_element = anchor._p
        for line in lines:
            if line in existing:
                continue
            p = insert_paragraph_after(doc, current_element, line, style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.first_line_indent = Pt(0)
            current_element = p._p


def insert_section(doc: Document, before_heading: str, items: list[tuple]) -> None:
    before = find_paragraph(doc, before_heading, after_index=150)
    anchor = before._p
    for item in items:
        heading, paragraphs, *maybe_table = item
        insert_paragraph_before(doc, anchor, heading, style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
        for paragraph_text in paragraphs:
            insert_paragraph_before(doc, anchor, paragraph_text, style="Normal")
        if maybe_table:
            rows, caption = maybe_table[0], maybe_table[1]
            cap = insert_paragraph_before(doc, anchor, caption, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
            cap.paragraph_format.first_line_indent = Pt(0)
            _table, _table_element = insert_table_before(doc, anchor, rows)
        insert_paragraph_before(doc, anchor, "", style="Normal")


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOC_PATH.with_name(f"kakabalowa_before_expand_sections_3_5_{timestamp}.docx")
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))
    update_mazmuny(doc)
    insert_section(doc, "4. Ulanyjy interfeýsiniň işlenip taýýarlanylyşy", SECTION_3)
    insert_section(doc, "5. Programma toplumyny synagdan geçirmek we netijeliligini bahalandyrmak", SECTION_4)
    insert_section(doc, "Zähmeti goramak we ekologik howpsuzlyk boýunça hasaplamalar", SECTION_5)
    doc.save(str(DOC_PATH))

    print(f"Backup: {backup}")
    print(f"Updated: {DOC_PATH}")
    for name, items in [("3", SECTION_3), ("4", SECTION_4), ("5", SECTION_5)]:
        total = 0
        for item in items:
            total += sum(word_count(p) for p in item[1])
        print(f"Added section {name}: approx {total} words")


if __name__ == "__main__":
    main()
