from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAIN_FILE = Path("kakabalowa.docx")
BACKUP_FILE = Path("kakabalowa_before_2_3.docx")


def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(paragraph, 14, True)
    return paragraph


def add_normal(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph, 14)
    return paragraph


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def build_section_doc():
    doc = Document()
    for style_name in ["Normal", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        if style_name.startswith("Heading"):
            style.font.bold = True

    add_heading(doc, "2.3. Maglumatlar bazasynyň taslanylyşy")

    paragraphs = [
        (
            "IoT Security Scanner programma toplumynda skanirleme netijelerini diňe ekranda görkezmek "
            "ýeterlik däldir. Administrator ýa-da diplom işi boýunça synag geçirýän ulanyjy öňki "
            "skanirlemeleriň netijelerini gaýtadan görüp bilmeli, tapylan gurluşlaryň sanyny deňeşdirmeli "
            "we zerur bolan ýagdaýynda hasabat taýýarlamaly. Şonuň üçin programma maglumatlary lokal "
            "maglumatlar bazasynda saklamak mümkinçiligi bilen taslanyldy."
        ),
        (
            "Maglumatlar bazasy hökmünde SQLite saýlanyldy. SQLite aýratyn serwer talap etmeýär we "
            "maglumatlary iot_security.db atly bir faýlyň içinde saklaýar. Bu çözgüt desktop programma "
            "üçin amatlydyr, sebäbi programma ulanyjynyň kompýuterinde özbaşdak işleýär. Maglumatlar "
            "bazasy programma ilkinji gezek işe girizilende ýa-da skanirleme netijeleri saklananda "
            "awtomatik döredilýär."
        ),
        (
            "Taslamada esasy maglumatlar scans atly tablisa ýazylýar. Bu tablisa ady boýunça bir skanirleme "
            "ýazgysy ýaly görünse-de, onuň içinde her bir tapylan gurluş aýratyn setir hökmünde saklanýar. "
            "Bir skanirleme wagtynda tapylan ähli gurluşlar umumy timestamp bahasy bilen birleşdirilýär. "
            "Şeýlelikde, programma soňra şol wagt belgisini ulanyp, bir skanirleme sessiýasyna degişli "
            "ähli gurluşlary toparlap görkezýär."
        ),
        (
            "Scans tablisasynyň esasy meýdanlary gurluşyň IP salgysyny, adyny, MAC salgysyny, açyk portlaryny, "
            "gowşaklyk ýagdaýyny, töwekgelçilik derejesini, maslahatyny, tapylan usulyny, gurluş görnüşini, "
            "öndürijisini, ulanylan adapteri we skanirlenen tor aralygyny saklaýar. Şeýle maglumatlar diňe "
            "tehniki hasaba alyş üçin däl, eýsem interfeýsde netijeleri düşündirmek, statistika çykarmak we "
            "PDF hasabat taýýarlamak üçin hem ulanylýar."
        ),
        (
            "Maglumatlar bazasy bilen işlemek database.py modulynyň içinde jemlenendir. init_db funksiýasy "
            "bazany we scans tablisasyny döredýär. migrate_scans_table funksiýasy programma täze wersiýa "
            "geçende ýetmeýän sütünleri awtomatik goşmaga mümkinçilik berýär. Bu aýratynlyk möhümdir, sebäbi "
            "programma ösdürilende discovery_method, device_type, vendor, adapter we network ýaly täze "
            "meýdanlar goşuldy."
        ),
        (
            "save_scan funksiýasy skanirleme tamamlanandan soň tapylan gurluşlaryň sanawyny maglumatlar "
            "bazasynda saklaýar. Her gurluş üçin portlar sanaw görnüşinden tekst görnüşine öwrülýär, sebäbi "
            "SQLite içinde bu maglumat ýönekeý setir hökmünde saklanýar. load_history funksiýasy bolsa "
            "saklanan netijeleri gaýtadan okap, olary timestamp boýunça toparlaýar. Şu funksiýanyň kömegi "
            "bilen Taryh sahypasy öňki skanirlemeleri görkezýär."
        ),
        (
            "get_scan_stats funksiýasy umumy statistikany almak üçin ulanylýar. Ol geçirilen skanirlemeleriň "
            "sanyny, tapylan gurluşlaryň umumy sanyny, gowşak gurluşlaryň sanyny we iň soňky skanirleme "
            "wagtyny hasaplaýar. Bu maglumatlar Panel sahypasynda ulanyja umumy ýagdaýy görkezmek üçin "
            "ulanylýar. clear_history funksiýasy bolsa zerur bolan ýagdaýynda skanirleme taryhyny arassalamaga "
            "mümkinçilik berýär."
        ),
        (
            "Şeýle maglumatlar bazasy taslamanyň amaly ähmiýetini ýokarlandyrýar. Ulanyjy bir gezeklik "
            "skanirleme bilen çäklenmän, öňki netijeleri saklap, olary soňra deňeşdirip bilýär. Mysal üçin, "
            "birinji skanirlemede Telnet porty açyk enjam tapylsa, düzedişden soň ikinji skanirleme geçirilip, "
            "portuň ýapylandygy ýa-da töwekgelçilik derejesiniň peselendigi barlanyp bilner."
        ),
    ]
    for text in paragraphs:
        add_normal(doc, text)

    caption = doc.add_paragraph("Tablisa 2.3 - scans tablisasynyň meýdanlary")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption, 14)

    headers = ["Meýdan", "Maglumat görnüşi", "Mazmuny"]
    rows = [
        ["id", "INTEGER", "Her setiriň awtomatik döredilýän aýratyn belgisi"],
        ["timestamp", "TEXT", "Skanirleme geçirilen wagty; bir sessiýadaky gurluşlary toparlaýar"],
        ["ip", "TEXT", "Tapylan gurluşyň IP salgysy"],
        ["name", "TEXT", "Gurluşyň host ady ýa-da IP boýunça görkezilýän ady"],
        ["mac", "TEXT", "Gurluşyň MAC salgysy, eger kesgitlenip bilinýän bolsa"],
        ["ports", "TEXT", "Açyk portlaryň vergül bilen ýazylan sanawy"],
        ["vulnerable", "BOOLEAN", "Gurluşda gowşaklyk alamatynyň bardygyny görkezýär"],
        ["risk_level", "TEXT", "Töwekgelçilik derejesi: pes, orta ýa-da ýokary"],
        ["recommendation", "TEXT", "Ulanyja berilýän howpsuzlyk maslahaty"],
        ["discovery_method", "TEXT", "Gurluşyň haýsy usul bilen tapylandygy: Nmap, ARP, Ping, Gateway"],
        ["device_type", "TEXT", "Gurluşyň görnüşi: Router, PC, Telefon/Unknown, IoT kandidat we ş.m."],
        ["vendor", "TEXT", "MAC salgysy boýunça kesgitlenen öndüriji ýa-da näbelli baha"],
        ["adapter", "TEXT", "Skanirleme geçirilen tor adapteriniň ady"],
        ["network", "TEXT", "Skanirlenen IP aralygy, mysal üçin 192.168.1.0/24"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table)

    add_normal(
        doc,
        "Maglumatlar bazasynyň şeýle gurluşy programma üçin ýeterlik derejede ýönekeý we şol bir wagtyň "
        "özünde funksionaldyr. Bir tablisa arkaly hem skanirleme sessiýalary, hem tapylan gurluşlar, hem "
        "töwekgelçilik maglumatlary saklanýar. Bu çemeleşme diplom taslamasynyň göwrümi üçin amatly bolup, "
        "programma interfeýsiniň Panel, Netijeler we Taryh sahypalaryny maglumat bilen üpjün edýär."
    )

    return doc


def insert_before_ecology(main_doc, section_doc):
    target = None
    for paragraph in main_doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Z") and "ekologik" in text and "hasaplamalar" in text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find insertion point before ecology/economic section")

    body = target._p.getparent()
    index = body.index(target._p)
    for element in list(section_doc.element.body):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(index, deepcopy(element))
        index += 1


def fix_heading_style(main_doc):
    for paragraph in main_doc.paragraphs:
        if paragraph.text.strip().startswith("2.3."):
            paragraph.style = "Heading 2"
            set_font(paragraph, 14, True)
            break


def main():
    if not MAIN_FILE.exists():
        raise FileNotFoundError(MAIN_FILE)

    main_doc = Document(str(MAIN_FILE))
    if any(p.text.strip().startswith("2.3.") for p in main_doc.paragraphs):
        print("Section 2.3 already exists; no changes made.")
        return

    shutil.copy2(MAIN_FILE, BACKUP_FILE)
    section_doc = build_section_doc()
    insert_before_ecology(main_doc, section_doc)
    fix_heading_style(main_doc)
    main_doc.save(str(MAIN_FILE))
    print(f"Updated {MAIN_FILE}")
    print(f"Backup saved as {BACKUP_FILE}")


if __name__ == "__main__":
    main()
